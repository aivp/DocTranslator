import threading
import subprocess
from docx import Document
from docx.oxml.ns import qn
from . import to_translate
from . import common
import os
import time
import datetime
import zipfile
import xml.etree.ElementTree as ET
from threading import Lock
from concurrent.futures import ThreadPoolExecutor, as_completed
import re
import tempfile
import shutil
import logging
from typing import List, Dict, Any, Optional
from ..utils.word_run_optimizer import SafeRunMerger, quick_optimize
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement

# 配置日志记录器
logger = logging.getLogger(__name__)

# 线程安全打印锁
print_lock = Lock()

def cleanup_temp_file(temp_path: str):
    """清理临时文件"""
    try:
        if temp_path and os.path.exists(temp_path) and 'optimized_' in os.path.basename(temp_path):
            # 删除临时文件
            os.remove(temp_path)
            # 删除临时目录
            temp_dir = os.path.dirname(temp_path)
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir, ignore_errors=True)
    except Exception as e:
        logger.error(f"清理临时文件失败: {str(e)}")

# 特殊符号和数学符号的正则表达式
SPECIAL_SYMBOLS_PATTERN = re.compile(
    r'^[★☆♥♦♣♠♀♂☼☾☽♪♫♬☑☒✓✔✕✖✗✘⊕⊗∞∑∏πΠ±×÷√∛∜∫∮∇∂∆∏∑√∝∞∟∠∡∢∣∤∥∦∧∨∩∪∫∬∭∮∯∰∱∲∳∴∵∶∷∸∹∺∻∼∽∾∿≀≁≂≃≄≅≆≇≈≉≊≋≌≍≎≏≐≑≒≓≔≕≖≗≘≙≚≛≜≝≞≟≠≡≢≣≤≥≦≧≨≩≪≫≬≭≮≯≰≱≲≳≴≵≶≷≸≹≺≻≼≽≾≿⊀⊁⊂⊃⊄⊅⊆⊇⊈⊉⊊⊋⊌⊍⊎⊏⊐⊑⊒⊓⊔⊕⊖⊗⊘⊙⊚⊛⊜⊝⊞⊟⊠⊡⊢⊣⊤⊥⊦⊧⊨⊩⊪⊫⊬⊭⊮⊯⊰⊱⊲⊳⊴⊵⊶⊷⊸⊹⊺⊻⊼⊽⊾⊿⋀⋁⋂⋃⋄⋅⋆⋇⋈⋉⋊⋋⋌⋍⋎⋏⋐⋑⋒⋓⋔⋕⋖⋗⋘⋙⋚⋛⋜⋝⋞⋟⋠⋡⋢⋣⋤⋥⋦⋧⋨⋩⋪⋫⋬⋭⋮⋯⋰⋱⋲⋳⋴⋵⋶⋷⋸⋹⋺⋻⋼⋽⋾⋿]+$')

# 纯数字和简单标点的正则表达式
NUMBERS_PATTERN = re.compile(r'^[\d\s\.,\-\+\*\/\(\)\[\]\{\}]+$')
TITLE_NUMBER_PATTERN = re.compile(r'^(\d+(?:\.\d+)*\.?)\s*$')  # 匹配标题编号如 "1."、"1.1"、"1.1.1"等

# 添加目录页码模式
TOC_PAGE_PATTERN = re.compile(r'[\s\-—_]+\d+$')


def check_if_image(run):
    """检查run是否包含图片"""
    namespaces = run.element.nsmap
    drawing = run.element.find('.//w:drawing', namespaces)
    if drawing is not None:
        blip = drawing.find('.//a:blip', {**namespaces, 'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
        if blip is not None:
            return True
    if run.element.find('.//w:object', namespaces) is not None:
        return True
    if run.element.find('.//w:pict', namespaces) is not None:
        return True
    return False

def check_if_textbox(run):
    """检查run是否包含文本框"""
    namespaces = {**run.element.nsmap, 'v': 'urn:schemas-microsoft-com:vml'}
    # 检查 DrawingML 文本框
    drawing = run.element.find('.//w:drawing', namespaces)
    if drawing is not None:
        txbx = drawing.find('.//w:txbxContent', {**namespaces, 'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        if txbx is not None:
            return True
    
    # 检查 VML 文本框
    vtextbox = run.element.find('.//v:textbox', namespaces)
    if vtextbox is not None:
        # VML 文本框内部通常有 w:txbxContent
        txbx = vtextbox.find('.//w:txbxContent', namespaces)
        if txbx is not None:
            return True
    
    return False


def clear_image(paragraph):
    """清除段落中的图片，但保留文本内容"""
    runs_to_remove = []
    for run in paragraph.runs:
        if check_if_image(run):
            runs_to_remove.append(run)
    
    # 从后往前删除，避免索引变化
    for run in reversed(runs_to_remove):
        try:
            run._element.getparent().remove(run._element)
        except:
            pass



def start(trans):
    """主入口函数，处理Word文档翻译"""
    # 硬编码线程数为30，忽略前端传入的配置
    max_threads = 30
    start_time = datetime.datetime.now()

    # ============== 检查是否使用Okapi方案 ==============
    use_okapi = trans.get('use_okapi', True)  # 默认使用Okapi
    
    logger.info(f"🔍 翻译方法选择调试:")
    logger.info(f"  配置的 use_okapi: {use_okapi}")
    logger.info(f"  配置的 threads: {trans.get('threads', '未设置')}")
    logger.info(f"  实际使用线程数: {max_threads}")
    
    if use_okapi:
        logger.info("🔄 使用 Okapi Framework 进行翻译（解决run切割问题）")
        return start_with_okapi(trans, start_time)
    else:
        logger.info("📝 使用传统方法进行翻译")
        return start_traditional(trans, start_time, max_threads)


def start_with_okapi(trans, start_time):
    """使用 Okapi Framework 进行翻译"""
    try:
        # 导入 Okapi 集成模块
        from .okapi_integration import OkapiWordTranslator, verify_okapi_installation
        
        # 验证 Okapi 安装
        if not verify_okapi_installation():
            logger.error("❌ Okapi 安装验证失败，回退到传统方法")
            # 硬编码线程数为30，忽略前端传入的配置
            max_threads = 30
            return start_traditional(trans, start_time, max_threads)
        
        # 如果用户选择了qwen-mt-plus，设置server为qwen
        if trans.get('model') == 'qwen-mt-plus':
            trans['server'] = 'qwen'
            logger.info("✅ 设置翻译服务为 Qwen")
        
        # 预加载术语库
        comparison_id = trans.get('comparison_id')
        if comparison_id:
            logger.info(f"📚 开始预加载术语库: {comparison_id}")
            from .main import get_comparison
            preloaded_terms = get_comparison(comparison_id)
            if preloaded_terms:
                logger.info(f"📚 术语库预加载成功: {len(preloaded_terms)} 个术语")
                # 将预加载的术语库添加到trans中
                trans['preloaded_terms'] = preloaded_terms
            else:
                logger.warning(f"📚 术语库预加载失败: {comparison_id}")
        
        # 创建 Okapi 翻译器
        translator = OkapiWordTranslator()
        logger.info("✅ Okapi 翻译器创建成功")
        
        # 设置翻译服务
        class OkapiTranslationService:
            def __init__(self, trans):
                self.trans = trans
            
            def batch_translate(self, texts, source_lang, target_lang):
                """批量翻译文本 - 使用多线程并行处理，支持术语库筛选"""
                from concurrent.futures import ThreadPoolExecutor, as_completed
                import threading
                
                translated_texts = [None] * len(texts)  # 预分配结果数组
                # 从前端配置获取最大线程数，默认为10
                # 硬编码线程数为30，忽略前端传入的配置
                max_workers = min(30, len(texts))
                
                logger.info(f"开始并行翻译 {len(texts)} 个文本，使用 {max_workers} 个线程（硬编码30）")
                
                # 进度更新相关变量
                completed_count = 0
                total_count = len(texts)
                progress_lock = threading.Lock()
                
                def update_progress():
                    """更新翻译进度"""
                    nonlocal completed_count
                    with progress_lock:
                        completed_count += 1
                        progress_percentage = min((completed_count / total_count) * 100, 100.0)
                        logger.info(f"翻译进度: {completed_count}/{total_count} ({progress_percentage:.1f}%)")
                        
                        # 更新数据库进度
                        try:
                            from .to_translate import db
                            db.execute("update translate set process=%s where id=%s", 
                                     str(format(progress_percentage, '.1f')), 
                                     self.trans['id'])
                            
                            # 如果进度达到100%，立即更新状态为已完成
                            if progress_percentage >= 100.0:
                                from datetime import datetime
                                import pytz
                                end_time = datetime.now(pytz.timezone('Asia/Shanghai'))
                                db.execute(
                                    "update translate set status='done',end_at=%s,process=100 where id=%s",
                                    end_time, self.trans['id']
                                )
                                logger.info("✅ 翻译完成，状态已更新为已完成")
                                
                        except Exception as e:
                            logger.error(f"更新进度失败: {str(e)}")
                
                def translate_single_text(index, text):
                    """翻译单个文本，支持术语库筛选"""
                    try:
                        # 检查是否有术语库配置
                        comparison_id = self.trans.get('comparison_id')
                        if comparison_id:
                            logger.debug(f"文本 {index} 使用术语库筛选: {comparison_id}")
                            
                            # 使用预加载的术语库进行筛选
                            preloaded_terms = self.trans.get('preloaded_terms')
                            if preloaded_terms:
                                # 记录术语库处理开始时间
                                term_start_time = time.time()
                                from .term_filter import optimize_terms_for_api
                                filtered_terms = optimize_terms_for_api(text, preloaded_terms, max_terms=50)
                                term_end_time = time.time()
                                term_duration = term_end_time - term_start_time
                                
                                logger.info(f"📚 术语库筛选用时: {term_duration:.3f}秒, 找到术语数: {len(filtered_terms) if filtered_terms else 0}")
                            else:
                                # 如果没有预加载的术语库，回退到原来的方法
                                logger.warning(f"文本 {index} 没有预加载的术语库，回退到数据库查询")
                                from .main import get_filtered_terms_for_text
                                filtered_terms = get_filtered_terms_for_text(text, comparison_id, max_terms=50)
                            
                            if filtered_terms:
                                logger.debug(f"文本 {index} 使用术语库")
                                # 创建临时翻译配置，包含筛选后的术语库
                                temp_trans = self.trans.copy()
                                temp_trans['filtered_terms'] = filtered_terms
                                translated = to_translate.translate_text(
                                    temp_trans, text, source_lang, target_lang
                                )
                            else:
                                logger.debug(f"文本 {index} 没有找到相关术语")
                                translated = to_translate.translate_text(
                                    self.trans, text, source_lang, target_lang
                                )
                        else:
                            logger.debug(f"文本 {index} 未使用术语库")
                            translated = to_translate.translate_text(
                                self.trans, text, source_lang, target_lang
                            )
                        
                        logger.debug(f"文本 {index} 翻译完成: {text[:50]}... -> {translated[:50]}...")
                        return index, translated, None
                    except Exception as e:
                        logger.error(f"文本 {index} 翻译失败: {e}")
                        return index, text, str(e)  # 失败时保持原文
                
                # 使用线程池并行翻译
                with ThreadPoolExecutor(max_workers=max_workers) as executor:
                    # 提交所有翻译任务
                    future_to_index = {
                        executor.submit(translate_single_text, i, text): i 
                        for i, text in enumerate(texts)
                    }
                    
                    # 收集结果
                    completed_count = 0
                    for future in as_completed(future_to_index):
                        index, translated_text, error = future.result()
                        translated_texts[index] = translated_text
                        completed_count += 1
                        
                        update_progress()
                
                logger.info(f"并行翻译完成，共翻译 {len(texts)} 个文本")
                return translated_texts
        
        translator.set_translation_service(OkapiTranslationService(trans))
        logger.info("✅ 翻译服务设置成功")
        
        # 语言映射：将中文语言名称转换为英文全拼
        def map_language_to_qwen_format(lang_name):
            """将中文语言名称映射为Qwen API需要的英文全拼格式"""
            # 处理空值和None的情况
            if not lang_name or lang_name.strip() == '':
                return 'auto'  # 源语言为空时返回auto
                
            language_mapping = {
                # 中文名称到英文全拼
                '中文': 'Chinese',
                '简体中文': 'Chinese',
                '繁体中文': 'Traditional Chinese',
                '英语': 'English',
                '英文': 'English',
                '俄语': 'Russian',
                '日语': 'Japanese',
                '韩语': 'Korean',
                '西班牙语': 'Spanish',
                '法语': 'French',
                '葡萄牙语': 'Portuguese',
                '德语': 'German',
                '意大利语': 'Italian',
                '泰语': 'Thai',
                '越南语': 'Vietnamese',
                '印度尼西亚语': 'Indonesian',
                '马来语': 'Malay',
                '阿拉伯语': 'Arabic',
                '印地语': 'Hindi',
                '希伯来语': 'Hebrew',
                '缅甸语': 'Burmese',
                '泰米尔语': 'Tamil',
                '乌尔都语': 'Urdu',
                '孟加拉语': 'Bengali',
                '波兰语': 'Polish',
                '荷兰语': 'Dutch',
                '罗马尼亚语': 'Romanian',
                '土耳其语': 'Turkish',
                '高棉语': 'Khmer',
                '老挝语': 'Lao',
                '粤语': 'Cantonese',
                '柬埔寨语': 'Khmer',
                '柬埔寨语（高棉语）': 'Khmer',
                '印尼语/马来语': 'Indonesian',
                '菲律宾语（他加禄语）': 'Tagalog',
                '菲律宾语': 'Tagalog',
                '他加禄语': 'Tagalog',
                # 英文全拼到自身 (确保英文全拼映射到自身)
                'Chinese': 'Chinese',
                'English': 'English',
                'Russian': 'Russian',
                'Japanese': 'Japanese',
                'Korean': 'Korean',
                'Spanish': 'Spanish',
                'French': 'French',
                'Portuguese': 'Portuguese',
                'German': 'German',
                'Italian': 'Italian',
                'Thai': 'Thai',
                'Vietnamese': 'Vietnamese',
                'Indonesian': 'Indonesian',
                'Malay': 'Malay',
                'Arabic': 'Arabic',
                'Hindi': 'Hindi',
                'Hebrew': 'Hebrew',
                'Burmese': 'Burmese',
                'Tamil': 'Tamil',
                'Urdu': 'Urdu',
                'Bengali': 'Bengali',
                'Polish': 'Polish',
                'Dutch': 'Dutch',
                'Romanian': 'Romanian',
                'Turkish': 'Turkish',
                'Khmer': 'Khmer',
                'Lao': 'Lao',
                'Cantonese': 'Cantonese',
                'Tagalog': 'Tagalog',
                # 语种编码到英文全拼
                'zh': 'Chinese',
                'en': 'English',
                'ja': 'Japanese',
                'ko': 'Korean',
                'fr': 'French',
                'de': 'German',
                'es': 'Spanish',
                'ru': 'Russian',
                'it': 'Italian',
                'ar': 'Arabic',
                'th': 'Thai',
                'vi': 'Vietnamese',
                'id': 'Indonesian',
                'ms': 'Malay',
                'tl': 'Tagalog',
                'my': 'Burmese',
                'km': 'Khmer',
                'lo': 'Lao',
                'pt': 'Portuguese',
                'hi': 'Hindi',
                'he': 'Hebrew',
                'ta': 'Tamil',
                'ur': 'Urdu',
                'bn': 'Bengali',
                'pl': 'Polish',
                'nl': 'Dutch',
                'ro': 'Romanian',
                'tr': 'Turkish',
                'yue': 'Cantonese',
            }
            return language_mapping.get(lang_name.strip(), lang_name.strip())
        
        # 获取并映射语言
        source_lang = map_language_to_qwen_format(trans.get('source_lang', ''))
        target_lang = map_language_to_qwen_format(trans.get('target_lang', '英语'))
        
        logger.info(f"🔍 语言映射调试:")
        logger.info(f"  原始源语言: {trans.get('source_lang', 'zh')}")
        logger.info(f"  原始目标语言: {trans.get('target_lang', 'en')}")
        logger.info(f"  映射后源语言: {source_lang}")
        logger.info(f"  映射后目标语言: {target_lang}")
        
        # 执行翻译
        success = translator.translate_document(
            trans['file_path'],
            trans['target_file'],
            source_lang,
            target_lang
        )
        
        if success:
            # 完成处理
            end_time = datetime.datetime.now()
            spend_time = common.display_spend(start_time, end_time)
            
            # 统计翻译的文本数量（这里简化处理）
            text_count = 1  # Okapi 以文档为单位处理
            
            if trans['run_complete']:
                to_translate.complete(trans, text_count, spend_time)
            
            logger.info(f"✅ Okapi 翻译完成，用时: {spend_time}")
            return True
        else:
            logger.error("❌ Okapi 翻译失败，回退到传统方法")
            return start_traditional(trans, start_time, 10)
            
    except Exception as e:
        logger.error(f"❌ Okapi 翻译出错: {e}，回退到传统方法")
        return start_traditional(trans, start_time, 10)


def start_traditional(trans, start_time, max_threads):
    """传统翻译方法（原有逻辑）"""
    # ============== Word文档翻译配置 ==============
    logger.info(f"Word文档翻译：使用 {trans.get('model', 'unknown')} 模型")
    
    # 如果用户选择了qwen-mt-plus，检查服务可用性
    if trans.get('model') == 'qwen-mt-plus':
        try:
            trans['server'] = 'qwen'
            # 建议线程数（Qwen并发已提升到1000次/分钟）
            if max_threads > 30:
                logger.info(f"建议: 当前线程数 {max_threads}，建议设置为 10-30 以获得最佳性能")
            elif max_threads < 5:
                logger.info(f"建议: 当前线程数 {max_threads}，可以适当增加到 10-30 以提升翻译速度")
            from .qwen_translate import check_qwen_availability
            qwen_available, qwen_message = check_qwen_availability()
            logger.info(f"Qwen服务检查: {qwen_message}")
            if not qwen_available:
                logger.warning("警告: Qwen服务不可用，但将继续尝试使用")
        except ImportError:
            logger.warning("警告: Qwen模块未找到，将使用默认翻译服务")
    


    # ============== 译文形式处理 ==============
    trans_type = trans.get('type', 'trans_text_only_inherit')  # 默认继承原版面
    logger.info(f"译文形式: {trans_type}")
    # ===========================================

    # ============== Word文档预处理 ==============
    logger.info("开始Word文档预处理...")
    
    # 创建临时文件路径
    temp_dir = tempfile.mkdtemp()
    base_name = os.path.basename(trans['file_path'])
    optimized_path = os.path.join(temp_dir, f"optimized_{base_name}")
    
    try:
        # 使用word-run-optimizer进行预处理
        stats = quick_optimize(trans['file_path'], optimized_path)
        logger.info(f"Word预处理完成: {stats.get('merged_runs', 0)} 个runs被合并")
        optimized_doc_path = optimized_path
    except Exception as e:
        logger.warning(f"Word预处理失败，使用原文档: {str(e)}")
        optimized_doc_path = trans['file_path']
    # ===========================================
    # optimized_doc_path = trans['file_path'] #启用以跳过预处理

    # 加载Word文档
    try:
        document = Document(optimized_doc_path)
    except Exception as e:
        to_translate.error(trans['id'], f"文档加载失败: {str(e)}")
        # 清理临时文件
        cleanup_temp_file(optimized_doc_path)
        return False

    # 提取需要翻译的文本
    texts = []
    
    try:
        # 直接使用多线程提取，简化配置
        # extract_threads = min(4, max_threads)  # 文本提取使用较少的线程
        extract_threads = max_threads
        extract_content_for_translation(document, trans['file_path'], texts, extract_threads)
    except Exception as e:
        to_translate.error(trans['id'], f"文本提取失败: {str(e)}")
        # 清理临时文件
        cleanup_temp_file(optimized_doc_path)
        return False

    # 过滤掉特殊符号和纯数字
    filtered_texts = []
    for i, item in enumerate(texts):
        if should_translate(item['text']):
            filtered_texts.append(item)
        else:
            # 对于不需要翻译的内容，标记为已完成
            item['complete'] = True
            with print_lock:
                logger.info(f"跳过翻译: {item['text'][:30]}..." if len(
                    item['text']) > 30 else f"跳过翻译: {item['text']}")

    # 多线程翻译
    run_translation(trans, filtered_texts, max_threads)

    # ============== 写入翻译结果 ==============
    # 检查是否启用自适应样式
    use_adaptive_styles = trans.get('adaptive_styles', False)  # 默认不启用
    
    if use_adaptive_styles:
        logger.info("启用样式自适应功能：字体大小和行间距将根据翻译后文本长度自动调整")
        text_count = apply_translations_with_adaptive_styles(document, texts)
    else:
        logger.info("使用原始样式：保持原始字体大小和行间距")
        text_count = apply_translations(document, texts)
    # ===========================================

    # 保存文档
    docx_path = trans['target_file']
    document.save(docx_path)

    # 智能run拼接已经在翻译过程中处理，这里不需要额外处理
    
    # 单个run的空格处理已经在翻译过程中完成，这里不需要额外处理

    # 处理批注等特殊元素
    update_special_elements(docx_path, texts)

    # 更新文档目录 - 使用LibreOffice更新
    try:
        logger.info("开始使用LibreOffice更新文档目录...")
        
        # 检查LibreOffice是否可用
        soffice_path = '/usr/bin/soffice'
        if not os.path.exists(soffice_path):
            logger.warning("LibreOffice未安装，跳过目录更新")
            return
        
        # 创建临时目录
        temp_dir = tempfile.mkdtemp()
        temp_docx = os.path.join(temp_dir, "temp_doc.docx")
        
        try:
            # 复制原文档到临时目录
            shutil.copy2(docx_path, temp_docx)
            
            # 使用LibreOffice的字段更新功能
            logger.info("使用LibreOffice转换更新目录...")
            cmd = [
                soffice_path, 
                '--headless', 
                '--norestore', 
                '--invisible',
                '--convert-to', 'docx:MS Word 2007 XML',
                '--outdir', temp_dir,
                temp_docx
            ]
            
            result = subprocess.run(cmd, capture_output=True, timeout=60)
            
            if result.returncode == 0:
                # 查找生成的文件
                generated_files = [f for f in os.listdir(temp_dir) if f.endswith('.docx') and f != 'temp_doc.docx']
                
                if generated_files:
                    generated_path = os.path.join(temp_dir, generated_files[0])
                    
                    # 备份原文件
                    backup_path = docx_path + ".backup"
                    if os.path.exists(docx_path):
                        shutil.copy2(docx_path, backup_path)
                        logger.info(f"原文件已备份到: {backup_path}")
                    
                    # 替换原文件
                    shutil.move(generated_path, docx_path)
                    logger.info("文档目录更新成功（LibreOffice转换方法）")
                    
                    # 清理备份文件
                    if os.path.exists(backup_path):
                        os.remove(backup_path)
                else:
                    logger.warning("LibreOffice未生成新文件")
            else:
                logger.warning(f"LibreOffice转换失败: {result.stderr.decode()}")
                
        finally:
            # 清理临时目录
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
        
        logger.info("LibreOffice目录更新完成")
        
    except Exception as e:
        logger.error(f"使用LibreOffice更新文档目录时发生错误: {str(e)}")
        logger.info("原始文档保持不变")

    # 完成处理
    end_time = datetime.datetime.now()
    spend_time = common.display_spend(start_time, end_time)
    
    # 清理临时文件
    cleanup_temp_file(optimized_doc_path)
    
    if trans['run_complete']:
        to_translate.complete(trans, text_count, spend_time)
    return True


def should_translate(text):
    """判断文本是否应该被翻译"""
    if not text or not text.strip():
        return False

    # 检查是否为标题编号（如 "1."、"2."、"3."）
    if TITLE_NUMBER_PATTERN.match(text.strip()):
        return True  # 标题编号需要处理，但不翻译
    
    # 检查是否为纯数字和简单标点
    if NUMBERS_PATTERN.match(text.strip()):
        return False

    # 检查是否为特殊符号
    if SPECIAL_SYMBOLS_PATTERN.match(text.strip()):
        return False

    return True


def extract_hyperlink_with_merge(hyperlink, texts):
    """提取超链接内容，使用保守run合并"""
    hyperlink_runs = list(hyperlink.runs)
    
    # 暂时注释掉run合并，让每个超链接保持独立
    # 使用保守的run合并策略
    # merged_runs = conservative_run_merge(hyperlink_runs)
    
    # 直接使用原始run，不合并
    merged_runs = []
    for run in hyperlink_runs:
        if check_text(run.text):
            merged_runs.append({
                'text': run.text,
                'type': 'merged',  # 改为merged类型，这样apply_translations才能处理
                'runs': [run]
            })
    
    for merged_item in merged_runs:
        if not check_text(merged_item['text']):
            continue
        
        texts.append({
            "text": merged_item['text'],
            "type": "merged_run",
            "merged_item": merged_item,
            "complete": False,
            "context_type": "hyperlink"  # 标记为超链接
        })


def extract_paragraph_with_merge(paragraph, texts, context_type, paragraph_index=None, total_paragraphs=None):
    """提取段落内容，使用保守run合并（不添加上下文）"""
    paragraph_runs = list(paragraph.runs)
    
    # 暂时注释掉主标题识别，避免表格中的段落被错误识别
    # 检查是否为第一个段落（通常是主标题）
    # is_main_title = paragraph_index == 0
    is_main_title = False  # 暂时禁用主标题识别
    
    # 暂时注释掉run合并，让每个段落保持独立
    # 使用保守的run合并策略，传递is_main_title参数
    # merged_runs = conservative_run_merge(paragraph_runs, is_main_title=is_main_title)
    
    # 直接使用原始run，不合并
    merged_runs = []
    for run in paragraph_runs:
        if check_text(run.text):
            merged_runs.append({
                'text': run.text,
                'type': 'merged',  # 改为merged类型，这样apply_translations才能处理
                'runs': [run]
            })
    
    # 简化日志输出，只在调试模式下显示详细信息
    if logger.isEnabledFor(logging.DEBUG):
        logger.debug(f"段落处理: 索引{paragraph_index}/{total_paragraphs}, 类型{context_type}, run数量{len(paragraph_runs)}")
    
    # 主标题处理已经被禁用，直接处理所有段落
    for merged_item in merged_runs:
        if not check_text(merged_item['text']):
            continue
        
        # 检查是否包含图片，如果包含则跳过
        has_image = False
        for run in merged_item['runs']:
            if check_if_image(run):
                has_image = True
                break
        
        if has_image:
            continue
        
        # 检查是否为标题编号（如 "1."、"2."、"3."）
        if TITLE_NUMBER_PATTERN.match(merged_item['text'].strip()):
            # 标题编号不翻译，但保留在文档中
            texts.append({
                "text": merged_item['text'],
                "type": "merged_run",
                "merged_item": merged_item,
                "complete": True,  # 标记为已完成，不需要翻译
                "context_type": context_type,
                "is_title_number": True  # 标记为标题编号
            })
            continue
        
        # 过滤纯数字
        if NUMBERS_PATTERN.match(merged_item['text'].strip()):
            continue
        
        # 检查是否可能是目录项（以 -数字 或 —数字 结尾）
        match = TOC_PAGE_PATTERN.search(merged_item['text'])
        if match and context_type == "body":  # 假设目录在body
            page_num = match.group(0)
            # 使用strip()去除前后空格，确保文本完整性
            main_text = merged_item['text'][:match.start()].strip()
            
            if main_text and should_translate(main_text):
                # 创建文本项
                texts.append({
                    "text": main_text,
                    "type": "merged_run",
                    "merged_item": merged_item,  # 共享runs，翻译后需小心替换
                    "complete": False,
                    "context_type": context_type,
                    "is_toc_text": True,
                    "original_page_num": page_num
                })
            
                # 创建页码项（不翻译）
                if should_translate(page_num):  # 额外检查，但纯数字应False
                    texts.append({
                        "text": page_num,
                        "type": "merged_run",
                        "merged_item": {'text': page_num, 'runs': merged_item['runs'][-1:], 'type': 'single'},  # 取最后一个run作为页码
                        "complete": False,  # 如果需要翻译
                        "context_type": context_type,
                        "is_toc_page": True
                    })
                else:
                    texts.append({
                        "text": page_num,
                        "type": "merged_run",
                        "merged_item": {'text': page_num, 'runs': merged_item['runs'][-1:], 'type': 'single'},
                        "complete": True,
                        "context_type": context_type,
                        "is_toc_page": True
                    })
        else:
            # 如果是第一个段落，标记为主标题
            if is_main_title:
                context_type = "main_title"
                logger.info(f"检测到主标题段落: '{merged_item['text']}'")
            
            texts.append({
                "text": merged_item['text'],
                "type": "merged_run",
                "merged_item": merged_item,
                "complete": False,
                "context_type": context_type,  # 标记类型
                "is_main_title": is_main_title  # 添加主标题标记
            })
    
    if logger.isEnabledFor(logging.DEBUG):
        logger.debug(f"段落处理完成: 索引{paragraph_index}")


def extract_content_for_translation(document, file_path, texts, max_threads=4):
    """提取需要翻译的内容，使用安全的多线程并行处理"""
    
    # 线程安全的文本列表
    from threading import Lock
    texts_lock = Lock()
    
    def add_text_safe(text_item):
        """线程安全地添加文本项"""
        with texts_lock:
            texts.append(text_item)
    
    def process_paragraphs_parallel():
        """并行处理段落（保持顺序）"""
        # 将段落分组，每组包含连续的段落
        paragraph_groups = []
        current_group = []
        
        for i, paragraph in enumerate(document.paragraphs):
            current_group.append((i, paragraph))
            
            # 每10个段落为一组，或者遇到表格时分组
            if len(current_group) >= 10 or (i < len(document.paragraphs) - 1 and 
                hasattr(document.paragraphs[i+1], '_element') and 
                document.paragraphs[i+1]._element.getnext() is not None):
                paragraph_groups.append(current_group)
                current_group = []
        
        if current_group:
            paragraph_groups.append(current_group)
        
        def process_paragraph_group(group):
            """处理一组连续的段落"""
            local_texts = []
            
            for index, paragraph in group:
                # 处理正文段落（跳过超链接run）
                # 直接调用extract_paragraph_with_merge，传递段落索引信息
                extract_paragraph_with_merge(paragraph, local_texts, "body", index, len(document.paragraphs))
                
                # 处理超链接
                for hyperlink in paragraph.hyperlinks:
                    extract_hyperlink_with_merge(hyperlink, local_texts)
                
                # 处理段落中的文本框
                for run in paragraph.runs:
                    if check_if_textbox(run):
                        # 处理 DrawingML
                        drawing_elem = run.element.find('.//w:drawing', run.element.nsmap)
                        if drawing_elem is not None:
                            txbx_elem = drawing_elem.find('.//w:txbxContent', drawing_elem.nsmap)
                            if txbx_elem is not None:
                                for p_elem in txbx_elem:
                                    if p_elem.tag == qn('w:p'):
                                        tb_para = Paragraph(p_elem, paragraph)
                                        extract_paragraph_with_merge(tb_para, local_texts, "textbox", index, len(document.paragraphs))
                        
                        # 处理 VML
                        namespaces = {**run.element.nsmap, 'v': 'urn:schemas-microsoft-com:vml'}
                        vtextbox_elem = run.element.find('.//v:textbox', namespaces)
                        if vtextbox_elem is not None:
                            txbx_elem = vtextbox_elem.find('.//w:txbxContent', namespaces)
                            if txbx_elem is not None:
                                for p_elem in txbx_elem:
                                    if p_elem.tag == qn('w:p'):
                                        tb_para = Paragraph(p_elem, paragraph)
                                        extract_paragraph_with_merge(tb_para, local_texts, "textbox", index, len(document.paragraphs))
            
            # 线程安全地添加到主列表
            for text_item in local_texts:
                add_text_safe(text_item)
        
        # 使用线程池并行处理段落组
        with ThreadPoolExecutor(max_workers=max_threads) as executor:
            futures = [executor.submit(process_paragraph_group, group) 
                      for group in paragraph_groups]
            for future in as_completed(futures):
                future.result()
    
    def process_tables_parallel():
        """并行处理表格（表格之间无依赖）"""
        def process_table(table):
            local_texts = []
            # 使用新的表格处理函数，包含布局调整
            process_table_with_layout_adjustment(table, local_texts)
            
            # 线程安全地添加到主列表
            for text_item in local_texts:
                add_text_safe(text_item)
        
        # 使用线程池并行处理表格
        with ThreadPoolExecutor(max_workers=max_threads) as executor:
            futures = [executor.submit(process_table, table) 
                      for table in document.tables]
            for future in as_completed(futures):
                future.result()
    
    def process_sections_parallel():
        """并行处理页眉页脚（section之间无依赖）"""
        def process_section(section):
            local_texts = []
            
            # 处理页眉
            for paragraph in section.header.paragraphs:
                extract_paragraph_with_merge(paragraph, local_texts, "header", 0, 1)  # 页眉段落索引为0
                # 处理页眉中的超链接
                for hyperlink in paragraph.hyperlinks:
                    extract_hyperlink_with_merge(hyperlink, local_texts)
                # 处理页眉中的文本框
                for run in paragraph.runs:
                    if check_if_textbox(run):
                        # 处理 DrawingML
                        drawing_elem = run.element.find('.//w:drawing', run.element.nsmap)
                        if drawing_elem is not None:
                            txbx_elem = drawing_elem.find('.//w:txbxContent', drawing_elem.nsmap)
                            if txbx_elem is not None:
                                for p_elem in txbx_elem:
                                    if p_elem.tag == qn('w:p'):
                                        tb_para = Paragraph(p_elem, paragraph)
                                        extract_paragraph_with_merge(tb_para, local_texts, "textbox", 0, 1)
                        
                        # 处理 VML
                        namespaces = {**run.element.nsmap, 'v': 'urn:schemas-microsoft-com:vml'}
                        vtextbox_elem = run.element.find('.//v:textbox', namespaces)
                        if vtextbox_elem is not None:
                            txbx_elem = vtextbox_elem.find('.//w:txbxContent', namespaces)
                            if txbx_elem is not None:
                                for p_elem in txbx_elem:
                                    if p_elem.tag == qn('w:p'):
                                        tb_para = Paragraph(p_elem, paragraph)
                                        extract_paragraph_with_merge(tb_para, local_texts, "textbox", 0, 1)

            # 处理页脚
            for paragraph in section.footer.paragraphs:
                extract_paragraph_with_merge(paragraph, local_texts, "footer", 0, 1)  # 页脚段落索引为0
                # 处理页脚中的超链接
                for hyperlink in paragraph.hyperlinks:
                    extract_hyperlink_with_merge(hyperlink, local_texts)
                # 处理页脚中的文本框
                for run in paragraph.runs:
                    if check_if_textbox(run):
                        # 处理 DrawingML
                        drawing_elem = run.element.find('.//w:drawing', run.element.nsmap)
                        if drawing_elem is not None:
                            txbx_elem = drawing_elem.find('.//w:txbxContent', drawing_elem.nsmap)
                            if txbx_elem is not None:
                                for p_elem in txbx_elem:
                                    if p_elem.tag == qn('w:p'):
                                        tb_para = Paragraph(p_elem, paragraph)
                                        extract_paragraph_with_merge(tb_para, local_texts, "textbox", 0, 1)
                        
                        # 处理 VML
                        namespaces = {**run.element.nsmap, 'v': 'urn:schemas-microsoft-com:vml'}
                        vtextbox_elem = run.element.find('.//v:textbox', namespaces)
                        if vtextbox_elem is not None:
                            txbx_elem = vtextbox_elem.find('.//w:txbxContent', namespaces)
                            if txbx_elem is not None:
                                for p_elem in txbx_elem:
                                    if p_elem.tag == qn('w:p'):
                                        tb_para = Paragraph(p_elem, paragraph)
                                        extract_paragraph_with_merge(tb_para, local_texts, "textbox", 0, 1)
            
            # 线程安全地添加到主列表
            for text_item in local_texts:
                add_text_safe(text_item)
        
        # 使用线程池并行处理sections
        with ThreadPoolExecutor(max_workers=max_threads) as executor:
            futures = [executor.submit(process_section, section) 
                      for section in document.sections]
            for future in as_completed(futures):
                future.result()
    
    def process_inline_shapes_parallel():
        """并行处理内嵌形状（inline shapes）"""
        def process_shape(shape):
            local_texts = []
            try:
                if hasattr(shape, 'has_text_frame') and shape.has_text_frame:
                    for paragraph in shape.text_frame.paragraphs:
                        extract_paragraph_with_merge(paragraph, local_texts, "textbox", 0, 1)  # 形状段落索引为0
                        # 处理形状中的超链接
                        for hyperlink in paragraph.hyperlinks:
                            extract_hyperlink_with_merge(hyperlink, local_texts)
            except Exception as e:
                logger.error(f"处理内嵌形状时出错: {str(e)}")
            
            # 线程安全地添加到主列表
            for text_item in local_texts:
                add_text_safe(text_item)
        
        # 使用线程池并行处理形状
        with ThreadPoolExecutor(max_workers=max_threads) as executor:
            futures = [executor.submit(process_shape, shape) 
                       for shape in document.inline_shapes]
            for future in as_completed(futures):
                try:
                    future.result()
                except Exception as e:
                    logger.error(f"内嵌形状线程异常: {str(e)}")
    
    # 按顺序执行各个部分，避免依赖问题
    logger.info("开始安全的多线程文本提取...")
    
    # 1. 并行处理段落（分组处理，保持顺序）
    process_paragraphs_parallel()
    
    # 2. 并行处理表格（表格之间无依赖）
    if document.tables:
        process_tables_parallel()
    
    # 3. 并行处理页眉页脚（section之间无依赖）
    if document.sections:
        process_sections_parallel()
    
    # 4. 处理内嵌形状（inline shapes，可能包含文本框）
    if document.inline_shapes:
        process_inline_shapes_parallel()
    
    # 5. 批注内容（单线程，因为需要ZIP操作）
    if hasattr(document, 'comments') and document.comments:
        extract_comments(file_path, texts)
    
    logger.info(f"文本提取完成，共提取 {len(texts)} 个文本项")


# 已删除单线程版本，统一使用多线程提取


def get_run_format_key(run):
    """获取run的格式标识，用于快速比较"""
    try:
        font_name = run.font.name if run.font.name else None
        font_size = run.font.size if run.font.size else None
        bold = run.bold
        italic = run.italic
        underline = run.underline
        color = run.font.color.rgb if run.font.color.rgb else None
        
        return (font_name, font_size, bold, italic, underline, color)
    except:
        return None

def are_runs_compatible(run1, run2, is_main_title=False):
    """检查两个run是否兼容，可以合并"""
    
    # 如果是主标题，放宽格式兼容性检查
    if is_main_title:
        # 主标题只需要检查字体大小是否相同
        try:
            size1 = run1.font.size.pt if run1.font.size else None
            size2 = run2.font.size.pt if run2.font.size else None
            
            # 如果字体大小相同或都为空，则认为兼容
            if size1 == size2 or (size1 is None and size2 is None):
                return True
            else:
                return False
        except:
            # 如果获取字体大小失败，默认兼容
            return True
    
    # 原有的兼容性检查逻辑
    try:
        # 获取格式信息
        format1 = get_run_format_key(run1)
        format2 = get_run_format_key(run2)
        
        # 如果格式完全相同，则兼容
        if format1 == format2:
            return True
        
        # 如果格式差异很小（比如只是字体名称不同），也可以合并
        # 这里可以根据需要调整兼容性判断的严格程度
        
        return False
    except:
        # 如果获取格式信息失败，默认不兼容
        return False


def conservative_run_merge(paragraph_runs, max_merge_length=1000, is_main_title=False):
    """保守的run合并策略"""
    
    merged = []
    current_group = []
    current_length = 0
    original_count = len([r for r in paragraph_runs if check_text(r.text)])
    merged_count = 0
    
    # 简化调试信息，只在调试模式下显示
    if logger.isEnabledFor(logging.DEBUG):
        logger.debug(f"conservative_run_merge: 输入{len(paragraph_runs)}个run, 有效{original_count}个, 主标题{is_main_title}")
    
    # 如果是主标题，强制合并所有run
    if is_main_title and original_count > 1:
        if logger.isEnabledFor(logging.DEBUG):
            logger.debug(f"主标题段落，强制合并所有{original_count}个run")
        
        # 收集所有有效run
        all_runs = [r for r in paragraph_runs if check_text(r.text)]
        
        # 创建合并后的项
        merged_item = merge_compatible_runs(all_runs)
        merged.append(merged_item)
        
        if logger.isEnabledFor(logging.DEBUG):
            logger.debug(f"conservative_run_merge完成: 合并后{len(merged)}个项")
        
        return merged
    
    # 原有的合并逻辑
    for i, run in enumerate(paragraph_runs):
        # 跳过包含图片的run
        if check_if_image(run):
            if logger.isEnabledFor(logging.DEBUG):
                logger.debug(f"Run {i}: '{run.text}' - 包含图片，跳过")
            # 保存当前组
            if current_group:
                merged.append(merge_compatible_runs(current_group))
                if len(current_group) > 1:
                    merged_count += len(current_group) - 1
                current_group = []
                current_length = 0
            
            # 图片run单独处理，但不添加到翻译列表
            continue
        
        if not check_text(run.text):
            if logger.isEnabledFor(logging.DEBUG):
                logger.debug(f"Run {i}: '{run.text}' - 无效文本，跳过")
            continue
        
        if logger.isEnabledFor(logging.DEBUG):
            logger.debug(f"Run {i}: '{run.text}' - 长度: {len(run.text)}, 当前组长度: {current_length}")
        
        # 只合并较短的run（通常是空格、标点、短词、短语）
        if len(run.text) <= 20 and current_length < max_merge_length:
            # 检查格式兼容性，传递is_main_title参数
            if not current_group or are_runs_compatible(current_group[-1], run, is_main_title):
                current_group.append(run)
                current_length += len(run.text)
                if logger.isEnabledFor(logging.DEBUG):
                    logger.debug(f"  添加到当前组，当前组: {[r.text for r in current_group]}")
                continue
            else:
                if logger.isEnabledFor(logging.DEBUG):
                    logger.debug(f"  格式不兼容，不合并")
        
        # 保存当前组
        if current_group:
            if logger.isEnabledFor(logging.DEBUG):
                logger.debug(f"保存当前组: {[r.text for r in current_group]}")
            merged.append(merge_compatible_runs(current_group))
            if len(current_group) > 1:
                merged_count += len(current_group) - 1  # 记录合并的run数量
            current_group = []
            current_length = 0
        
        # 当前run单独处理
        if logger.isEnabledFor(logging.DEBUG):
            logger.debug(f"Run {i} 单独处理: '{run.text}'")
        merged.append({
            'text': run.text,
            'runs': [run],
            'type': 'single'
        })
    
    if current_group:
        if logger.isEnabledFor(logging.DEBUG):
            logger.debug(f"保存最后的组: {[r.text for r in current_group]}")
        merged.append(merge_compatible_runs(current_group))
        if len(current_group) > 1:
            merged_count += len(current_group) - 1
    
    # 打印合并统计信息
    if merged_count > 0:
        with print_lock:
            logger.info(f"Run合并优化: 原始{original_count}个run -> 合并后{len(merged)}个，减少了{merged_count}个API调用")
    
    if logger.isEnabledFor(logging.DEBUG):
        logger.debug(f"conservative_run_merge完成: 合并后{len(merged)}个项")
    
    return merged


def smart_space_preservation(original_runs, translated_text):
    """智能保持原始run之间的空格关系"""
    if not original_runs:
        return translated_text
    
    # 分析原始run之间的空格关系
    space_pattern = []
    for i, run in enumerate(original_runs):
        if i > 0:
            prev_run = original_runs[i-1]
            # 检查是否有空格分隔
            if prev_run.text.endswith(' ') or run.text.startswith(' '):
                space_pattern.append(True)  # 有空格分隔
            else:
                space_pattern.append(False)  # 无空格分隔
        space_pattern.append(False)  # 第一个run
    
    # 如果原始run之间没有空格分隔，翻译后也不添加
    if not any(space_pattern):
        return translated_text
    
    # 如果有空格分隔，在适当位置添加空格
    # 这里可以根据需要实现更复杂的空格处理逻辑
    return translated_text


def merge_compatible_runs(run_group):
    """合并兼容的run组，保持原始格式（不添加空格）"""
    
    # 直接连接文本，不添加额外空格
    # 空格处理将在翻译后的分配阶段进行
    merged_text = "".join(run.text for run in run_group)
    
    return {
        'text': merged_text,
        'runs': run_group,
        'type': 'merged'
    }


def extract_paragraph_with_context(paragraph, texts, context_type):
    """为正文段落添加上下文，使用保守的run合并策略"""
    
    # 过滤掉超链接的run，避免重复处理
    paragraph_runs = []
    for run in paragraph.runs:
        # 检查run是否属于超链接
        is_hyperlink_run = False
        for hyperlink in paragraph.hyperlinks:
            if run in hyperlink.runs:
                is_hyperlink_run = True
                break
        
        if not is_hyperlink_run:
            paragraph_runs.append(run)
    
    # 暂时注释掉run合并，让每个段落保持独立
    # 使用保守的run合并策略
    # merged_runs = conservative_run_merge(paragraph_runs)
    
    # 直接使用原始run，不合并
    merged_runs = []
    for run in paragraph_runs:
        if check_text(run.text):
            merged_runs.append({
                'text': run.text,
                'type': 'merged',  # 改为merged类型，这样apply_translations才能处理
                'runs': [run]
            })
    
    for i, merged_item in enumerate(merged_runs):
        if not check_text(merged_item['text']):
            continue
        
        # 检查是否包含图片，如果包含图片则跳过，避免在图片后添加空行
        has_image = False
        for run in merged_item['runs']:
            if check_if_image(run):
                has_image = True
                break
        
        if has_image:
            continue
        
        # 判断是否是段落边界
        is_start = i == 0
        is_end = i == len(merged_runs) - 1
        
        # 根据边界情况获取上下文
        if is_start:
            # 段落开头：只要后文
            context_before = ""
            context_after = get_context_after_merged(merged_runs, i, 2)
        elif is_end:
            # 段落结尾：只要前文
            context_before = get_context_before_merged(merged_runs, i, 2)
            context_after = ""
        else:
            # 正常情况：前后文都要
            context_before = get_context_before_merged(merged_runs, i, 2)
            context_after = get_context_after_merged(merged_runs, i, 2)
        
        # 构建带上下文的翻译文本
        context_parts = []
        if context_before:
            context_parts.append(f"[前文: {context_before}]")
        context_parts.append(merged_item['text'])
        if context_after:
            context_parts.append(f"[后文: {context_after}]")
        
        # 使用空字符串连接，避免添加额外空格
        context_text = "".join(context_parts)
        
        texts.append({
            "text": merged_item['text'],           # 原始文本
            "context_text": context_text,  # 带上下文的文本
            "type": "merged_run",
            "merged_item": merged_item,  # 包含runs列表
            "complete": False,
            "context_type": context_type,  # 标记为正文
            "is_boundary": is_start or is_end
        })


def get_context_before(runs, current_index, window_size):
    """获取前文上下文"""
    start_index = max(0, current_index - window_size)
    context_runs = runs[start_index:current_index]
    
    # 收集文本
    context_texts = []
    for run in context_runs:
        if check_text(run.text):
            context_texts.append(run.text)
    
    return ''.join(context_texts)


def get_context_after(runs, current_index, window_size):
    """获取后文上下文"""
    end_index = min(len(runs), current_index + window_size + 1)
    context_runs = runs[current_index + 1:end_index]
    
    # 收集文本
    context_texts = []
    for run in context_runs:
        if check_text(run.text):
            context_texts.append(run.text)
    
    return ''.join(context_texts)


def get_context_before_merged(merged_runs, current_index, window_size):
    """获取合并run的前文上下文，保持原始空格格式"""
    start_index = max(0, current_index - window_size)
    context_items = merged_runs[start_index:current_index]
    
    # 收集文本，保持原始格式
    context_texts = []
    for item in context_items:
        if check_text(item['text']):
            context_texts.append(item['text'])
    
    # 使用空字符串连接，保持原始空格
    return ''.join(context_texts)


def get_context_after_merged(merged_runs, current_index, window_size):
    """获取合并run的后文上下文，保持原始空格格式"""
    end_index = min(len(merged_runs), current_index + window_size + 1)
    context_items = merged_runs[current_index + 1:end_index]
    
    # 收集文本，保持原始格式
    context_texts = []
    for item in context_items:
        if check_text(item['text']):
            context_texts.append(item['text'])
    
    # 使用空字符串连接，保持原始空格
    return ''.join(context_texts)


def extract_comments(file_path, texts):
    """提取文档中的批注"""
    try:
        with zipfile.ZipFile(file_path, 'r') as docx:
            if 'word/comments.xml' in docx.namelist():
                with docx.open('word/comments.xml') as f:
                    tree = ET.parse(f)
                    root = tree.getroot()

                    # 定义命名空间
                    namespaces = {
                        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                    }

                    # 查找所有批注
                    for comment in root.findall('.//w:comment', namespaces) or root.findall(
                            './/{*}comment'):
                        # 尝试不同方式获取ID
                        comment_id = None
                        for attr_name in [
                            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id',
                            'id']:
                            if attr_name in comment.attrib:
                                comment_id = comment.attrib[attr_name]
                                break

                        if not comment_id:
                            continue

                        # 收集批注中的所有文本
                        text_elements = comment.findall('.//w:t', namespaces) or comment.findall(
                            './/{*}t')
                        for t_elem in text_elements:
                            if t_elem.text and check_text(t_elem.text):
                                texts.append({
                                    "text": t_elem.text,
                                    "type": "comment",
                                    "comment_id": comment_id,
                                    "complete": False,
                                    "original_text": t_elem.text  # 保存原始文本用于匹配
                                })
    except Exception as e:
        logger.error(f"提取批注时出错: {str(e)}")


def run_translation(trans, texts, max_threads):
    """执行多线程翻译"""
    if not texts:
        logger.info("没有需要翻译的内容")
        return

    event = threading.Event()
    
    with print_lock:
        logger.info(f"开始翻译 {len(texts)} 个文本片段")
        logger.info(f"翻译服务: {trans.get('server', 'unknown')}")  # 确认使用的翻译服务

    # 进度更新相关变量
    completed_count = 0
    total_count = len(texts)
    progress_lock = threading.Lock()
    
    def update_progress():
        """更新翻译进度"""
        nonlocal completed_count
        with progress_lock:
            completed_count += 1
            progress_percentage = min((completed_count / total_count) * 100, 100.0)
            logger.info(f"翻译进度: {completed_count}/{total_count} ({progress_percentage:.1f}%)")
            
            # 更新数据库进度
            try:
                from .to_translate import db
                db.execute("update translate set process=%s where id=%s", 
                         str(format(progress_percentage, '.1f')), 
                         trans['id'])
                
                # 如果进度达到100%，立即更新状态为已完成
                if progress_percentage >= 100.0:
                    from datetime import datetime
                    import pytz
                    end_time = datetime.now(pytz.timezone('Asia/Shanghai'))
                    db.execute(
                        "update translate set status='done',end_at=%s,process=100 where id=%s",
                        end_time, trans['id']
                    )
                    logger.info("✅ 翻译完成，状态已更新为已完成")
                    
            except Exception as e:
                logger.error(f"更新进度失败: {str(e)}")

    # 使用线程池执行翻译任务
    with ThreadPoolExecutor(max_workers=max_threads) as executor:
        # 提交所有翻译任务
        futures = []
        for i in range(len(texts)):
            future = executor.submit(to_translate.get, trans, event, texts, i)
            futures.append(future)
            with print_lock:
                logger.info(f"提交翻译任务 {i}")
        
        # 等待所有任务完成
        for future in as_completed(futures):
            try:
                future.result()  # 获取结果，如果有异常会抛出
                update_progress()  # 更新进度
            except Exception as e:
                with print_lock:
                    logger.error(f"翻译任务执行异常: {str(e)}")
                if not event.is_set():
                    event.set()  # 设置事件，通知其他线程停止

    with print_lock:
        logger.info("所有翻译任务已完成")


def calculate_adaptive_font_size(original_text, translated_text, original_font_size):
    """根据翻译后文本长度计算自适应字体大小"""
    try:
        if not original_font_size:
            return None
        
        # 计算文本长度比例，使用strip()去除前后空格确保准确计算
        original_length = len(original_text.strip())
        translated_length = len(translated_text.strip())
        
        if original_length == 0:
            return original_font_size
        
        # 计算长度比例
        length_ratio = translated_length / original_length
        
        # 如果翻译后文本变长，适当缩小字体
        if length_ratio > 1.2:  # 文本长度增加超过20%
            # 根据长度比例计算新的字体大小，但不要小于原始大小的70%
            new_size = max(original_font_size * 0.7, original_font_size / length_ratio)
            return int(new_size)
        elif length_ratio < 0.8:  # 文本长度减少超过20%
            # 如果文本变短，可以适当增大字体，但不要超过原始大小的120%
            new_size = min(original_font_size * 1.2, original_font_size / length_ratio)
            return int(new_size)
        else:
            # 长度变化不大，保持原始字体大小
            return original_font_size
            
    except Exception as e:
        logger.error(f"计算自适应字体大小失败: {str(e)}")
        return original_font_size


def calculate_adaptive_line_spacing(original_text, translated_text, original_line_spacing):
    """根据翻译后文本长度计算自适应行间距"""
    try:
        # 计算文本行数变化
        original_lines = len(original_text.split('\n'))
        translated_lines = len(translated_text.split('\n'))
        
        if original_lines == 0:
            return original_line_spacing
        
        # 计算行数比例
        line_ratio = translated_lines / original_lines
        
        # 如果行数增加，适当增加行间距
        if line_ratio > 1.1:  # 行数增加超过10%
            # 增加行间距，但不要超过原始行间距的150%
            new_spacing = min(original_line_spacing * 1.5, original_line_spacing * line_ratio)
            return new_spacing
        elif line_ratio < 0.9:  # 行数减少超过10%
            # 如果行数减少，可以适当减少行间距，但不要少于原始行间距的80%
            new_spacing = max(original_line_spacing * 0.8, original_line_spacing * line_ratio)
            return new_spacing
        else:
            # 行数变化不大，保持原始行间距
            return original_line_spacing
            
    except Exception as e:
        logger.error(f"计算自适应行间距失败: {str(e)}")
        return original_line_spacing


def apply_adaptive_styles(run, original_text, translated_text, context_type=None):
    """应用自适应样式到run"""
    try:
        # 获取原始字体大小
        original_font_size = None
        if run.font.size:
            try:
                original_font_size = run.font.size.pt
            except:
                original_font_size = None
        
        # 计算自适应字体大小
        if original_font_size:
            adaptive_font_size = calculate_adaptive_font_size(original_text, translated_text, original_font_size)
            if adaptive_font_size and adaptive_font_size != original_font_size:
                # 如果是文本框，进一步缩小20%
                if context_type == 'textbox':
                    adaptive_font_size = int(adaptive_font_size * 0.8)
                from docx.shared import Pt
                run.font.size = Pt(adaptive_font_size)
                logger.info(f"字体大小自适应: {original_font_size}pt -> {adaptive_font_size}pt")
            elif context_type == 'textbox':
                # 对于文本框，即使长度没变，也缩小20%
                adaptive_font_size = int(original_font_size * 0.8)
                from docx.shared import Pt
                run.font.size = Pt(adaptive_font_size)
                logger.info(f"文本框字体固定缩小: {original_font_size}pt -> {adaptive_font_size}pt")
        
        # 获取段落对象并应用行间距自适应
        try:
            paragraph = run._element.getparent()
            if hasattr(paragraph, 'paragraph_format') and paragraph.paragraph_format:
                # 获取原始行间距
                original_line_spacing = paragraph.paragraph_format.line_spacing
                if original_line_spacing:
                    adaptive_line_spacing = calculate_adaptive_line_spacing(original_text, translated_text, original_line_spacing)
                    if adaptive_line_spacing and adaptive_line_spacing != original_line_spacing:
                        paragraph.paragraph_format.line_spacing = adaptive_line_spacing
                        logger.info(f"行间距自适应: {original_line_spacing} -> {adaptive_line_spacing}")
        except Exception as e:
            logger.error(f"应用行间距自适应失败: {str(e)}")
                
    except Exception as e:
        logger.error(f"应用自适应样式失败: {str(e)}")


def apply_translations_with_adaptive_styles(document, texts):
    """应用翻译结果到文档，同时应用自适应样式"""
    text_count = 0

    for item in texts:
        if not item.get('complete', False):
            continue

        if item['type'] == 'run':
            # 检查run是否包含图片，如果包含则跳过
            if check_if_image(item['run']):
                continue
            
            # 获取原始文本和翻译后文本
            original_text = item['run'].text
            translated_text = item.get('text', original_text)
            
            # 应用翻译结果
            item['run'].text = translated_text
            
            # 应用自适应样式
            apply_adaptive_styles(item['run'], original_text, translated_text)
            
            text_count += 1
        elif item['type'] == 'merged_run':
            # 处理合并的run，需要将翻译结果分配回原始run
            merged_item = item['merged_item']
            translated_text = item.get('text', merged_item['text'])
            original_text = merged_item['text']
            
            # 检查是否包含图片，如果包含则跳过
            has_image = False
            for run in merged_item['runs']:
                if check_if_image(run):
                    has_image = True
                    break
            
            if has_image:
                continue
            
            if merged_item['type'] == 'merged':
                # 合并的run组，需要智能分配翻译结果
                distribute_translation_to_runs_with_adaptive_styles(merged_item['runs'], translated_text, original_text)
            else:
                # 单个run，直接替换
                run = merged_item['runs'][0]
                run.text = translated_text
                apply_adaptive_styles(run, original_text, translated_text)
            
            text_count += 1
        elif item['type'] == 'single_run':
            # 处理单个run，直接应用翻译结果
            logger.info(f"=== 处理single_run类型 ===")
            merged_item = item['merged_item']
            translated_text = item.get('text', merged_item['text'])
            original_text = merged_item['text']
            
            # 检查是否包含图片，如果包含则跳过
            has_image = False
            for run in merged_item['runs']:
                if check_if_image(run):
                    has_image = True
                    break
            
            if has_image:
                continue
            
            # 直接替换run文本
            run = merged_item['runs'][0]
            run.text = translated_text
            logger.info(f"single_run文本替换: '{original_text}' -> '{translated_text}'")
            
            text_count += 1
        elif item['type'] == 'main_title_runs':
            # 处理主标题的多个run，在run之间添加空格
            logger.info(f"=== 主标题run空格处理开始 ===")
            runs = item['runs']
            translated_text = item.get('text', "")
            
            logger.info(f"主标题原文: '{translated_text}'")
            logger.info(f"Run数量: {len(runs)}")
            
            # 将翻译文本按run数量分配，并在run之间添加空格
            words = translated_text.split()
            logger.info(f"翻译文本分割后的单词: {words}")
            
            if len(words) >= len(runs):
                # 平均分配单词到每个run
                words_per_run = len(words) // len(runs)
                remainder = len(words) % len(runs)
                
                current_word_index = 0
                for i, run in enumerate(runs):
                    # 计算当前run应该包含的单词数
                    if i < remainder:
                        run_word_count = words_per_run + 1
                    else:
                        run_word_count = words_per_run
                    
                    # 提取当前run的单词
                    run_words = words[current_word_index:current_word_index + run_word_count]
                    run_text = ' '.join(run_words)
                    
                    # 更新run文本
                    run.text = run_text
                    current_word_index += run_word_count
                    
                    logger.info(f"主标题run {i}: '{run_text}' (分配单词: {run_words})")
                    
                    # 检查是否以标点符号结尾，如果不是就加空格
                    if not run_text.rstrip().endswith(('.,;:!?()[]{}"\'-')):
                        run.text = run_text + ' '
                        logger.info(f"完整标题run {i}不是标点结尾，加结尾空格: '{run_text}' -> '{run.text}'")
                    else:
                        logger.info(f"完整标题run {i}以标点结尾，不加空格: '{run_text}'")
            else:
                # 单词数量少于run数量，简单分配
                logger.info(f"单词数量({len(words)})少于run数量({len(runs)})，进行简单分配")
                for i, run in enumerate(runs):
                    if i < len(words):
                        run.text = words[i]
                        # 检查是否以标点符号结尾，如果不是就加空格
                        if not words[i].rstrip().endswith(('.,;:!?()[]{}"\'-')):
                            run.text = words[i] + ' '
                            logger.info(f"完整标题run {i}不是标点结尾，加结尾空格: '{words[i]}' -> '{run.text}'")
                    else:
                        run.text = ""
                    logger.info(f"完整标题run {i}: '{run.text}'")
            
            logger.info("=== 主标题run空格处理完成 ===")
            text_count += 1

    # 翻译完成后，重新调整所有表格的布局
    try:
        for table in document.tables:
            adjust_table_layout_for_translation(table)
    except Exception as e:
        logger.error(f"最终调整表格布局时出错: {str(e)}")

    # 智能run拼接已经在翻译过程中处理，这里不需要额外处理

    return text_count


def distribute_translation_to_runs_with_adaptive_styles(runs, translated_text, original_text):
    """将翻译结果智能分配回原始run，同时应用自适应样式"""
    
    # 如果只有一个run，直接替换
    if len(runs) == 1:
        runs[0].text = translated_text
        apply_adaptive_styles(runs[0], original_text, translated_text)
        return
    
    # 计算原始文本的总长度
    original_total_length = sum(len(run.text) for run in runs)
    
    # 如果翻译后文本长度变化不大，按比例分配
    if abs(len(translated_text) - original_total_length) <= 2:
        # 按原始比例分配
        current_pos = 0
        for run in runs:
            original_ratio = len(run.text) / original_total_length
            target_length = int(len(translated_text) * original_ratio)
            
            if current_pos < len(translated_text):
                end_pos = min(current_pos + target_length, len(translated_text))
                allocated_text = translated_text[current_pos:end_pos]
                run.text = allocated_text
                
                # 应用自适应样式
                apply_adaptive_styles(run, run.text, allocated_text)
                
                current_pos = end_pos
            else:
                run.text = ""
    else:
        # 长度变化较大，尝试按空格和标点分割
        words = translated_text.split()
        if len(words) >= len(runs):
            # 按run数量分配单词
            for i, run in enumerate(runs):
                if i < len(words):
                    allocated_text = words[i]
                    run.text = allocated_text
                    apply_adaptive_styles(run, run.text, allocated_text)
                else:
                    run.text = ""
        else:
            # 单词数量少于run数量，平均分配
            chars_per_run = len(translated_text) // len(runs)
            for i, run in enumerate(runs):
                start_pos = i * chars_per_run
                end_pos = start_pos + chars_per_run if i < len(runs) - 1 else len(translated_text)
                allocated_text = translated_text[start_pos:end_pos]
                run.text = allocated_text
                apply_adaptive_styles(run, run.text, allocated_text)


def update_special_elements(docx_path, texts):
    """更新批注等特殊元素"""
    # 筛选出需要处理的批注
    comment_texts = [t for t in texts if t.get('type') == 'comment' and t.get('complete')]
    if not comment_texts:
        return  # 没有需要处理的批注

    temp_path = f"temp_{os.path.basename(docx_path)}"

    try:
        with zipfile.ZipFile(docx_path, 'r') as zin, \
                zipfile.ZipFile(temp_path, 'w') as zout:

            for item in zin.infolist():
                with zin.open(item) as file:
                    if item.filename == 'word/comments.xml':
                        try:
                            tree = ET.parse(file)
                            root = tree.getroot()

                            # 定义命名空间
                            namespaces = {
                                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                            }

                            # 查找所有批注
                            for comment in root.findall('.//w:comment', namespaces) or root.findall(
                                    './/{*}comment'):
                                # 尝试不同方式获取ID
                                comment_id = None
                                for attr_name in [
                                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id',
                                    'id']:
                                    if attr_name in comment.attrib:
                                        comment_id = comment.attrib[attr_name]
                                        break

                                if not comment_id:
                                    continue

                                # 查找匹配的翻译结果
                                matching_texts = [t for t in comment_texts if
                                                  t.get('comment_id') == comment_id]
                                if not matching_texts:
                                    continue

                                # 更新批注文本
                                text_elements = comment.findall('.//w:t',
                                                                namespaces) or comment.findall(
                                    './/{*}t')
                                for i, t_elem in enumerate(text_elements):
                                    if i < len(matching_texts):
                                        # 找到匹配的原始文本
                                        for match in matching_texts:
                                            if match.get('original_text') == t_elem.text:
                                                t_elem.text = match.get('text', t_elem.text)
                                                break

                            # 写入修改后的XML
                            modified_xml = ET.tostring(root, encoding='utf-8', xml_declaration=True)
                            zout.writestr(item.filename, modified_xml)
                        except Exception as e:
                            logger.error(f"处理批注时出错: {str(e)}")
                            # 如果解析失败，直接复制原文件
                            file.seek(0)
                            zout.writestr(item.filename, file.read())
                    else:
                        # 直接复制其他文件
                        file.seek(0)
                        zout.writestr(item.filename, file.read())

        # 替换原始文件
        os.replace(temp_path, docx_path)
    except Exception as e:
        logger.error(f"更新批注时出错: {str(e)}")
        # 确保临时文件被删除
        if os.path.exists(temp_path):
            os.remove(temp_path)



def check_text(text):
    """检查文本是否有效（非空且非纯标点）- 优化版本"""
    if not text or len(text) == 0:
        return False
    
    # 检查是否为标题编号（如 "1."、"2."、"3."）
    if TITLE_NUMBER_PATTERN.match(text.strip()):
        return True  # 标题编号应该保留
    
    # 快速检查：如果第一个字符是字母或数字，很可能需要翻译
    if text[0].isalnum():
        return True
    
    # 快速检查：如果全是空格，不需要翻译
    if text.isspace():
        return False
    
    # 快速检查：如果长度很短且全是标点，不需要翻译
    if len(text) <= 3 and all(not c.isalnum() for c in text):
        return False
    
    # 最后检查：使用正则表达式
    if SPECIAL_SYMBOLS_PATTERN.match(text):
        return False
    
    if NUMBERS_PATTERN.match(text):
        return False
    
    return not common.is_all_punc(text)


def apply_translations(document, texts):
    """应用翻译结果到文档，完全保留原始格式"""
    text_count = 0

    for item in texts:
        if not item.get('complete', False):
            continue

        if item['type'] == 'run':
            # 检查run是否包含图片，如果包含则跳过
            if check_if_image(item['run']):
                continue
            # 直接替换run文本，保留所有格式
            item['run'].text = item.get('text', item['run'].text)
            # 如果是文本框，应用自适应样式
            if item.get('context_type') == 'textbox':
                original_text = item['run'].text  # 注意：这里 original_text 应该是替换前的，但由于已替换，使用 item['text'] 作为 translated
                translated_text = item['run'].text
                apply_adaptive_styles(item['run'], original_text, translated_text, context_type='textbox')
            text_count += 1
        elif item['type'] == 'merged_run':
            # 处理合并的run，需要将翻译结果分配回原始run
            merged_item = item['merged_item']
            translated_text = item.get('text', merged_item['text'])
            
            # 如果是标题编号，直接跳过，保持原样
            if item.get('is_title_number'):
                logger.info(f"跳过标题编号: '{merged_item['text']}'")
                continue
            
            # 如果是目录文本，附加原页码
            if item.get('is_toc_text'):
                translated_text += item['original_page_num']
            
            # 检查是否包含图片，如果包含则跳过
            has_image = False
            for run in merged_item['runs']:
                if check_if_image(run):
                    has_image = True
                    break
            
            if has_image:
                continue
            
            if merged_item['type'] == 'merged':
                # 合并的run组，需要智能分配翻译结果
                distribute_translation_to_runs(merged_item['runs'], translated_text, merged_item['text'])
            else:
                # 单个run，直接替换
                merged_item['runs'][0].text = translated_text
            
            # 如果是文本框，应用自适应样式到每个run
            if item.get('context_type') == 'textbox':
                original_text = merged_item['text']  # 原始合并文本
                # 对于每个run，应用自适应（简化：对第一个run应用整体调整）
                for run in merged_item['runs']:
                    apply_adaptive_styles(run, original_text, translated_text, context_type='textbox')
            
            text_count += 1

    # 翻译完成后，重新调整所有表格的布局
    try:
        for table in document.tables:
            adjust_table_layout_for_translation(table)
    except Exception as e:
        logger.error(f"最终调整表格布局时出错: {str(e)}")

    return text_count


def distribute_translation_to_runs(runs, translated_text, original_text=None):
    """将翻译后的文本分配到各个run中"""
    if not runs:
        return
    
    logger.info(f"=== 开始分配翻译结果到runs ===")
    logger.info(f"原始文本: '{original_text}'")
    logger.info(f"翻译文本: '{translated_text}'")
    logger.info(f"Run数量: {len(runs)}")
    
    # 检查是否为主标题的多个run
    is_main_title = False
    if len(runs) > 1 and original_text:
        # 检查原始文本是否包含多个中文字符，可能是完整标题
        chinese_char_count = sum(1 for c in original_text if '\u4e00' <= c <= '\u9fff')
        if chinese_char_count >= 6 and len(original_text) >= 8:  # 至少6个中文字符，总长度至少8
            is_main_title = True
            logger.info(f"检测到主标题run组，原始文本: '{original_text}' (包含{chinese_char_count}个中文字符)")
    
    # 主标题识别已被禁用，直接使用智能空格分布
    logger.info(f"调用智能空格分布函数...")
    smart_text_distribution_with_spaces(runs, translated_text, original_text)
    logger.info(f"=== 翻译结果分配完成 ===")


def debug_spacing_analysis(original_text, translated_text, runs):
    """调试空格分析，帮助诊断问题"""
    logger.info(f"=== 空格分析调试 ===")
    logger.info(f"原始文本: '{original_text}' (长度: {len(original_text)})")
    logger.info(f"翻译文本: '{translated_text}' (长度: {len(translated_text)})")
    logger.info(f"Run数量: {len(runs)}")
    
    # 检测语言
    original_is_chinese = is_chinese_text(original_text)
    translated_is_english = is_english_text(translated_text)
    logger.info(f"原始文本是否中文: {original_is_chinese}")
    logger.info(f"翻译文本是否英文: {translated_is_english}")
    
    # 检测空格需求
    needs_spaces, spacing_type = detect_language_and_spacing_needs(original_text, translated_text)
    logger.info(f"需要添加空格: {needs_spaces}")
    logger.info(f"空格类型: {spacing_type}")
    
    # 分析原始run
    logger.info(f"原始run内容:")
    for i, run in enumerate(runs):
        logger.info(f"  Run {i}: '{run.text}' (长度: {len(run.text)})")
    
    # 分析翻译后文本的单词
    words = translated_text.split()
    logger.info(f"翻译后单词: {words}")
    logger.info(f"单词数量: {len(words)}")
    
    # 分析空格问题
    if needs_spaces and spacing_type == "chinese_to_english":
        logger.info(f"检测到中译英，需要添加空格")
        # 检查是否有连在一起的单词
        for word in words:
            if len(word) > 15:  # 检查是否有异常长的单词
                logger.info(f"  发现长单词: '{word}' (长度: {len(word)})")
                # 尝试分析这个长单词
                if any(c.isupper() for c in word[1:]):  # 检查是否有大写字母（可能是多个单词连在一起）
                    logger.info(f"    可能包含多个单词，建议添加空格")
    
    logger.info("==================")


def smart_text_distribution_with_spaces(runs, translated_text, original_text):
    """智能文本分布，根据语言和空格需求决定处理方式"""
    
    # 调试信息
    debug_spacing_analysis(original_text, translated_text, runs)
    
    # 检测语言和空格需求
    needs_spaces, spacing_type = detect_language_and_spacing_needs(original_text, translated_text)
    
    if needs_spaces:
        # 中译英，需要添加空格
        if len(runs) == 1:
            # 单个run，直接赋值
            runs[0].text = translated_text
            
            # 检查是否以标点符号结尾，如果不是就加空格
            if not translated_text.rstrip().endswith(('.,;:!?()[]{}"\'-')):
                runs[0].text = translated_text + ' '
                logger.info(f"单run不是标点结尾，加结尾空格: '{translated_text}' -> '{runs[0].text}'")
            else:
                logger.info(f"单run以标点结尾，不加空格: '{translated_text}'")
        else:
            # 多个run，使用智能run拼接
            logger.info(f"多run情况，使用智能run拼接")
            smart_run_concatenation(runs, translated_text)
    else:
        # 不需要添加空格，保持原有空格
        distribute_preserving_original_spaces(runs, translated_text, original_text)
    
    # 调试分配结果
    logger.info(f"分配后的run内容:")
    for i, run in enumerate(runs):
        logger.info(f"  Run {i}: '{run.text}' (长度: {len(run.text)})")
    logger.info("==================")


def smart_run_concatenation(runs, translated_text):
    """智能run拼接，在需要的地方自动添加空格"""
    logger.info(f"=== 智能run拼接开始 ===")
    logger.info(f"翻译文本: '{translated_text}'")
    logger.info(f"Run数量: {len(runs)}")
    
    # 将翻译文本按run数量分配
    words = translated_text.split()
    logger.info(f"翻译文本分割后的单词: {words}")
    
    if len(words) >= len(runs):
        # 平均分配单词到每个run
        words_per_run = len(words) // len(runs)
        remainder = len(words) % len(runs)
        
        current_word_index = 0
        for i, run in enumerate(runs):
            # 计算当前run应该包含的单词数
            if i < remainder:
                run_word_count = words_per_run + 1
            else:
                run_word_count = words_per_run
            
            # 提取当前run的单词
            run_words = words[current_word_index:current_word_index + run_word_count]
            run_text = ' '.join(run_words)
            
            # 更新run文本
            run.text = run_text
            current_word_index += run_word_count
            
            logger.info(f"Run {i}: '{run_text}'")
            
            # 检查是否以标点符号结尾，如果不是就加空格
            if not run_text.rstrip().endswith(('.,;:!?()[]{}"\'-')):
                run.text = run_text + ' '
                logger.info(f"Run {i}不是标点结尾，加结尾空格: '{run_text}' -> '{run.text}'")
            else:
                logger.info(f"Run {i}以标点结尾，不加空格: '{run_text}'")
    else:
        # 单词数量少于run数量，简单分配
        for i, run in enumerate(runs):
            if i < len(words):
                run.text = words[i]
                # 检查是否以标点符号结尾，如果不是就加空格
                if not words[i].rstrip().endswith(('.,;:!?()[]{}"\'-')):
                    run.text = words[i] + ' '
                    logger.info(f"Run {i}不是标点结尾，加结尾空格: '{words[i]}' -> '{run.text}'")
                else:
                    logger.info(f"Run {i}以标点结尾，不加空格: '{words[i]}'")
            else:
                run.text = ""
                logger.info(f"Run {i}: 空run")
            logger.info(f"Run {i}最终结果: '{run.text}'")
    
    # 现在在run之间添加空格
    logger.info(f"=== 开始添加run间空格 ===")
    for i in range(1, len(runs)):
        current_run = runs[i]
        prev_run = runs[i-1]
        
        # 检查前一个run是否以标点符号结尾
        prev_text = prev_run.text.strip()
        current_text = current_run.text.strip()
        
        if not prev_text or not current_text:
            continue
            
        prev_ends_with_punct = prev_text[-1] in '.,;:!?'
        current_starts_with_punct = current_text[0] in '.,;:!?'
        
        # 如果前一个run不以标点结尾，且当前run不以标点开头，就在当前run前面加空格
        if not prev_ends_with_punct and not current_starts_with_punct:
            if not current_text.startswith(' '):
                current_run.text = ' ' + current_text
                logger.info(f"在run {i}前面加空格: '{current_run.text}'")
        else:
            logger.info(f"run {i}不需要加空格 (前一个以标点结尾: {prev_ends_with_punct}, 当前以标点开头: {current_starts_with_punct})")
    
    # 检查每个run的结尾，如果不是标点结尾就加空格
    logger.info(f"=== 检查每个run结尾，添加结尾空格 ===")
    for i, run in enumerate(runs):
        if not run.text.strip():
            continue
            
        # 检查是否以标点符号结尾
        run_text = run.text.strip()
        ends_with_punct = run_text[-1] in '.,;:!?()[]{}"\'-'
        
        if not ends_with_punct:
            # 不是标点结尾，在结尾加空格
            run.text = run_text + ' '
            logger.info(f"Run {i}不是标点结尾，加结尾空格: '{run_text}' -> '{run.text}'")
        else:
            logger.info(f"Run {i}以标点结尾，不加空格: '{run_text}'")
    
    logger.info("=== 智能run拼接完成 ===")


def detect_language_and_spacing_needs(original_text, translated_text):
    """更智能地检测语言类型和空格需求"""
    # 检测原始文本语言
    original_is_chinese = is_chinese_text(original_text)
    translated_is_english = is_english_text(translated_text)
    translated_is_chinese = is_chinese_text(translated_text)
    
    # 中译英：需要添加空格
    if original_is_chinese and translated_is_english:
        return True, "chinese_to_english"
    
    # 英译中：通常不需要空格
    if translated_is_chinese and not original_is_chinese:
        return False, "english_to_chinese"
    
    # 其他情况：保持原始空格格式
    return False, "preserve_original"


def is_chinese_text(text):
    """检测是否为中文文本"""
    if not text:
        return False
    
    # 统计中文字符数量
    chinese_chars = 0
    total_chars = 0
    
    for char in text:
        if char.isspace() or char.isdigit() or char in '.,;:!?()[]{}"\'-':
            continue
        
        total_chars += 1
        # 检查是否为中文字符（Unicode范围）
        if '\u4e00' <= char <= '\u9fff':
            chinese_chars += 1
    
    # 如果中文字符占比超过50%，认为是中文文本
    return total_chars > 0 and chinese_chars / total_chars > 0.5


def is_english_text(text):
    """检测是否为英文文本"""
    if not text:
        return False
    
    # 统计英文字符数量
    english_chars = 0
    total_chars = 0
    
    for char in text:
        if char.isspace() or char.isdigit() or char in '.,;:!?()[]{}"\'-':
            continue
        
        total_chars += 1
        # 检查是否为英文字符
        if char.isalpha() and char.isascii():
            english_chars += 1
    
    # 如果英文字符占比超过70%，认为是英文文本
    return total_chars > 0 and english_chars / total_chars > 0.7


def detect_spacing_needs(original_text, translated_text):
    """检测是否需要添加空格（改进版本）"""
    needs_spaces, spacing_type = detect_language_and_spacing_needs(original_text, translated_text)
    
    if spacing_type == "chinese_to_english":
        # 中译英：强制添加空格
        return True
    elif spacing_type == "english_to_chinese":
        # 英译中：保持原始格式
        return False
    else:
        # 其他情况：基于空格存在性判断
        original_has_spaces = ' ' in original_text
        translated_has_spaces = ' ' in translated_text
        
        # 如果原始文本没有空格但翻译后有空格，说明需要添加空格
        if not original_has_spaces and translated_has_spaces:
            return True
        
        # 如果原始文本有空格但翻译后没有，说明需要保持空格
        if original_has_spaces and not translated_has_spaces:
            return False
        
        # 如果都有空格，检查是否需要重新格式化
        if original_has_spaces and translated_has_spaces:
            # 比较空格模式，如果差异很大，可能需要重新格式化
            original_word_count = len(original_text.split())
            translated_word_count = len(translated_text.split())
            if abs(original_word_count - translated_word_count) > 2:
                return True
    
    return False


def distribute_with_proper_spacing(runs, translated_text):
    """为需要添加空格的文本分配run，确保单词间有适当空格（专门处理中译英）"""
    # 暂时注释掉有问题的函数调用
    # 使用专门的run间空格处理函数
    # ensure_proper_spacing_between_runs(runs, translated_text)
    
    # 简单的空格处理
    if len(runs) == 1:
        runs[0].text = translated_text
    else:
        # 多run情况，简单分配
        words = translated_text.split()
        for i, run in enumerate(runs):
            if i < len(words):
                run.text = words[i]
            else:
                run.text = ""


def distribute_preserving_original_spaces(runs, translated_text, original_text):
    """保持原始空格格式的分配"""
    # 计算原始文本的总长度
    original_total_length = sum(len(run.text) for run in runs)
    
    # 如果翻译后文本长度变化不大，按比例分配
    if abs(len(translated_text) - original_total_length) <= 2:
        # 按原始比例分配
        current_pos = 0
        for run in runs:
            original_ratio = len(run.text) / original_total_length
            target_length = int(len(translated_text) * original_ratio)
            
            if current_pos < len(translated_text):
                end_pos = min(current_pos + target_length, len(translated_text))
                run.text = translated_text[current_pos:end_pos]
                current_pos = end_pos
            else:
                run.text = ""
    else:
        # 长度变化较大，按字符平均分配
        chars_per_run = len(translated_text) // len(runs)
        for i, run in enumerate(runs):
            start_pos = i * chars_per_run
            end_pos = start_pos + chars_per_run if i < len(runs) - 1 else len(translated_text)
            run.text = translated_text[start_pos:end_pos]


def adjust_table_layout_for_translation(table):
    """调整表格布局以适应翻译后的文本长度"""
    try:
        from docx.shared import Inches, Cm
        
        # 获取表格的列数和行数
        num_cols = len(table.columns)
        num_rows = len(table.rows)
        
        if num_cols == 0:
            return
        
        # 计算每列的理想宽度
        # 根据列数分配可用宽度，留出一些边距
        available_width = 6.0  # 假设页面宽度为6英寸
        margin = 0.5  # 左右边距
        usable_width = available_width - margin * 2
        
        # 为每列分配宽度，可以根据内容调整
        column_widths = []
        for col_idx in range(num_cols):
            # 计算该列所有单元格的最大文本长度
            max_text_length = 0
            for row_idx in range(num_rows):
                if col_idx < len(table.rows[row_idx].cells):
                    cell = table.rows[row_idx].cells[col_idx]
                    for paragraph in cell.paragraphs:
                        text_length = len(paragraph.text)
                        max_text_length = max(max_text_length, text_length)
            
            # 根据文本长度计算列宽（中文字符大约需要0.1英寸宽度）
            # 最小列宽为0.5英寸，最大为2.0英寸
            estimated_width = max(0.5, min(2.0, max_text_length * 0.1))
            column_widths.append(estimated_width)
        
        # 调整列宽
        for col_idx, width in enumerate(column_widths):
            if col_idx < len(table.columns):
                # 设置列宽
                table.columns[col_idx].width = Inches(width)
                
                # 同时设置该列所有单元格的宽度
                for row_idx in range(num_rows):
                    if col_idx < len(table.rows[row_idx].cells):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.width = Inches(width)
        
        # 设置表格的自动调整属性
        table.autofit = True
        
        # 设置表格样式，确保内容不会超出边界
        table.style = 'Table Grid'
        
    except Exception as e:
        logger.error(f"调整表格布局时出错: {str(e)}")


def process_table_with_layout_adjustment(table, local_texts):
    """处理表格并调整布局"""
    # 处理表格内容
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                extract_paragraph_with_merge(paragraph, local_texts, "table", 0, 1)  # 表格段落索引为0
                # 处理表格中的超链接
                for hyperlink in paragraph.hyperlinks:
                    extract_hyperlink_with_merge(hyperlink, local_texts)
                # 处理表格单元格中的文本框
                for run in paragraph.runs:
                    if check_if_textbox(run):
                        # 处理 DrawingML
                        drawing_elem = run.element.find('.//w:drawing', run.element.nsmap)
                        if drawing_elem is not None:
                            txbx_elem = drawing_elem.find('.//w:txbxContent', drawing_elem.nsmap)
                            if txbx_elem is not None:
                                for p_elem in txbx_elem:
                                    if p_elem.tag == qn('w:p'):
                                        tb_para = Paragraph(p_elem, paragraph)
                                        extract_paragraph_with_merge(tb_para, local_texts, "textbox", 0, 1)
                        
                        # 处理 VML
                        namespaces = {**run.element.nsmap, 'v': 'urn:schemas-microsoft-com:vml'}
                        vtextbox_elem = run.element.find('.//v:textbox', namespaces)
                        if vtextbox_elem is not None:
                            txbx_elem = vtextbox_elem.find('.//w:txbxContent', namespaces)
                            if txbx_elem is not None:
                                for p_elem in txbx_elem:
                                    if p_elem.tag == qn('w:p'):
                                        tb_para = Paragraph(p_elem, paragraph)
                                        extract_paragraph_with_merge(tb_para, local_texts, "textbox", 0, 1)
    
    # 调整表格布局
    adjust_table_layout_for_translation(table)


def convert_chinese_punctuation_to_english(text):
    """将中文标点符号转换为英文标点符号"""
    if not text:
        return text
    
    # 中文标点到英文标点的映射
    punctuation_map = {
        '，': ',',  # 逗号
        '。': '.',  # 句号
        '！': '!',  # 感叹号
        '？': '?',  # 问号
        '；': ';',  # 分号
        '：': ':',  # 冒号
        '（': '(',  # 左括号
        '）': ')',  # 右括号
        '【': '[',  # 左方括号
        '】': ']',  # 右方括号
        '《': '<',  # 左尖括号
        '》': '>',  # 右尖括号
        '"': '"',   # 中文引号
        '"': '"',   # 中文引号
        ''': "'",   # 中文单引号
        ''': "'",   # 中文单引号
        '…': '...', # 省略号
        '—': '-',   # 破折号
        '－': '-',  # 全角连字符
        '　': ' ',  # 全角空格
    }
    
    # 执行转换
    converted_text = text
    for chinese, english in punctuation_map.items():
        converted_text = converted_text.replace(chinese, english)
    
    # 如果文本有变化，记录日志
    if converted_text != text:
        logger.info(f"标点符号转换: '{text}' -> '{converted_text}'")
    
    return converted_text



