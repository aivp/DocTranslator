# -*- coding: utf-8 -*-
import os
import time
import datetime
import traceback
from pathlib import Path
import threading
from docx import Document
from docx.oxml.ns import qn
from . import to_translate
from . import common
import zipfile
import xml.etree.ElementTree as ET
from threading import Lock
import re
import shutil
import fitz
import json
import logging

from ..utils.doc2x import Doc2XService

# 线程安全打印锁
print_lock = Lock()

# 特殊符号和数学符号的正则表达式
SPECIAL_SYMBOLS_PATTERN = re.compile(
    r'^[★☆♥♦♣♠♀♂☼☾☽♪♫♬☑☒✓✔✕✖✗✘⊕⊗∞∑∏πΠ±×÷√∛∜∫∮∇∂∆∏∑√∝∞∟∠∡∢∣∤∥∦∧∨∩∪∫∬∭∮∯∰∱∲∳∴∵∶∷∸∹∺∻∼∽∾∿≀≁≂≃≄≅≆≇≈≉≊≋≌≍≎≏≐≑≒≓≔≕≖≗≘≙≚≛≜≝≞≟≠≡≢≣≤≥≦≧≨≩≪≫≬≭≮≯≰≱≲≳≴≵≶≷≸≹≺≻≼≽≾≿⊀⊁⊂⊃⊄⊅⊆⊇⊈⊉⊊⊋⊌⊍⊎⊏⊐⊑⊒⊓⊔⊕⊖⊗⊘⊙⊚⊛⊜⊝⊞⊟⊠⊡⊢⊣⊤⊥⊦⊧⊨⊩⊪⊫⊬⊭⊮⊯⊰⊱⊲⊳⊴⊵⊶⊷⊸⊹⊺⊻⊼⊽⊾⊿⋀⋁⋂⋃⋄⋅⋆⋇⋈⋉⋊⋋⋌⋍⋎⋏⋐⋑⋒⋓⋔⋕⋖⋗⋘⋙⋚⋛⋜⋝⋞⋟⋠⋡⋢⋣⋤⋥⋦⋧⋨⋩⋪⋫⋬⋭⋮⋯⋰⋱⋲⋳⋴⋵⋶⋷⋸⋹⋺⋻⋼⋽⋾⋿]+$')

# 纯数字和简单标点的正则表达式
NUMBERS_PATTERN = re.compile(r'^[\d\s\.,\-\+\*\/\(\)\[\]\{\}]+$')

def check_docx_quality(docx_path):
    """检查转换后的DOCX文件质量，分析编码和文本内容"""
    try:
        print("\n=== 开始DOCX文件质量检查 ===")
        print("文件路径: " + docx_path)
        
        # 检查文件基本信息
        file_size = os.path.getsize(docx_path)
        print("文件大小: " + str(file_size) + " 字节")
        
        if file_size < 1000:  # 小于1KB可能有问题
            print("⚠️  警告: 文件大小异常，可能转换失败")
            return False
        
        # 尝试加载文档
        try:
            document = Document(docx_path)
            print(f"✅ 文档加载成功")
        except Exception as e:
            print(f"❌ 文档加载失败: {str(e)}")
            return False
        
        # 分析文档结构
        paragraph_count = len(document.paragraphs)
        table_count = len(document.tables)
        section_count = len(document.sections)
        
        print(f"文档结构: {paragraph_count} 个段落, {table_count} 个表格, {section_count} 个节")
        
        # 分析前几个段落的文本内容
        print(f"\n--- 前5个段落内容分析 ---")
        chinese_chars = 0
        total_chars = 0
        sample_texts = []
        
        for i, para in enumerate(document.paragraphs[:5]):
            if para.text.strip():
                text = para.text.strip()
                sample_texts.append(text)
                
                # 统计字符
                para_chinese = sum(1 for char in text if '\u4e00' <= char <= '\u9fff')
                para_total = len(text)
                chinese_chars += para_chinese
                total_chars += para_total
                
                # 显示段落内容（限制长度避免日志过长）
                display_text = text[:100] + "..." if len(text) > 100 else text
                print(f"段落{i+1}: '{display_text}'")
                print(f"  长度: {para_total}, 中文字符: {para_chinese}")
                
                # 显示编码信息
                try:
                    encoded = text.encode('utf-8')
                    print(f"  UTF-8编码: {encoded}")
                except Exception as e:
                    print(f"  编码检查失败: {str(e)}")
        
        # 分析表格内容
        if table_count > 0:
            print(f"\n--- 表格内容分析 ---")
            table_chinese = 0
            table_total = 0
            
            for i, table in enumerate(document.tables[:2]):  # 只分析前2个表格
                print(f"表格{i+1}:")
                for row_idx, row in enumerate(table.rows[:3]):  # 只分析前3行
                    for col_idx, cell in enumerate(row.cells[:3]):  # 只分析前3列
                        if cell.text.strip():
                            text = cell.text.strip()
                            cell_chinese = sum(1 for char in text if '\u4e00' <= char <= '\u9fff')
                            cell_total = len(text)
                            table_chinese += cell_chinese
                            table_total += cell_total
                            
                            if cell_total > 0:
                                display_text = text[:50] + "..." if len(text) > 50 else text
                                print(f"  单元格[{row_idx+1},{col_idx+1}]: '{display_text}' (中文字符: {cell_chinese})")
            
            chinese_chars += table_chinese
            total_chars += table_total
            print(f"表格总计: {table_chinese} 个中文字符, {table_total} 个总字符")
        
        # 统计结果
        print(f"\n--- 质量检查结果 ---")
        if total_chars > 0:
            chinese_ratio = (chinese_chars / total_chars) * 100
            print(f"中文字符比例: {chinese_chars}/{total_chars} ({chinese_ratio:.1f}%)")
            
            if chinese_ratio < 10:
                print("⚠️  警告: 中文字符比例过低，可能存在编码问题")
            elif chinese_ratio > 80:
                print("✅ 中文字符比例正常")
            else:
                print("ℹ️  中文字符比例中等")
        else:
            print("⚠️  警告: 未发现任何文本内容")
        
        # 检查是否有明显的问题
        if paragraph_count == 0 and table_count == 0:
            print("❌ 严重问题: 文档没有任何内容")
            return False
        
        if chinese_chars == 0 and total_chars > 0:
            print("⚠️  警告: 有文本内容但没有中文字符，可能存在编码问题")
        
        print(f"=== DOCX质量检查完成 ===\n")
        return True
        
    except Exception as e:
        print(f"❌ DOCX质量检查失败: {str(e)}")
        traceback.print_exc()
        return False


def get_doc2x_save_dir():
    # 获取基础存储目录
    base_dir = Path(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))).parent.absolute()

    # 创建日期子目录（YYYY-MM-DD）
    date_str = datetime.datetime.now().strftime('%Y-%m-%d')
    save_dir = base_dir / "storage" / "doc2x_results" / date_str

    save_dir.mkdir(parents=True, exist_ok=True)

    return str(save_dir)


def start(trans):
    """PDF翻译入口"""
    print("🚨 DEBUG: PDF翻译函数被调用！")
    print("🚨 DEBUG: 这是强制输出测试")
    
    try:
        # 开始时间
        start_time = datetime.datetime.now()
        print(f"=== 开始PDF翻译任务 ===")
        print(f"任务ID: {trans['id']}")
        print(f"源文件: {trans['file_path']}")
        print(f"目标文件: {trans['target_file']}")
        print(f"开始时间: {start_time}")
        print("=" * 50)
        
        # 检查PDF翻译方法设置
        # 优先使用trans中的pdf_translate_method，如果没有则从系统设置中获取
        pdf_translate_method = trans.get('pdf_translate_method')
        if not pdf_translate_method:
            pdf_translate_method = get_pdf_translate_method()
        print(f"📋 PDF翻译方法: {pdf_translate_method}")
        
        # 根据设置选择翻译方法
        if pdf_translate_method == 'direct':
            print("🎯 使用直接PDF翻译方法")
            return start_direct_pdf_translation(trans)
        else:
            print("🎯 使用Doc2x转换后翻译方法")
            return start_doc2x_pdf_translation(trans)

    except Exception as e:
        # 打印详细错误信息
        print("❌ PDF翻译过程出错: " + str(e))
        print("详细错误信息:")
        traceback.print_exc()
        # 确保错误状态被正确记录
        to_translate.error(trans['id'], "PDF翻译过程出错: " + str(e))
        return False


def start_doc2x_pdf_translation(trans):
    """Doc2x转换后翻译方法（原有逻辑）"""
    try:
        # 开始时间
        start_time = datetime.datetime.now()
        print(f"=== 开始Doc2x PDF翻译任务 ===")
        print(f"任务ID: {trans['id']}")
        print(f"源文件: {trans['file_path']}")
        print(f"目标文件: {trans['target_file']}")
        print(f"开始时间: {start_time}")
        print("=" * 50)
        
        # 立即更新任务状态为"changing"，设置PDF转换初始进度0%
        try:
            from .to_translate import db
            db.execute("update translate set status='changing', process='0' where id=%s", trans['id'])
            print("✅ 已更新任务状态为changing，进度0%（开始PDF转换）")
        except Exception as e:
            print(f"⚠️  更新任务状态失败: {str(e)}")

        # 检查文件是否存在
        original_path = Path(trans['file_path'])
        print(f"检查源文件是否存在: {original_path}")
        if not original_path.exists():
            print(f"❌ 源文件不存在: {trans['file_path']}")
            to_translate.error(trans['id'], "文件不存在: " + trans['file_path'])
            return False
        print(f"✅ 源文件存在，大小: {os.path.getsize(original_path)} 字节")

        # 确保目标目录存在
        target_dir = os.path.dirname(trans['target_file'])
        print(f"创建目标目录: {target_dir}")
        os.makedirs(target_dir, exist_ok=True)
        print(f"✅ 目标目录已准备")

        # 获取doc2x保存目录
        doc2x_save_dir = get_doc2x_save_dir()
        print(f"📁 doc2x保存目录: {doc2x_save_dir}")

        # 设置转换后的Word文件路径（保存在doc2x_results目录下）
        docx_filename = f"{original_path.stem}_doc2x.docx"
        docx_path = os.path.join(doc2x_save_dir, docx_filename)
        print(f"📄 转换后的DOCX文件路径: {docx_path}")

        # 如果文件已存在，先删除
        if os.path.exists(docx_path):
            try:
                os.remove(docx_path)
                print(f"已删除已存在的文件: {docx_path}")
            except Exception as e:
                print(f"删除已存在的文件失败: {str(e)}")
                # 使用时间戳创建唯一文件名
                docx_filename = f"{original_path.stem}_doc2x_{int(time.time())}.docx"
                docx_path = os.path.join(doc2x_save_dir, docx_filename)

        print(f"转换后的DOCX文件将保存在: {docx_path}")

        # 使用Doc2X服务将PDF转换为DOCX
        print(f"开始将PDF转换为DOCX: {original_path}")
        import logging
        logger = logging.getLogger(__name__)
        
        # 记录Doc2X开始时间
        doc2x_start_time = datetime.datetime.now()
        logger.info(f"🚀 Doc2X转换开始时间: {doc2x_start_time}")

        # 获取API密钥
        api_key = trans.get('doc2x_api_key', '')
        print(f"🔑 检查API密钥: {'已设置' if api_key else '未设置'}")
        if not api_key:
            print(f"❌ 缺少Doc2X API密钥")
            to_translate.error(trans['id'], "缺少Doc2X API密钥")
            return False
        print(f"✅ API密钥已设置")

        # 1. 启动转换任务
        print(f"🚀 开始启动Doc2X转换任务...")
        task_start_time = datetime.datetime.now()
        logger.info(f"📤 Doc2X任务启动时间: {task_start_time}")
        try:
            uid = Doc2XService.start_task(api_key, str(original_path))
            task_end_time = datetime.datetime.now()
            task_duration = task_end_time - task_start_time
            logger.info(f"📤 Doc2X任务启动耗时: {task_duration.total_seconds():.2f}秒")
            print(f"✅ Doc2X任务启动成功，UID: {uid}")
            
            # Doc2X任务启动成功，设置为changing状态
            try:
                from .to_translate import db
                print(f"🔍 准备更新任务状态: 任务ID={trans['id']}, 新状态=changing, 新进度=0")
                result = db.execute("update translate set status='changing', process='0' where id=%s", trans['id'])
                print(f"🔍 数据库更新结果: {result}")
                print("✅ 已更新任务状态为changing，进度0%（Doc2X任务启动成功）")
            except Exception as e:
                print(f"❌ 更新任务状态失败: {str(e)}")
                # 如果状态更新失败，记录错误但不中断流程
                import traceback
                traceback.print_exc()
                
        except Exception as e:
            print(f"❌ 启动Doc2X任务失败: {str(e)}")
            to_translate.error(trans['id'], f"启动Doc2X任务失败: {str(e)}")
            return False

        # 2. 等待解析完成
        max_retries = 60  # 最多等待10分钟
        retry_count = 0
        parse_start_time = datetime.datetime.now()
        logger.info(f"🔄 Doc2X解析开始时间: {parse_start_time}")
        
        while retry_count < max_retries:
            try:
                status_info = Doc2XService.check_parse_status(api_key, uid)
                if status_info['status'] == 'success':
                    parse_end_time = datetime.datetime.now()
                    parse_duration = parse_end_time - parse_start_time
                    logger.info(f"🔄 Doc2X解析完成时间: {parse_end_time}")
                    logger.info(f"⏱️  Doc2X解析总耗时: {parse_duration.total_seconds():.2f}秒")
                    print(f"PDF解析成功: {uid}")
                    
                    # Doc2X解析完成，但保持changing状态，继续导出和下载
                    print(f"PDF解析成功: {uid}")
                    # 不在这里切换状态，等所有Doc2X阶段完成后再切换
                    break
                elif status_info['status'] == 'failed':
                    to_translate.error(trans['id'],
                                       f"PDF解析失败: {status_info.get('detail', '未知错误')}")
                    return False

                # 模拟进度：基于重试次数计算进度（0-90%）
                # 前90%用于解析阶段，最后10%留给导出和下载
                simulated_progress = min(90, int((retry_count / max_retries) * 90))
                print(f"🔍 模拟进度: {simulated_progress}% (重试次数: {retry_count}/{max_retries})")
                
                # 实时更新数据库进度
                try:
                    print(f"🔍 准备更新进度: 任务ID={trans['id']}, 新进度={simulated_progress}%")
                    result = db.execute("update translate set process=%s where id=%s", str(simulated_progress), trans['id'])
                    print(f"🔍 进度更新结果: {result}")
                    print(f"✅ 已更新进度为 {simulated_progress}%")
                except Exception as e:
                    print(f"❌ 更新进度失败: {str(e)}")
                    import traceback
                    traceback.print_exc()

                # 等待10秒后再次检查
                time.sleep(10)
                retry_count += 1
            except Exception as e:
                to_translate.error(trans['id'], f"检查解析状态失败: {str(e)}")
                return False

        if retry_count >= max_retries:
            to_translate.error(trans['id'], "PDF解析超时")
            return False

        # 3. 触发导出
        export_start_time = datetime.datetime.now()
        logger.info(f"📤 Doc2X导出开始时间: {export_start_time}")
        
        # 开始导出阶段，进度设为95%（changing状态下的进度）
        print("📤 开始导出阶段")
        try:
            from .to_translate import db
            db.execute("update translate set process='95' where id=%s", trans['id'])
            print("✅ 已更新进度为95%（开始导出）")
        except Exception as e:
            print("⚠️  更新进度失败: " + str(e))
            
        try:
            export_success = Doc2XService.trigger_export(api_key, uid, original_path.stem)
            if not export_success:
                to_translate.error(trans['id'], "触发导出失败")
                return False
            export_end_time = datetime.datetime.now()
            export_duration = export_end_time - export_start_time
            logger.info(f"📤 Doc2X导出耗时: {export_duration.total_seconds():.2f}秒")
            print(f"已触发导出: {uid}")
        except Exception as e:
            to_translate.error(trans['id'], f"触发导出失败: {str(e)}")
            return False

        # 4. 等待导出完成并下载
        download_start_time = datetime.datetime.now()
        logger.info(f"📥 Doc2X下载开始时间: {download_start_time}")
        
        # 开始下载阶段，进度设为98%（changing状态下的进度）
        print("📥 开始下载阶段")
        try:
            db.execute("update translate set process='98' where id=%s", trans['id'])
            print("✅ 已更新进度为98%（开始下载）")
        except Exception as e:
            print("⚠️  更新进度失败: " + str(e))
            
        try:
            download_url = Doc2XService.check_export_status(api_key, uid)
            print(f"获取到下载链接: {download_url}")

            # 下载文件
            download_success = Doc2XService.download_file(download_url, docx_path)
            if not download_success:
                to_translate.error(trans['id'], "下载转换后的DOCX文件失败")
                return False
            download_end_time = datetime.datetime.now()
            download_duration = download_end_time - download_start_time
            logger.info(f"📥 Doc2X下载耗时: {download_duration.total_seconds():.2f}秒")
            print(f"DOCX文件下载成功: {docx_path}")
            
            # 下载完成，进度设为100%（changing状态下的进度）
            try:
                db.execute("update translate set process='100' where id=%s", trans['id'])
                print("✅ 已更新进度为100%（下载完成）")
            except Exception as e:
                print("⚠️  更新进度失败: " + str(e))
            
            # Doc2X所有阶段完成，切换到process状态开始翻译
            try:
                print(f"🔍 准备切换到process状态: 任务ID={trans['id']}, 新状态=process, 新进度=0")
                result = db.execute("update translate set status='process', process='0' where id=%s", trans['id'])
                print(f"🔍 状态切换结果: {result}")
                print("✅ 已更新状态为process，进度0%（Doc2X全部完成，开始翻译）")
            except Exception as e:
                print(f"❌ 更新状态失败: {str(e)}")
                import traceback
                traceback.print_exc()
                
        except Exception as e:
            to_translate.error(trans['id'], f"下载DOCX文件失败: {str(e)}")
            return False

        # 确认文件存在
        if not os.path.exists(docx_path):
            to_translate.error(trans['id'], "转换后的DOCX文件不存在")
            return False

        # 4.5. 检查转换后的DOCX文件质量
        print(f"开始检查转换后的DOCX文件质量...")
        if not check_docx_quality(docx_path):
            print("⚠️  DOCX文件质量检查发现问题，但继续处理...")
        else:
            print("✅ DOCX文件质量检查通过")
            
        # 记录Doc2X完成时间并计算耗时
        doc2x_end_time = datetime.datetime.now()
        doc2x_duration = doc2x_end_time - doc2x_start_time
        logger.info(f"✅ Doc2X转换完成时间: {doc2x_end_time}")
        logger.info(f"⏱️  Doc2X转换总耗时: {doc2x_duration}")
        logger.info(f"📊 Doc2X转换耗时详情:")
        logger.info(f"   - 开始时间: {doc2x_start_time}")
        logger.info(f"   - 完成时间: {doc2x_end_time}")
        logger.info(f"   - 总耗时: {doc2x_duration}")
        logger.info(f"   - 耗时秒数: {doc2x_duration.total_seconds():.2f}秒")

        # 5. 使用Word翻译逻辑处理DOCX文件
        # 创建一个新的trans对象，包含DOCX文件路径
        docx_trans = trans.copy()
        docx_trans['file_path'] = docx_path

        # 设置目标文件路径（如果是PDF，添加.docx扩展名）
        target_file = trans['target_file']
        if target_file.lower().endswith('.pdf'):
            target_file = target_file + '.docx'
        docx_trans['target_file'] = target_file

        # 加载Word文档
        try:
            document = Document(docx_path)
            print("成功加载DOCX文档: " + docx_path)
        except Exception as e:
            to_translate.error(trans['id'], f"文档加载失败: {str(e)}")
            return False

        # 提取需要翻译的文本
        texts = []
        extract_content_for_translation(document, docx_path, texts)
        print("从DOCX提取了 " + str(len(texts)) + " 个文本片段")

        # 过滤掉特殊符号和纯数字
        filtered_texts = []
        skipped_count = 0
        for i, item in enumerate(texts):
            if should_translate(item['text']):
                filtered_texts.append(item)
            else:
                # 对于不需要翻译的内容，标记为已完成
                item['complete'] = True
                text = item['text']
                reason = get_skip_reason(text)
                skipped_count += 1
                
                # 限制日志长度，避免输出过多
                if skipped_count <= 50:  # 只显示前50个跳过的项目
                    display_text = text[:50] + "..." if len(text) > 50 else text
                    print("跳过翻译: '" + display_text + "' - 原因: " + reason)
                elif skipped_count == 51:
                    print("... 还有更多跳过的项目，不再显示详细原因 ...")

        print("过滤后需要翻译的文本片段: " + str(len(filtered_texts)))
        print("跳过的文本片段: " + str(skipped_count))

        # 使用Okapi进行XLIFF转换，然后Qwen翻译
        print("🔄 使用Okapi进行XLIFF转换，然后Qwen翻译...")
        
        try:
            # 导入 Okapi 集成模块
            from .okapi_integration import OkapiWordTranslator, verify_okapi_installation
            
            # 验证 Okapi 安装
            if not verify_okapi_installation():
                print("❌ Okapi 安装验证失败，回退到传统方法")
                run_translation(docx_trans, filtered_texts, max_threads=30)
            else:
                print("✅ Okapi 安装验证成功，使用XLIFF转换方案")
                
                # 创建 Okapi 翻译器
                translator = OkapiWordTranslator()
                print("✅ Okapi 翻译器创建成功")
                
                # 设置翻译服务（Qwen）
                if docx_trans.get('model') == 'qwen-mt-plus':
                    docx_trans['server'] = 'qwen'
                    print("✅ 设置翻译服务为 Qwen")
                
                # 预加载术语库
                comparison_id = docx_trans.get('comparison_id')
                if comparison_id:
                    print("📚 开始预加载术语库: " + str(comparison_id))
                    from .main import get_comparison
                    preloaded_terms = get_comparison(comparison_id)
                    if preloaded_terms:
                        print("📚 术语库预加载成功: " + str(len(preloaded_terms)) + " 个术语")
                        docx_trans['preloaded_terms'] = preloaded_terms
                    else:
                        print("📚 术语库预加载失败: " + str(comparison_id))
                
                # 设置翻译服务
                translator.set_translation_service(docx_trans)
                print("✅ Okapi 翻译服务设置成功")
                
                # 语言映射：将中文语言名称转换为英文全拼
                def map_language_to_qwen_format(lang_name):
                    language_mapping = {
                        '中文': 'Chinese',
                        '英语': 'English',
                        '日语': 'Japanese',
                        '韩语': 'Korean',
                        '法语': 'French',
                        '德语': 'German',
                        '西班牙语': 'Spanish',
                        '俄语': 'Russian',
                        '阿拉伯语': 'Arabic',
                        '葡萄牙语': 'Portuguese',
                        '意大利语': 'Italian',
                        '泰语': 'Thai',
                        '越南语': 'Vietnamese',
                        '印尼语': 'Indonesian',
                        '马来语': 'Malay',
                        '菲律宾语': 'Filipino',
                        '缅甸语': 'Burmese',
                        '柬埔寨语': 'Khmer',
                        '老挝语': 'Lao',
                        '柬语': 'Khmer'
                    }
                    return language_mapping.get(lang_name.strip(), lang_name.strip())
                
                # 获取并映射语言
                source_lang = "auto"  # 写死为auto，让API自动检测源语言
                target_lang = map_language_to_qwen_format(docx_trans.get('lang', '英语'))
                
                print("🔍 语言映射调试:")
                print("  原始目标语言: " + str(docx_trans.get('lang', '英语')))
                print("  映射后目标语言: " + target_lang)
                
                # 执行翻译：Okapi转换XLIFF，Qwen翻译，然后合并
                success = translator.translate_document(
                    docx_path,
                    target_file,
                    source_lang,
                    target_lang
                )
                
                if success:
                    print("✅ Okapi XLIFF转换 + Qwen翻译完成")
                    
                    # 更新进度为100%（翻译完成）
                    try:
                        db.execute("update translate set process='100' where id=%s", trans['id'])
                        print("✅ 已更新进度为100%（翻译完成）")
                    except Exception as e:
                        print("⚠️  更新进度失败: " + str(e))
                    
                    # 完成处理
                    end_time = datetime.datetime.now()
                    spend_time = common.display_spend(start_time, end_time)
                    
                    # 统计翻译的文本数量（这里简化处理）
                    text_count = len(filtered_texts)  # 使用过滤后的文本数量
                    
                    if docx_trans['run_complete']:
                        to_translate.complete(docx_trans, text_count, spend_time)
                    
                    print("✅ Okapi XLIFF转换 + Qwen翻译完成，用时: " + spend_time)
                    return True
                else:
                    print("❌ Okapi XLIFF转换 + Qwen翻译失败，回退到传统方法")
                    run_translation(docx_trans, filtered_texts, max_threads=30)
                    
        except Exception as e:
            print("❌ Okapi XLIFF转换 + Qwen翻译出错: " + str(e) + "，回退到传统方法")
            run_translation(docx_trans, filtered_texts, max_threads=30)

        # 写入翻译结果（完全保留原始格式）
        text_count = apply_translations(document, texts)
        print("应用了 " + str(text_count) + " 个翻译结果")

        # 保存文档
        try:
            document.save(target_file)
            print("翻译后的文档保存成功: " + target_file)
        except Exception as e:
            to_translate.error(trans['id'], "保存文档失败: " + str(e))
            return False

        # 处理批注等特殊元素
        update_special_elements(target_file, texts)

        # 7. 完成处理
        end_time = datetime.datetime.now()
        spend_time = common.display_spend(start_time, end_time)
        
        # 更新进度为100%（传统方法翻译完成）
        try:
            db.execute("update translate set process='100' where id=%s", trans['id'])
            print("✅ 已更新进度为100%（传统方法翻译完成）")
        except Exception as e:
            print("⚠️  更新进度失败: " + str(e))
        
        if trans['run_complete']:
            to_translate.complete(trans, text_count, spend_time)
        print("PDF翻译任务完成: " + str(trans['id']))
        return True

    except Exception as e:
        # 打印详细错误信息
        print("❌ PDF翻译过程出错: " + str(e))
        print("详细错误信息:")
        traceback.print_exc()
        # 确保错误状态被正确记录
        to_translate.error(trans['id'], "PDF翻译过程出错: " + str(e))
        return False


def check_image(run):
    """检查run是否包含图片"""
    try:
        # 检查run中是否有图片元素
        for element in run._element:
            if element.tag.endswith('drawing') or element.tag.endswith('picture'):
                return True
        return False
    except:
        return False


def are_runs_compatible(run1, run2):
    """检查两个run是否兼容（最严格的兼容性检查）"""
    try:
        # 1. 字体名称必须完全相同
        if run1.font.name != run2.font.name:
            return False
        
        # 2. 字体大小必须完全相同
        if run1.font.size != run2.font.size:
            return False
        
        # 3. 粗体状态必须完全相同
        if run1.font.bold != run2.font.bold:
            return False
        
        # 4. 斜体状态必须完全相同
        if run1.font.italic != run2.font.italic:
            return False
        
        # 5. 下划线状态必须完全相同
        if run1.font.underline != run2.font.underline:
            return False
        
        # 6. 删除线状态必须完全相同
        if run1.font.strike != run2.font.strike:
            return False
        
        # 7. 字体颜色必须完全相同
        if run1.font.color.rgb != run2.font.color.rgb:
            return False
        
        # 8. 背景色必须完全相同
        if run1.font.highlight_color != run2.font.highlight_color:
            return False
        
        # 9. 上标/下标状态必须完全相同
        if run1.font.superscript != run2.font.superscript:
            return False
        
        # 10. 小型大写字母状态必须完全相同
        if run1.font.small_caps != run2.font.small_caps:
            return False
        
        return True
    except:
        # 如果任何检查失败，默认不兼容
        return False


def extract_paragraph_with_merge(paragraph, texts, context_type):
    """提取段落内容，使用保守run合并（不添加上下文）"""
    paragraph_runs = list(paragraph.runs)
    
    # 使用保守的run合并策略
    merged_runs = conservative_run_merge(paragraph_runs)
    
    for merged_item in merged_runs:
        if not check_text(merged_item['text']):
            continue
        
        # 检查是否包含图片，如果包含则跳过
        has_image = False
        for run in merged_item['runs']:
            if check_image(run):
                has_image = True
                break
        
        if has_image:
            continue
        
        texts.append({
            "text": merged_item['text'],
            "type": "merged_run",
            "merged_item": merged_item,
            "complete": False,
            "context_type": context_type  # 标记类型
        })


def conservative_run_merge(paragraph_runs, max_merge_length=500):
    """最严格的run合并策略"""
    
    merged = []
    current_group = []
    current_length = 0
    original_count = len([r for r in paragraph_runs if check_text(r.text)])
    merged_count = 0
    
    for run in paragraph_runs:
        # 跳过包含图片的run
        if check_image(run):
            # 保存当前组
            if current_group:
                merged.append(merge_compatible_runs(current_group))
                if len(current_group) > 1:
                    merged_count += 1
                current_group = []
                current_length = 0
            
            # 图片run单独处理，但不添加到翻译列表
            continue
        
        if not check_text(run.text):
            continue
        
        # 最严格条件：只合并极短的run（通常是空格、标点、单个字符）
        if len(run.text) <= 5 and current_length < max_merge_length:
            # 检查格式兼容性（最严格检查）
            if not current_group or are_runs_compatible(current_group[-1], run):
                # 额外检查：确保合并后的文本不会过长
                if current_length + len(run.text) <= 10:
                    current_group.append(run)
                    current_length += len(run.text)
                    continue
        
        # 保存当前组
        if current_group:
            merged.append(merge_compatible_runs(current_group))
            if len(current_group) > 1:
                merged_count += 1
            current_group = []
            current_length = 0
        
        # 当前run单独处理
        merged.append({
            'text': run.text,
            'runs': [run],
            'type': 'single'
        })
    
    if current_group:
        merged.append(merge_compatible_runs(current_group))
        if len(current_group) > 1:
            merged_count += 1
    
    # 打印合并统计信息
    if merged_count > 0:
        print(f"Run合并优化（最严格策略）: 原始{original_count}个run -> 合并后{len(merged)}个，减少了{merged_count}个API调用")
    else:
        print(f"Run合并（最严格策略）: 未找到可合并的run，保持原始{original_count}个run")
    
    return merged


def merge_compatible_runs(run_group):
    """合并兼容的run组（最严格策略）"""
    
    # 合并文本
    merged_text = "".join(run.text for run in run_group)
    
    # 额外安全检查：确保合并后的文本不会过长
    if len(merged_text) > 20:
        print(f"⚠️  警告: 合并后文本过长({len(merged_text)}字符)，取消合并")
        # 如果合并后过长，返回单个run
        return {
            'text': run_group[0].text,
            'runs': [run_group[0]],
            'type': 'single'
        }
    
    # 额外检查：确保所有run的格式完全一致
    first_run = run_group[0]
    for run in run_group[1:]:
        if not are_runs_compatible(first_run, run):
            print(f"⚠️  警告: 检测到格式不一致，取消合并")
            return {
                'text': run_group[0].text,
                'runs': [run_group[0]],
                'type': 'single'
            }
    
    return {
        'text': merged_text,
        'runs': run_group,
        'type': 'merged'
    }


def get_skip_reason(text):
    """获取文本被跳过的原因"""
    if not text or len(text) == 0:
        return "空文本"
    
    if len(text.strip()) <= 2:
        text_stripped = text.strip()
        
        # 检查中文字符
        if all('\u4e00' <= char <= '\u9fff' for char in text_stripped):
            return f"短中文字符'{text_stripped}'（已允许翻译）"
        
        # 检查数字
        if text_stripped.isdigit():
            return f"短数字'{text_stripped}'（已允许翻译）"
        
        # 检查单位符号
        if all(char in '•℃VAhH' for char in text_stripped):
            return f"常用单位符号'{text_stripped}'（已允许翻译）"
        
        # 检查中英文混合
        if len(text_stripped) == 2:
            has_chinese = any('\u4e00' <= char <= '\u9fff' for char in text_stripped)
            if has_chinese:
                return f"中英文混合短文本'{text_stripped}'（已允许翻译）"
        
        return f"短文本'{text_stripped}'（无意义）"
    
    if SPECIAL_SYMBOLS_PATTERN.match(text):
        return "纯特殊符号"
    
    if NUMBERS_PATTERN.match(text):
        return "纯数字和简单标点"
    
    if common.is_all_punc(text):
        return "纯标点符号"
    
    return "其他原因"


def should_translate(text):
    """判断文本是否需要翻译（改进的过滤逻辑）"""
    if not text or len(text) == 0:
        return False

    # 对于短文本（≤2字符），进行特殊处理
    if len(text.strip()) <= 2:
        text_stripped = text.strip()
        
        # 允许短中文字符（通常是有意义的）
        if all('\u4e00' <= char <= '\u9fff' for char in text_stripped):
            return True
        
        # 允许短数字组合
        if text_stripped.isdigit():
            return True
        
        # 允许常用单位符号组合
        if all(char in '•℃VAhH' for char in text_stripped):
            return True
        
        # 允许中英文混合的短文本（如"图1"、"表2"）
        if len(text_stripped) == 2:
            # 检查是否包含中文字符
            has_chinese = any('\u4e00' <= char <= '\u9fff' for char in text_stripped)
            if has_chinese:
                return True
        
        # 过滤其他短文本
        return False

    # 过滤纯特殊符号
    if SPECIAL_SYMBOLS_PATTERN.match(text):
        return False

    # 过滤纯数字和简单标点
    if NUMBERS_PATTERN.match(text):
        return False

    # 过滤全是标点符号的文本
    if common.is_all_punc(text):
        return False

    return True


def extract_content_for_translation(document, file_path, texts):
    """提取需要翻译的内容，使用run合并优化"""
    # 正文段落
    for paragraph in document.paragraphs:
        extract_paragraph_with_merge(paragraph, texts, "paragraph")
        
        # 处理超链接（单独处理，避免与普通run混合）
        for hyperlink in paragraph.hyperlinks:
            for run in hyperlink.runs:
                if check_text(run.text):
                    texts.append({
                        "text": run.text,
                        "type": "run",
                        "run": run,
                        "complete": False
                    })

    # 表格内容
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    extract_paragraph_with_merge(paragraph, texts, "table_cell")
                    
                    # 处理表格中的超链接
                    for hyperlink in paragraph.hyperlinks:
                        for run in hyperlink.runs:
                            if check_text(run.text):
                                texts.append({
                                    "text": run.text,
                                    "type": "run",
                                    "run": run,
                                    "complete": False
                                })

    # 页眉页脚
    for section in document.sections:
        # 处理页眉
        for paragraph in section.header.paragraphs:
            extract_paragraph_with_merge(paragraph, texts, "header")
            
            # 处理页眉中的超链接
            for hyperlink in paragraph.hyperlinks:
                for run in hyperlink.runs:
                    if check_text(run.text):
                        texts.append({
                            "text": run.text,
                            "type": "run",
                            "run": run,
                            "complete": False
                        })

        # 处理页脚
        for paragraph in section.footer.paragraphs:
            extract_paragraph_with_merge(paragraph, texts, "footer")
            
            # 处理页脚中的超链接
            for hyperlink in paragraph.hyperlinks:
                for run in hyperlink.runs:
                    if check_text(run.text):
                        texts.append({
                            "text": run.text,
                            "type": "run",
                            "run": run,
                            "complete": False
                        })

    # 批注内容
    extract_comments(file_path, texts)


def extract_comments(file_path, texts):
    """提取文档中的批注"""
    try:
        with zipfile.ZipFile(file_path, 'r') as docx:
            if 'word/comments.xml' in docx.namelist():
                with docx.open('word/comments.xml') as f:
                    tree = ET.parse(f)
                    root = tree.getroot()

                    # 定义命名空间
                    namespaces = {
                        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                    }

                    # 查找所有批注
                    for comment in root.findall('.//w:comment', namespaces) or root.findall(
                            './/{*}comment'):
                        # 尝试不同方式获取ID
                        comment_id = None
                        for attr_name in [
                            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id',
                            'id']:
                            if attr_name in comment.attrib:
                                comment_id = comment.attrib[attr_name]
                                break

                        if not comment_id:
                            continue

                        # 收集批注中的所有文本
                        text_elements = comment.findall('.//w:t', namespaces) or comment.findall(
                            './/{*}t')
                        for t_elem in text_elements:
                            if t_elem.text and check_text(t_elem.text):
                                texts.append({
                                    "text": t_elem.text,
                                    "type": "comment",
                                    "comment_id": comment_id,
                                    "complete": False,
                                    "original_text": t_elem.text  # 保存原始文本用于匹配
                                })
    except Exception as e:
        print(f"提取批注时出错: {str(e)}")


def run_translation(trans, texts, max_threads=30):
    # 硬编码线程数为30，忽略前端传入的配置
    """执行多线程翻译"""
    if not texts:
        print("没有需要翻译的内容")
        return

    event = threading.Event()
    run_index = 0
    active_count = threading.activeCount()
    
    # 进度更新相关变量
    completed_count = 0
    total_count = len(texts)
    progress_lock = threading.Lock()
    
    def update_progress():
        """更新翻译进度"""
        nonlocal completed_count
        with progress_lock:
            completed_count += 1
            progress_percentage = min((completed_count / total_count) * 100, 100.0)
            print(f"翻译进度: {completed_count}/{total_count} ({progress_percentage:.1f}%)")
            
            # 更新数据库进度
            try:
                from .to_translate import db
                db.execute("update translate set process=%s where id=%s", 
                         str(format(progress_percentage, '.1f')), 
                         trans['id'])
                
            except Exception as e:
                print(f"更新进度失败: {str(e)}")

    print(f"开始翻译 {len(texts)} 个文本片段")

    while run_index < len(texts):
        if threading.activeCount() < max_threads + active_count and not event.is_set():
            thread = threading.Thread(
                target=to_translate.get,
                args=(trans, event, texts, run_index)
            )
            thread.start()
            print(f"启动翻译线程 {run_index}")
            run_index += 1
        time.sleep(0.1)

    # 等待翻译完成，并监控进度
    last_completed_count = 0
    while not all(t.get('complete') for t in texts) and not event.is_set():
        # 检查完成的文本数量
        current_completed = sum(1 for t in texts if t.get('complete', False))
        if current_completed > last_completed_count:
            # 有新的文本完成，更新进度
            completed_count = current_completed
            progress_percentage = min((completed_count / total_count) * 100, 100.0)
            print(f"翻译进度: {completed_count}/{total_count} ({progress_percentage:.1f}%)")
            
            # 更新数据库进度
            try:
                from .to_translate import db
                db.execute("update translate set process=%s where id=%s", 
                         str(format(progress_percentage, '.1f')), 
                         trans['id'])
                
            except Exception as e:
                print(f"更新进度失败: {str(e)}")
            
            last_completed_count = current_completed
        
        time.sleep(1)

    print("所有翻译任务已完成")


def apply_translations(document, texts):
    """应用翻译结果到文档，完全保留原始格式"""
    text_count = 0

    for item in texts:
        if not item.get('complete', False):
            continue

        if item['type'] == 'run':
            # 直接替换run文本，保留所有格式
            item['run'].text = item.get('text', item['run'].text)
            text_count += 1
        elif item['type'] == 'merged_run':
            # 处理合并的run，需要将翻译结果分配回原始run
            merged_item = item['merged_item']
            translated_text = item.get('text', merged_item['text'])
            
            if merged_item['type'] == 'merged':
                # 合并的run组，需要智能分配翻译结果
                distribute_translation_to_runs(merged_item['runs'], translated_text)
            else:
                # 单个run，直接替换
                merged_item['runs'][0].text = translated_text
            
            text_count += 1

    return text_count


def distribute_translation_to_runs(runs, translated_text):
    """将翻译结果智能分配回原始run，保持格式"""
    
    # 如果只有一个run，直接替换
    if len(runs) == 1:
        runs[0].text = translated_text
        return
    
    # 计算原始文本的总长度
    original_total_length = sum(len(run.text) for run in runs)
    
    # 如果翻译后文本长度变化不大，按比例分配
    if abs(len(translated_text) - original_total_length) <= 2:
        # 按原始比例分配
        current_pos = 0
        for run in runs:
            original_ratio = len(run.text) / original_total_length
            target_length = int(len(translated_text) * original_ratio)
            
            if current_pos < len(translated_text):
                end_pos = min(current_pos + target_length, len(translated_text))
                run.text = translated_text[current_pos:end_pos]
                current_pos = end_pos
            else:
                run.text = ""
    else:
        # 长度变化较大，尝试按空格和标点分割
        words = translated_text.split()
        if len(words) >= len(runs):
            # 按run数量分配单词
            for i, run in enumerate(runs):
                if i < len(words):
                    run.text = words[i]
                else:
                    run.text = ""
        else:
            # 单词数量少于run数量，平均分配
            chars_per_run = len(translated_text) // len(runs)
            for i, run in enumerate(runs):
                start_pos = i * chars_per_run
                end_pos = start_pos + chars_per_run if i < len(runs) - 1 else len(translated_text)
                run.text = translated_text[start_pos:end_pos]


def update_special_elements(docx_path, texts):
    """更新批注等特殊元素"""
    # 筛选出需要处理的批注
    comment_texts = [t for t in texts if t.get('type') == 'comment' and t.get('complete')]
    if not comment_texts:
        return  # 没有需要处理的批注

    temp_path = f"temp_{int(time.time())}_{os.path.basename(docx_path)}"

    try:
        with zipfile.ZipFile(docx_path, 'r') as zin, \
                zipfile.ZipFile(temp_path, 'w') as zout:

            for item in zin.infolist():
                with zin.open(item) as file:
                    if item.filename == 'word/comments.xml':
                        try:
                            tree = ET.parse(file)
                            root = tree.getroot()

                            # 定义命名空间
                            namespaces = {
                                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                            }

                            # 查找所有批注
                            for comment in root.findall('.//w:comment', namespaces) or root.findall(
                                    './/{*}comment'):
                                # 尝试不同方式获取ID
                                comment_id = None
                                for attr_name in [
                                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id',
                                    'id']:
                                    if attr_name in comment.attrib:
                                        comment_id = comment.attrib[attr_name]
                                        break

                                if not comment_id:
                                    continue

                                # 查找匹配的翻译结果
                                matching_texts = [t for t in comment_texts if
                                                  t.get('comment_id') == comment_id]
                                if not matching_texts:
                                    continue

                                # 更新批注文本
                                text_elements = comment.findall('.//w:t',
                                                                namespaces) or comment.findall(
                                    './/{*}t')
                                for i, t_elem in enumerate(text_elements):
                                    if i < len(matching_texts):
                                        # 找到匹配的原始文本
                                        for match in matching_texts:
                                            if match.get('original_text') == t_elem.text:
                                                t_elem.text = match.get('text', t_elem.text)
                                                break

                            # 写入修改后的XML
                            modified_xml = ET.tostring(root, encoding='utf-8', xml_declaration=True)
                            zout.writestr(item.filename, modified_xml)
                        except Exception as e:
                            print(f"处理批注时出错: {str(e)}")
                            # 如果解析失败，直接复制原文件
                            file.seek(0)
                            zout.writestr(item.filename, file.read())
                    else:
                        # 直接复制其他文件
                        file.seek(0)
                        zout.writestr(item.filename, file.read())

        # 替换原始文件
        os.replace(temp_path, docx_path)
    except Exception as e:
        print(f"更新批注时出错: {str(e)}")
        # 确保临时文件被删除
        if os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except:
                pass


def check_text(text):
    """检查文本是否有效（非空且非纯标点）"""
    return text and len(text) > 0 and not common.is_all_punc(text)


def get_pdf_translate_method():
    """获取PDF翻译方法设置"""
    try:
        from app.models.setting import Setting
        pdf_method_setting = Setting.query.filter_by(
            group='other_setting',
            alias='pdf_translate_method',
            deleted_flag='N'
        ).first()
        return pdf_method_setting.value if pdf_method_setting else 'direct'
    except Exception as e:
        logging.warning(f"获取PDF翻译方法设置失败: {e}")
        return 'direct'  # 默认使用直接翻译


def start_direct_pdf_translation(trans):
    """直接PDF翻译方法"""
    try:
        print("🚀 开始直接PDF翻译流程")
        
        # 更新任务状态为处理中，但不设置初始进度
        try:
            from .to_translate import db
            db.execute("update translate set status='process', process='0' where id=%s", trans['id'])
            print("✅ 已更新任务状态为process，进度0%（开始PDF处理）")
        except Exception as e:
            print(f"⚠️ 更新任务状态失败: {str(e)}")
        
        # 检查文件是否存在
        original_path = Path(trans['file_path'])
        if not original_path.exists():
            print(f"❌ 源文件不存在: {trans['file_path']}")
            to_translate.error(trans['id'], "文件不存在: " + trans['file_path'])
            return False
        
        # 确保目标目录存在
        target_dir = os.path.dirname(trans['target_file'])
        os.makedirs(target_dir, exist_ok=True)
        
        # 检测PDF页数，决定使用哪种翻译方法
        try:
            import fitz
            doc = fitz.open(str(original_path))
            total_pages = doc.page_count
            doc.close()
            print(f"📄 PDF总页数: {total_pages}")
            
            if total_pages > 25:
                print("📊 检测到大文件（超过25页），使用多线程分批处理")
                return start_large_pdf_translation(trans, total_pages)
            else:
                print("📊 检测到小文件（25页以内），使用标准处理")
                return start_small_pdf_translation(trans)
                
        except Exception as e:
            print(f"⚠️ 检测PDF页数失败: {e}，使用标准处理")
            return start_small_pdf_translation(trans)
            
    except Exception as e:
        print(f"❌ 直接PDF翻译异常: {str(e)}")
        to_translate.error(trans['id'], "直接PDF翻译异常: " + str(e))
        return False


def start_small_pdf_translation(trans):
    """小文件PDF翻译方法（25页以内）"""
    try:
        print("🎯 使用小文件翻译方法")
        
        original_path = Path(trans['file_path'])
        
        # 创建翻译函数
        def translate_func(text):
            """翻译函数，使用现有的翻译逻辑"""
            try:
                # 使用现有的翻译逻辑
                from .to_translate import translate_text
                return translate_text(trans, text)
            except Exception as e:
                logging.warning("翻译失败，使用原文: " + str(e))
                return text
        
        # 创建直接PDF翻译器
        translator = DirectPDFTranslator(
            input_pdf_path=str(original_path),
            target_lang=trans.get('lang', 'zh'),  # 使用 'lang' 字段与翻译函数一致
            user_id=trans.get('user_id')  # 传递用户ID用于临时文件隔离
        )
        
        # 执行完整翻译流程
        result_file = translator.run_complete_translation(
            trans=trans,
            output_file=trans['target_file']
        )
        
        if result_file and os.path.exists(result_file):
            print(f"✅ 小文件PDF翻译完成: {result_file}")
            
            # 更新任务状态为完成
            try:
                from .to_translate import db
                db.execute("update translate set status='done', process='100' where id=%s", trans['id'])
                print("✅ 已更新任务状态为done，进度100%")
            except Exception as e:
                print(f"⚠️ 更新任务状态失败: {str(e)}")
            
            return True
        else:
            print(f"❌ 小文件PDF翻译失败")
            to_translate.error(trans['id'], "小文件PDF翻译失败")
            return False

    except Exception as e:
        print(f"❌ 小文件PDF翻译过程出错: {str(e)}")
        traceback.print_exc()
        to_translate.error(trans['id'], "小文件PDF翻译过程出错: " + str(e))
        return False


def start_large_pdf_translation(trans, total_pages):
    """大文件PDF翻译方法（超过20页）"""
    try:
        print("🎯 使用大文件多线程翻译方法")
        print(f"📊 总页数: {total_pages}")
        
        original_path = Path(trans['file_path'])
        
        # 导入大文件翻译器
        from .large_pdf_translator import LargePDFTranslator
        
        # 创建大文件翻译器，使用与小PDF相同的线程配置
        translator = LargePDFTranslator(
            input_pdf_path=str(original_path),
            batch_size=5,  # 减小批次大小，降低内存占用
            max_workers=30,  # 与小PDF保持一致，使用系统默认30线程
            target_lang=trans.get('lang', 'zh'),  # 使用 'lang' 字段与翻译函数一致
            user_id=trans.get('user_id')  # 传递用户ID用于临时文件隔离
        )
        
        # 执行完整翻译流程
        result_file = translator.run_complete_translation(
            trans=trans,
            output_file=trans['target_file']
        )
        
        if result_file and os.path.exists(result_file):
            print(f"✅ 大文件PDF翻译完成: {result_file}")
            
            # 更新任务状态为完成
            try:
                from .to_translate import db
                db.execute("update translate set status='done', process='100' where id=%s", trans['id'])
                print("✅ 已更新任务状态为done，进度100%")
            except Exception as e:
                print(f"⚠️ 更新任务状态失败: {str(e)}")
            
            return True
        else:
            print("❌ 大文件PDF翻译失败")
            to_translate.error(trans['id'], "大文件PDF翻译失败")
            return False

    except Exception as e:
        print(f"❌ 大文件PDF翻译过程出错: {str(e)}")
        traceback.print_exc()
        to_translate.error(trans['id'], "大文件PDF翻译过程出错: " + str(e))
        return False


class DirectPDFTranslator:
    """
    直接PDF翻译器 - 支持中文字符显示，保持原始样式，无背景覆盖
    基于pdf-translator-final项目集成
    """
    
    def __init__(self, input_pdf_path, target_lang='zh', user_id=None):
        self.input_pdf_path = input_pdf_path
        self.target_lang = target_lang
        self.user_id = user_id
        self.doc = None
        self.extracted_texts = []
        
    def step1_split_pdf(self, output_dir):
        """步骤1: 拆分PDF为文本JSON和无文本PDF"""
        print("=" * 60)
        print("步骤1: 拆分PDF为文本JSON和无文本PDF")
        print("=" * 60)
        
        try:
            # 1. 打开PDF
            print("1. 打开PDF...")
            self.doc = fitz.open(self.input_pdf_path)
            print(f"   打开了 {self.doc.page_count} 页的PDF")
            
            # 2. 提取文本信息
            print("\n2. 提取文本信息...")
            self.extracted_texts = []
            
            for page_num in range(self.doc.page_count):
                page = self.doc[page_num]
                print(f"   处理第 {page_num + 1} 页...")
                
                # 提取文本块
                text_blocks = page.get_text("dict", flags=fitz.TEXTFLAGS_TEXT)["blocks"]
                page_texts = []
                
                for block in text_blocks:
                    if "lines" in block:
                        for line in block["lines"]:
                            for span in line["spans"]:
                                text = span["text"].strip()
                                if text:
                                    text_info = {
                                        "text": text,
                                        "bbox": span["bbox"],
                                        "size": span["size"],
                                        "color": span["color"],
                                        "font": span["font"]
                                    }
                                    page_texts.append(text_info)
                
                page_data = {
                    "page_number": page_num,
                    "texts": page_texts
                }
                self.extracted_texts.append(page_data)
                print(f"   提取了 {len(page_texts)} 个文本块")
            
            # 3. 保存提取的文本
            print("\n3. 保存提取的文本...")
            extracted_texts_file = os.path.join(output_dir, "extracted_texts.json")
            with open(extracted_texts_file, 'w', encoding='utf-8') as f:
                json.dump(self.extracted_texts, f, ensure_ascii=False, indent=2)
            print(f"✅ 提取的文本已保存到: {extracted_texts_file}")
            
            # 4. 创建无文本PDF
            print("\n4. 创建无文本PDF...")
            no_text_doc = fitz.open()
            
            for page_num in range(self.doc.page_count):
                page = self.doc[page_num]
                new_page = no_text_doc.new_page(width=page.rect.width, height=page.rect.height)
                
                # 复制页面内容（图片、背景等）
                new_page.show_pdf_page(page.rect, self.doc, page_num)
                
                # 使用精确方法删除所有文本
                text_blocks = page.get_text("dict", flags=fitz.TEXTFLAGS_TEXT)["blocks"]
                text_count = 0
                
                for block in text_blocks:
                    if "lines" in block:
                        for line in block["lines"]:
                            for span in line["spans"]:
                                bbox = span["bbox"]
                                text = span["text"].strip()
                                if text:
                                    try:
                                        # 使用add_redact_annot但不填充，避免白色遮挡
                                        redact_annot = new_page.add_redact_annot(bbox, fill=None)
                                        text_count += 1
                                    except Exception as e:
                                        logging.warning(f"删除文本失败: {e}")
                                        # 如果失败，尝试用透明填充
                                        try:
                                            redact_annot = new_page.add_redact_annot(bbox, fill=(0, 0, 0, 0))
                                            text_count += 1
                                        except Exception as e2:
                                            logging.warning(f"透明填充也失败: {e2}")
                
                # 应用删除操作
                try:
                    new_page.apply_redactions()
                    print(f"   第 {page_num + 1} 页删除了 {text_count} 个文本块")
                except Exception as e:
                    logging.warning(f"应用删除操作失败: {e}")
                    print(f"   ⚠️ 第 {page_num + 1} 页应用删除操作失败: {e}")
            
            no_text_pdf_file = os.path.join(output_dir, "no_text.pdf")
            no_text_doc.save(no_text_pdf_file)
            no_text_doc.close()
            print(f"✅ 无文本PDF已保存到: {no_text_pdf_file}")
            
            return extracted_texts_file, no_text_pdf_file
            
        except Exception as e:
            logging.error(f"拆分PDF时出错: {e}")
            raise
    
    def step2_translate_texts(self, extracted_texts_file, trans, output_dir):
        """步骤2: 使用多线程翻译JSON中的文本"""
        print("\n" + "=" * 60)
        print("步骤2: 使用多线程翻译JSON中的文本")
        print("=" * 60)
        
        try:
            # 1. 加载提取的文本
            print("1. 加载提取的文本...")
            with open(extracted_texts_file, 'r', encoding='utf-8') as f:
                extracted_texts = json.load(f)
            
            print("   加载了 " + str(len(extracted_texts)) + " 页的文本数据")
            
            # 2. 准备多线程翻译数据
            print("\n2. 准备多线程翻译数据...")
            texts_for_translation = []
            text_mapping = {}  # 用于映射翻译结果回原始位置
            
            for page_idx, page_data in enumerate(extracted_texts):
                page_num = page_data["page_number"]
                for text_idx, text_info in enumerate(page_data["texts"]):
                    original_text = text_info["text"]
                    if original_text and original_text.strip():
                        # 创建翻译任务
                        translation_task = {
                            'text': original_text,
                            'complete': False,
                            'page_idx': page_idx,
                            'text_idx': text_idx,
                            'original_info': text_info
                        }
                        texts_for_translation.append(translation_task)
                        text_mapping[(page_idx, text_idx)] = translation_task
            
            print("   准备翻译 " + str(len(texts_for_translation)) + " 个文本片段")
            
            # 3. 使用多线程翻译
            print("\n3. 开始多线程翻译...")
            if texts_for_translation:
                # 使用现有的多线程翻译系统
                run_translation(trans, texts_for_translation, max_threads=30)
                print("   多线程翻译完成")
            else:
                print("   没有需要翻译的文本")
            
            # 4. 重新组织翻译结果
            print("\n4. 重新组织翻译结果...")
            translated_texts = []
            
            for page_data in extracted_texts:
                page_num = page_data["page_number"]
                translated_page_data = {"page_number": page_num, "texts": []}
                
                for text_idx, text_info in enumerate(page_data["texts"]):
                    original_text = text_info["text"]
                    if original_text and original_text.strip():
                        # 获取翻译结果
                        translation_task = text_mapping.get((extracted_texts.index(page_data), text_idx))
                        if translation_task and translation_task.get('complete'):
                            translated_text = translation_task.get('text', original_text)
                            print("   ✅ 翻译: '" + original_text[:20] + "...' -> '" + translated_text[:20] + "...'")
                        else:
                            translated_text = original_text
                            print("   ⚠️ 翻译失败，使用原文: '" + original_text[:20] + "...'")
                        
                        # 创建翻译后的文本信息
                        translated_text_info = text_info.copy()
                        translated_text_info["text"] = translated_text
                        translated_text_info["original_text"] = original_text
                        translated_page_data["texts"].append(translated_text_info)
                    else:
                        translated_page_data["texts"].append(text_info)
                
                translated_texts.append(translated_page_data)
            
            # 5. 保存翻译后的文本
            print("\n5. 保存翻译后的文本...")
            translated_texts_file = os.path.join(output_dir, "translated_texts.json")
            with open(translated_texts_file, 'w', encoding='utf-8') as f:
                json.dump(translated_texts, f, ensure_ascii=False, indent=2)
            print("✅ 翻译后的文本已保存到: " + translated_texts_file)
            
            return translated_texts_file
            
        except Exception as e:
            logging.error("翻译文本时出错: " + str(e))
            raise
    
    def step3_fill_translated_texts(self, translated_texts_file, no_text_pdf_file, output_file):
        """步骤3: 使用insert_htmlbox回填翻译后的文本"""
        print("\n" + "=" * 60)
        print("步骤3: 使用insert_htmlbox回填翻译后的文本")
        print("=" * 60)
        
        try:
            # 1. 加载翻译后的文本
            print("1. 加载翻译后的文本...")
            with open(translated_texts_file, 'r', encoding='utf-8') as f:
                translated_texts = json.load(f)
            
            print(f"   加载了 {len(translated_texts)} 页的翻译文本数据")
            
            # 2. 打开无文本PDF
            print("\n2. 打开无文本PDF...")
            doc = fitz.open(no_text_pdf_file)
            print(f"   打开了 {doc.page_count} 页的PDF")
            
            # 3. 回填翻译文本
            print("\n3. 回填翻译文本...")
            
            for page_data in translated_texts:
                page_num = page_data["page_number"]
                page = doc[page_num]
                print(f"   处理第 {page_num + 1} 页...")
                
                for text_info in page_data["texts"]:
                    text = text_info["text"]
                    if text and text.strip():
                        bbox = text_info["bbox"]
                        font_size = text_info["size"]
                        color = text_info["color"]
                        
                        # 标准化颜色
                        if isinstance(color, (int, float)):
                            if color == 0:
                                color = (0, 0, 0)  # 黑色
                            elif color == 16777215:
                                color = (1, 1, 1)  # 白色
                            else:
                                # 转换为RGB
                                r = ((color >> 16) & 0xFF) / 255.0
                                g = ((color >> 8) & 0xFF) / 255.0
                                b = (color & 0xFF) / 255.0
                                color = (r, g, b)
                        elif isinstance(color, (list, tuple)) and len(color) >= 3:
                            color = tuple(color[:3])
                        else:
                            color = (0, 0, 0)  # 默认黑色
                        
                        # 使用insert_htmlbox方法（支持中文，避免背景覆盖）
                        try:
                            # 创建文本框
                            textbox = fitz.Rect(bbox[0], bbox[1], bbox[2], bbox[3])
                            
                            # 计算文本长度和box宽度，决定是否需要换行
                            box_width = bbox[2] - bbox[0]
                            box_height = bbox[3] - bbox[1]
                            
                            # 估算每行字符数（根据字体大小）
                            chars_per_line = max(1, int(box_width / (font_size * 0.6)))  # 0.6是经验值

                            wrapped_text = text
                            
                            # 构建HTML文本，使用完全透明的背景
                            html_text = f"""
                            <div style="
                                font-size: {font_size}px;
                                color: rgb({int(color[0]*255)}, {int(color[1]*255)}, {int(color[2]*255)});
                                font-family: sans-serif;
                                line-height: 1.0;
                                margin: 0;
                                padding: 0;
                                background: none;
                                background-color: transparent;
                                background-image: none;
                                border: none;
                                outline: none;
                                box-shadow: none;
                            ">
                                {wrapped_text}
                            </div>
                            """
                            
                            # 使用insert_htmlbox插入文本
                            page.insert_htmlbox(textbox, html_text)
                            logging.info(f"✅ 文本插入成功: '{text[:20]}...'")
                            print(f"   ✅ 文本插入成功: '{text[:20]}...'")
                        except Exception as e:
                            logging.error(f"文本插入失败: {e}")
                            print(f"   ❌ 文本插入失败: '{text[:20]}...' - {e}")
                    else:
                        logging.warning(f"文本为空，跳过: '{text}'")
            
            # 4. 保存最终PDF
            print("\n4. 保存最终PDF...")
            doc.save(output_file)
            doc.close()
            print(f"✅ 最终翻译后的PDF已保存到: {output_file}")
            
            return output_file
            
        except Exception as e:
            logging.error(f"回填PDF时出错: {e}")
            raise
    
    def run_complete_translation(self, trans, output_file):
        """运行完整的翻译流程"""
        print("🚀 开始完整的PDF翻译流程")
        print("=" * 60)
        
        # 临时文件列表，用于最后清理
        temp_files = []
        
        try:
            # 创建输出目录
            output_dir = os.path.dirname(output_file)
            os.makedirs(output_dir, exist_ok=True)
            
            # 为每个翻译任务创建唯一的临时目录，避免批量翻译时文件冲突
            import uuid
            if self.user_id:
                # 使用用户ID创建隔离目录
                temp_dir_name = f"temp_user_{self.user_id}_{uuid.uuid4().hex[:8]}"
            else:
                # 如果没有用户ID，使用默认方式
                temp_dir_name = f"temp_{uuid.uuid4().hex[:8]}"
            temp_dir = os.path.join(output_dir, temp_dir_name)
            os.makedirs(temp_dir, exist_ok=True)
            print(f"📁 创建临时目录: {temp_dir}")
            print(f"👤 用户ID: {self.user_id if self.user_id else 'N/A'}")
            
            # 步骤1: 拆分PDF
            extracted_texts_file, no_text_pdf_file = self.step1_split_pdf(temp_dir)
            temp_files.extend([extracted_texts_file, no_text_pdf_file])
            
            # 步骤2: 翻译文本
            translated_texts_file = self.step2_translate_texts(extracted_texts_file, trans, temp_dir)
            temp_files.append(translated_texts_file)
            
            # 步骤3: 回填翻译文本
            final_pdf_file = self.step3_fill_translated_texts(translated_texts_file, no_text_pdf_file, output_file)
            
            # 步骤4: 压缩PDF文件
            print("\n" + "=" * 60)
            print("步骤4: 压缩PDF文件")
            print("=" * 60)
            
            # 生成压缩后的文件名
            base_name = os.path.splitext(output_file)[0]
            compressed_pdf_file = base_name + "_compressed.pdf"
            
            # 压缩PDF
            optimized_pdf_file = self._optimize_pdf_size(final_pdf_file, compressed_pdf_file)
            
            if optimized_pdf_file and os.path.exists(optimized_pdf_file):
                # 删除原始合成PDF，只保留压缩后的PDF
                try:
                    os.remove(final_pdf_file)
                    print("✅ 已删除原始合成PDF: " + os.path.basename(final_pdf_file))
                except Exception as e:
                    logging.warning("删除原始PDF失败: " + str(e))
                    print("⚠️ 删除原始PDF失败: " + str(e))
                
                # 将压缩后的PDF重命名为最终输出文件
                try:
                    os.rename(optimized_pdf_file, output_file)
                    print("✅ 压缩后的PDF已重命名为最终输出文件")
                    final_pdf_file = output_file
                except Exception as e:
                    logging.warning("重命名压缩PDF失败: " + str(e))
                    print("⚠️ 重命名压缩PDF失败: " + str(e))
                    final_pdf_file = optimized_pdf_file
            else:
                print("⚠️ PDF压缩失败，使用原始PDF")
                final_pdf_file = final_pdf_file
            
            # 删除原始上传文件，保留翻译后的文件（带UUID后缀）
            original_file = self.input_pdf_path
            if os.path.exists(original_file) and original_file != output_file:
                try:
                    os.remove(original_file)
                    print("✅ 已删除原始上传文件: " + os.path.basename(original_file))
                except Exception as e:
                    logging.warning("删除原始上传文件失败: " + str(e))
                    print("⚠️ 删除原始上传文件失败: " + str(e))
            
            # 确保输出文件保持UUID后缀，避免多次翻译冲突
            # 不重命名文件，保持数据库路径与实际文件一致
            
            print("\n" + "=" * 60)
            print("🎉 完整PDF翻译流程完成!")
            print("=" * 60)
            print(f"📄 输入文件: {self.input_pdf_path}")
            print(f"📝 提取文本: {extracted_texts_file}")
            print(f"🔄 翻译文本: {translated_texts_file}")
            print(f"📄 无文本PDF: {no_text_pdf_file}")
            print(f"🎯 最终输出: {final_pdf_file}")
            print("=" * 60)
            
            # 清理临时文件和目录
            self._cleanup_temp_files(temp_files, temp_dir)
            
            return final_pdf_file
            
        except Exception as e:
            logging.error(f"完整翻译流程失败: {e}")
            # 即使出错也要清理临时文件和目录
            self._cleanup_temp_files(temp_files, temp_dir)
            raise
        finally:
            if self.doc:
                self.doc.close()
    
    def _optimize_pdf_size(self, input_pdf_path, output_pdf_path=None):
        """优化PDF文件大小"""
        print("\n" + "=" * 60)
        print("PDF文件大小优化")
        print("=" * 60)
        
        if not os.path.exists(input_pdf_path):
            print("❌ 文件不存在: " + input_pdf_path)
            return None
        
        # 获取原始文件大小
        original_size = os.path.getsize(input_pdf_path)
        print("原始文件大小: " + str(original_size) + " 字节 (" + str(original_size/1024/1024) + " MB)")
        
        doc = None
        doc2 = None
        try:
            # 打开PDF
            print("\n1. 打开PDF...")
            doc = fitz.open(input_pdf_path)
            print("   打开了 " + str(doc.page_count) + " 页的PDF")
            
            # 设置输出文件名
            if output_pdf_path is None:
                base_name = os.path.splitext(input_pdf_path)[0]
                output_pdf_path = base_name + "_optimized.pdf"
            
            # 优化选项
            print("\n2. 应用优化选项...")
            
            # 方法1: 使用压缩选项保存
            print("   方法1: 使用压缩选项...")
            doc.save(
                output_pdf_path,
                garbage=4,        # 垃圾回收
                deflate=True,     # 压缩
                clean=True,       # 清理
                encryption=fitz.PDF_ENCRYPT_NONE  # 无加密
            )
            
            # 检查优化后的文件大小
            optimized_size = os.path.getsize(output_pdf_path)
            print("   优化后文件大小: " + str(optimized_size) + " 字节 (" + str(optimized_size/1024/1024) + " MB)")
            
            # 计算压缩率
            compression_ratio = (1 - optimized_size / original_size) * 100
            print("   压缩率: " + str(compression_ratio) + "%")
            
            # 方法2: 如果还是太大，尝试更激进的优化
            if optimized_size > 800000:  # 如果还是超过800KB
                print("\n   方法2: 应用更激进的优化...")
                aggressive_output = os.path.splitext(output_pdf_path)[0] + "_aggressive.pdf"
                
                # 重新打开并应用更激进的优化
                doc2 = fitz.open(input_pdf_path)
                doc2.save(
                    aggressive_output,
                    garbage=4,
                    deflate=True,
                    clean=True,
                    encryption=fitz.PDF_ENCRYPT_NONE,
                    # 更激进的优化选项
                    ascii=False,      # 二进制模式
                    expand=0,         # 不展开
                    no_new_id=True,   # 不生成新ID
                )
                
                # 检查激进优化后的文件大小
                aggressive_size = os.path.getsize(aggressive_output)
                print("   激进优化后文件大小: " + str(aggressive_size) + " 字节 (" + str(aggressive_size/1024/1024) + " MB)")
                
                aggressive_ratio = (1 - aggressive_size / original_size) * 100
                print("   激进压缩率: " + str(aggressive_ratio) + "%")
                
                # 选择更小的文件
                if aggressive_size < optimized_size:
                    print("   ✅ 激进优化效果更好，使用: " + aggressive_output)
                    output_pdf_path = aggressive_output
                    optimized_size = aggressive_size
            
            print("\n✅ 优化完成! 输出文件: " + output_pdf_path)
            print("📊 文件大小对比:")
            print("   原始: " + str(original_size) + " 字节 (" + str(original_size/1024/1024) + " MB)")
            print("   优化: " + str(optimized_size) + " 字节 (" + str(optimized_size/1024/1024) + " MB)")
            print("   节省: " + str(original_size - optimized_size) + " 字节 (" + str((1 - optimized_size/original_size)*100) + "%)")
            
            return output_pdf_path
            
        except Exception as e:
            logging.error("优化PDF时出错: " + str(e))
            print("❌ 优化PDF时出错: " + str(e))
            raise
        finally:
            # 确保所有文档对象被正确关闭
            if doc2 is not None:
                try:
                    doc2.close()
                    logging.debug("激进优化PDF文档已关闭")
                except Exception as close_error:
                    logging.warning(f"关闭激进优化PDF文档时出错: {close_error}")
            
            if doc is not None:
                try:
                    doc.close()
                    logging.debug("优化PDF文档已关闭")
                except Exception as close_error:
                    logging.warning(f"关闭优化PDF文档时出错: {close_error}")

    def _cleanup_temp_files(self, temp_files, temp_dir=None):
        """清理临时文件和目录"""
        print("\n🧹 开始清理临时文件...")
        cleaned_count = 0
        
        # 清理临时文件
        for temp_file in temp_files:
            if temp_file and os.path.exists(temp_file):
                try:
                    os.remove(temp_file)
                    print("✅ 已删除临时文件: " + os.path.basename(temp_file))
                    cleaned_count += 1
                except Exception as e:
                    logging.warning("删除临时文件失败: " + temp_file + " - " + str(e))
                    print("⚠️ 删除临时文件失败: " + os.path.basename(temp_file) + " - " + str(e))
        
        # 清理临时目录
        if temp_dir and os.path.exists(temp_dir):
            try:
                import shutil
                shutil.rmtree(temp_dir)
                print("✅ 已删除临时目录: " + os.path.basename(temp_dir))
                cleaned_count += 1
            except Exception as e:
                logging.warning("删除临时目录失败: " + temp_dir + " - " + str(e))
                print("⚠️ 删除临时目录失败: " + os.path.basename(temp_dir) + " - " + str(e))
        
        print("🧹 临时文件清理完成，共清理 " + str(cleaned_count) + " 个文件/目录")
