# -*- coding: utf-8 -*-
"""
Okapi Framework 集成模块

使用 Okapi Framework 进行 Word 文档翻译，解决 run 切割过细的问题。
通过 XLIFF 格式保持完整格式和上下文。

作者：Claude
版本：1.0.0
"""

import os
import subprocess
import tempfile
import shutil
import logging
import time
import xml.etree.ElementTree as ET
from contextlib import contextmanager
from typing import Optional, Dict, Any, List
import zipfile
import re

# 配置日志
logger = logging.getLogger(__name__)

def clean_xliff_for_translation(text: str) -> tuple[str, dict]:
    """
    清理 XLIFF 文本中的格式标签，为翻译做准备
    只清理明显的格式标签，保留Okapi的标准XLIFF标签
    
    Args:
        text: 包含格式标签的XLIFF文本
        
    Returns:
        tuple: (清理后的文本, 代码映射字典)
    """
    # 需要清理的标签模式（这些是Word文档的原始格式标签）
    format_patterns = [
        r'<run\d+>', r'</run\d+>',  # Word运行标签
        r'<tags\d+/>',              # Word格式标签
        r'<lendof\d+\|>',           # 段落结束标签
        r'<[a-z]+\d*[^>]*/>',       # 其他自闭合标签
    ]
    
    # 保留的Okapi标准XLIFF标签（这些不应该被清理）
    # <bpt>, <ept>, <ph>, <g>, <x> 等是Okapi的标准标签
    
    combined_pattern = '|'.join(format_patterns)
    matches = re.findall(combined_pattern, text)
    
    # 创建代码映射
    code_mapping = {}
    clean_text = text
    
    for i, match in enumerate(matches):
        placeholder = f"__FORMAT_TAG_{i}__"
        code_mapping[placeholder] = match
        clean_text = clean_text.replace(match, placeholder, 1)
    
    logger.info(f"清理格式标签: {len(matches)} 个标签")
    logger.info(f"原文: {text[:100]}...")
    logger.info(f"清理后: {clean_text[:100]}...")
    
    return clean_text, code_mapping

def restore_xliff_tags(translated_text: str, code_mapping: dict) -> str:
    """
    恢复格式标签到翻译后的文本中
    
    Args:
        translated_text: 翻译后的文本
        code_mapping: 代码映射字典
        
    Returns:
        str: 恢复格式标签后的文本
    """
    restored_text = translated_text
    
    for placeholder, original_code in code_mapping.items():
        # 将占位符替换回原始的格式标签
        restored_text = restored_text.replace(placeholder, original_code)
    
    logger.info(f"恢复格式标签: {len(code_mapping)} 个标签")
    logger.info(f"翻译文本: {translated_text[:100]}...")
    logger.info(f"恢复后: {restored_text[:100]}...")
    
    return restored_text

class OkapiIntegrationError(Exception):
    """Okapi 集成错误"""
    pass

class DockerOkapiIntegration:
    """Docker 环境中的 Okapi 集成类"""
    
    def __init__(self, okapi_home: str = "/opt/okapi"):
        """
        初始化 Okapi 集成
        
        Args:
            okapi_home: Okapi 安装目录
        """
        self.okapi_home = okapi_home
        
        # 获取 Tikal 脚本
        self.tikal_script = self._get_tikal_script()
        
        logger.info(f"Okapi 集成初始化完成: {okapi_home}")
        # 调试：显示所有找到的文件
        self._debug_okapi_installation()
    
    def _debug_okapi_installation(self):
        """调试 Okapi 安装"""
        import glob
        
        logger.info(f"=== Okapi 安装调试信息 ===")
        logger.info(f"Okapi 目录: {self.okapi_home}")
        logger.info(f"目录是否存在: {os.path.exists(self.okapi_home)}")
        
        if os.path.exists(self.okapi_home):
            # 列出目录内容
            try:
                dir_contents = os.listdir(self.okapi_home)
                logger.info(f"目录内容: {dir_contents}")
                
                # 查找所有 JAR 文件
                all_jars = glob.glob(os.path.join(self.okapi_home, "**", "*.jar"), recursive=True)
                logger.info(f"所有 JAR 文件: {all_jars}")
                
                # 查找所有可执行文件
                all_executables = glob.glob(os.path.join(self.okapi_home, "**", "*"), recursive=True)
                executables = [f for f in all_executables if os.path.isfile(f) and os.access(f, os.X_OK)]
                logger.info(f"可执行文件: {executables}")
                
            except Exception as e:
                logger.error(f"读取目录内容失败: {e}")
        
        logger.info("=== 调试信息结束 ===")
    
    def _get_tikal_script(self) -> str:
        """获取 Tikal 脚本路径"""
        tikal_script = os.path.join(self.okapi_home, "tikal.sh")
        
        if not os.path.exists(tikal_script):
            raise OkapiIntegrationError(f"Tikal 脚本不存在: {tikal_script}")
        
        if not os.access(tikal_script, os.X_OK):
            raise OkapiIntegrationError(f"Tikal 脚本无执行权限: {tikal_script}")
        
        logger.info(f"使用 Tikal 脚本: {tikal_script}")
        return tikal_script
    
    def _verify_java_environment(self) -> bool:
        """验证 Java 环境是否可用"""
        try:
            result = subprocess.run(
                ["java", "-version"],
                capture_output=True,
                timeout=10
            )
            if result.returncode != 0:
                raise OkapiIntegrationError("Java 不可用")
            
            logger.info("Java 环境验证通过")
            return True
            
        except Exception as e:
            raise OkapiIntegrationError(f"Java 环境验证失败: {e}")
    
    @contextmanager
    def temp_workspace(self):
        """创建临时工作空间"""
        temp_dir = tempfile.mkdtemp(prefix="okapi_work_")
        try:
            logger.debug(f"创建临时工作空间: {temp_dir}")
            yield temp_dir
        finally:
            # 自动清理临时文件
            shutil.rmtree(temp_dir, ignore_errors=True)
            logger.debug(f"清理工作空间: {temp_dir}")
    
    def extract_to_xliff(self, input_file: str, output_xliff: str, 
                        source_lang: str = "zh", target_lang: str = "en",
                        java_opts: Optional[str] = None) -> bool:
        """
        使用 Tikal 提取 Word 文档到 XLIFF
        
        根据 Tikal 文档，提取会生成 input_file.xlf 文件
        
        Args:
            input_file: 输入的 Word 文档路径
            output_xliff: 输出的 XLIFF 文件路径
            source_lang: 源语言代码
            target_lang: 目标语言代码
            java_opts: Java 选项
            
        Returns:
            bool: 是否成功
        """
        try:
            # 使用 Tikal 脚本 - 默认生成 XLIFF 1.2，但我们会处理为 2.0
            cmd = [self.tikal_script, "-x", input_file, "-sl", source_lang, "-tl", target_lang]
            logger.info(f"使用 Tikal 脚本执行提取")
            logger.info(f"执行提取命令: {' '.join(cmd)}")
            
            # 记录执行时间
            start_time = time.time()
            
            # 执行进程
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=120,  # 2分钟超时
                cwd=os.path.dirname(input_file)  # 在输入文件所在目录执行
            )
            
            duration = time.time() - start_time
            logger.info(f"Tikal 提取完成，用时: {duration:.2f}秒")
            
            # 详细的结果处理
            if result.returncode == 0:
                # 根据 Tikal 文档，输出文件为 input_file.xlf
                expected_xliff = f"{input_file}.xlf"
                
                if os.path.exists(expected_xliff):
                    # 复制到指定位置
                    shutil.copy2(expected_xliff, output_xliff)
                    logger.info(f"✅ Tikal 提取成功: {output_xliff}")
                    if result.stdout:
                        logger.debug(f"STDOUT: {result.stdout}")
                    return True
                else:
                    logger.error(f"❌ 预期的 XLIFF 文件不存在: {expected_xliff}")
                    return False
            else:
                logger.error(f"❌ Tikal 提取失败 (返回码: {result.returncode})")
                logger.error(f"STDOUT: {result.stdout}")
                logger.error(f"STDERR: {result.stderr}")
                return False
                
        except subprocess.TimeoutExpired:
            logger.error("❌ Tikal 执行超时")
            return False
        except FileNotFoundError:
            logger.error("❌ Java 或 Tikal 未找到")
            return False
        except Exception as e:
            logger.error(f"❌ 执行过程中出错: {e}")
            return False
    
    def merge_from_xliff(self, original_file: str, xliff_file: str, 
                        output_file: str, java_opts: Optional[str] = None) -> bool:
        """
        将翻译后的 XLIFF 合并回 Word 文档
        
        根据 Tikal 文档，合并时需要：
        1. XLIFF 文件名为：original_file.xlf
        2. 原始文件在同一目录
        3. 输出文件为：original_file.out.docx
        
        Args:
            original_file: 原始 Word 文档路径
            xliff_file: 翻译后的 XLIFF 文件路径
            output_file: 输出的 Word 文档路径
            java_opts: Java 选项
            
        Returns:
            bool: 是否成功
        """
        try:
            # 确保原始文件在 XLIFF 文件的同一目录
            xliff_dir = os.path.dirname(xliff_file)
            original_filename = os.path.basename(original_file)
            original_in_xliff_dir = os.path.join(xliff_dir, original_filename)
            
            # 复制原始文件到 XLIFF 目录（如果不在同一目录）
            if original_file != original_in_xliff_dir:
                shutil.copy2(original_file, original_in_xliff_dir)
                logger.info(f"复制原始文件到 XLIFF 目录: {original_in_xliff_dir}")
            
            # 根据 Tikal 官方文档，XLIFF 文件名应该是原始文件名加上 .xlf 扩展名
            # 例如：myFile.html -> myFile.html.xlf
            correct_xliff_name = f"{original_filename}.xlf"
            correct_xliff_path = os.path.join(xliff_dir, correct_xliff_name)
            
            # 如果当前 XLIFF 文件名不正确，重命名它
            if os.path.basename(xliff_file) != correct_xliff_name:
                shutil.copy2(xliff_file, correct_xliff_path)
                logger.info(f"重命名 XLIFF 文件为: {correct_xliff_path}")
                xliff_file = correct_xliff_path
            else:
                # 如果文件名已经正确，使用当前路径
                xliff_file = correct_xliff_path
            
            # 确保文件确实存在且有正确的扩展名
            if not os.path.exists(xliff_file):
                logger.error(f"XLIFF 文件不存在: {xliff_file}")
                return False
            
            # 检查文件扩展名
            if not xliff_file.endswith('.xlf'):
                logger.error(f"文件扩展名不正确: {xliff_file}")
                return False
            
            logger.info(f"使用 XLIFF 文件进行合并: {xliff_file}")
            
            # 列出目录中的所有文件进行调试
            logger.info(f"目录 {xliff_dir} 中的文件:")
            for f in os.listdir(xliff_dir):
                logger.info(f"  - {f}")
            
            # 使用 Tikal 脚本 - 根据官方文档，合并命令不需要 -f 参数
            cmd = [self.tikal_script, "-m", xliff_file, "-trace"]
            logger.info(f"使用 Tikal 脚本执行合并")
            
            logger.info(f"执行合并命令: {' '.join(cmd)}")
            
            # 记录执行时间
            start_time = time.time()
            
            # 执行脚本
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=120,
                cwd=xliff_dir  # 在 XLIFF 文件所在目录执行
            )
            
            duration = time.time() - start_time
            logger.info(f"Tikal 合并完成，用时: {duration:.2f}秒")
            
            if result.returncode == 0:
                # 检查目录中生成的所有文件
                logger.info(f"Tikal 合并成功，检查输出文件...")
                logger.info(f"目录 {xliff_dir} 中的所有文件:")
                for f in os.listdir(xliff_dir):
                    logger.info(f"  - {f}")
                
                # 查找可能的输出文件
                # 根据原始文件扩展名确定输出文件扩展名
                original_ext = os.path.splitext(original_filename)[1].lower()
                if original_ext == '.pptx':
                    output_ext = '.pptx'
                    expected_pattern = '.out.pptx'
                elif original_ext in ['.docx', '.doc']:
                    output_ext = '.docx'
                    expected_pattern = '.out.docx'
                else:
                    # 默认使用原始文件扩展名
                    output_ext = original_ext
                    expected_pattern = f'.out{original_ext}'
                
                possible_outputs = []
                for f in os.listdir(xliff_dir):
                    # 查找匹配的输出文件（.out.扩展名）
                    if f.endswith(expected_pattern) and f != original_filename:
                        possible_outputs.append(f)
                    # 也查找同扩展名的文件（兼容其他可能的命名方式）
                    elif f.endswith(output_ext) and f != original_filename and '.out.' in f:
                        possible_outputs.append(f)
                
                if possible_outputs:
                    # 使用找到的第一个输出文件
                    output_filename = possible_outputs[0]
                    actual_output = os.path.join(xliff_dir, output_filename)
                    shutil.copy2(actual_output, output_file)
                    logger.info(f"✅ XLIFF 合并成功: {output_file}")
                    logger.info(f"实际输出文件: {actual_output}")
                    if result.stdout:
                        logger.debug(f"STDOUT: {result.stdout}")
                    return True
                else:
                    logger.error(f"❌ 未找到输出文件")
                    logger.error(f"期望的文件: {original_filename}{expected_pattern}")
                    logger.error(f"目录中的所有文件:")
                    for f in os.listdir(xliff_dir):
                        logger.error(f"  - {f}")
                    return False
            else:
                logger.error(f"❌ XLIFF 合并失败: {result.stderr}")
                logger.error(f"STDOUT: {result.stdout}")
                return False
                
        except Exception as e:
            logger.error(f"❌ 合并过程出错: {e}")
            return False
    
    def parse_xliff_content(self, xliff_file: str) -> List[Dict[str, Any]]:
        """
        解析 XLIFF 文件内容，提取需要翻译的文本
        
        Args:
            xliff_file: XLIFF 文件路径
            
        Returns:
            List[Dict]: 翻译单元列表
        """
        try:
            logger.info(f"开始解析 XLIFF 文件: {xliff_file}")
            
            # 检查文件是否存在
            if not os.path.exists(xliff_file):
                logger.error(f"XLIFF 文件不存在: {xliff_file}")
                return []
            
            # 读取文件内容进行调试
            with open(xliff_file, 'r', encoding='utf-8') as f:
                content = f.read()
                logger.info(f"XLIFF 文件大小: {len(content)} 字符")
                logger.info(f"XLIFF 文件内容前1000字符: {content[:1000]}")
                
                # 检查是否包含 trans-unit 或 unit
                if 'trans-unit' in content:
                    logger.info("文件中包含 trans-unit 标签")
                if 'unit' in content:
                    logger.info("文件中包含 unit 标签")
                if 'source' in content:
                    logger.info("文件中包含 source 标签")
                if 'target' in content:
                    logger.info("文件中包含 target 标签")
            
            tree = ET.parse(xliff_file)
            root = tree.getroot()
            
            logger.info(f"XLIFF 根元素: {root.tag}")
            
            # 支持 XLIFF 1.2 和 2.0 格式
            translation_units = []
            
            # 检查 XLIFF 版本
            if 'urn:oasis:names:tc:xliff:document:2.0' in root.tag:
                logger.info("解析 XLIFF 2.0 格式")
                namespaces = {'xliff': 'urn:oasis:names:tc:xliff:document:2.0'}
                unit_xpath = './/xliff:unit'
                source_xpath = './/xliff:source'
                target_xpath = './/xliff:target'
            else:
                logger.info("解析 XLIFF 1.2 格式")
                namespaces = {'xliff': 'urn:oasis:names:tc:xliff:document:1.2'}
                unit_xpath = './/xliff:trans-unit'
                source_xpath = './/xliff:source'
                target_xpath = './/xliff:target'
            
            units = root.findall(unit_xpath, namespaces)
            logger.info(f"找到 {len(units)} 个翻译单元")
            
            for unit in units:
                unit_id = unit.get('id', '')
                
                source_elem = unit.find(source_xpath, namespaces)
                target_elem = unit.find(target_xpath, namespaces)
                
                if source_elem is not None:
                    # 只提取<bpt>和<ept>标签之间的纯文本内容
                    # 不提取<bpt>和<ept>标签内部的格式信息
                    text_parts = []
                    
                    # 添加source元素开头的直接文本（如果有）
                    if source_elem.text:
                        text_parts.append(source_elem.text)
                    
                    # 遍历所有子元素，只提取实际文本内容
                    for child in source_elem:
                        # 跳过<bpt>和<ept>标签本身，但提取它们后面的tail文本
                        if child.tail:
                            text_parts.append(child.tail)
                    
                    source_text = ''.join(text_parts).strip()
                    
                    if source_text and len(source_text) > 0:
                        translation_units.append({
                            'id': unit_id,
                            'source': source_text,  # 纯文本用于翻译
                            'target': ''.join(target_elem.itertext()).strip() if target_elem is not None else '',
                            'unit_element': unit
                        })
            
            logger.info(f"解析到 {len(translation_units)} 个翻译单元")
            return translation_units
            
        except Exception as e:
            logger.error(f"解析 XLIFF 文件失败: {e}")
            return []

    def parse_xliff_with_placeholders(self, xliff_file: str) -> List[Dict[str, Any]]:
        """
        解析 XLIFF 文件内容，使用占位符标记格式边界
        
        Args:
            xliff_file: XLIFF 文件路径
            
        Returns:
            List[Dict]: 翻译单元列表，包含占位符信息
        """
        try:
            logger.info(f"开始解析 XLIFF 文件（使用占位符）: {xliff_file}")
            
            if not os.path.exists(xliff_file):
                logger.error(f"XLIFF 文件不存在: {xliff_file}")
                return []
            
            tree = ET.parse(xliff_file)
            root = tree.getroot()
            
            # 支持 XLIFF 1.2 和 2.0 格式
            if 'urn:oasis:names:tc:xliff:document:2.0' in root.tag:
                logger.info("解析 XLIFF 2.0 格式")
                namespaces = {'xliff': 'urn:oasis:names:tc:xliff:document:2.0'}
                unit_xpath = './/xliff:unit'
                source_xpath = './/xliff:source'
                target_xpath = './/xliff:target'
            else:
                logger.info("解析 XLIFF 1.2 格式")
                namespaces = {'xliff': 'urn:oasis:names:tc:xliff:document:1.2'}
                unit_xpath = './/xliff:trans-unit'
                source_xpath = './/xliff:source'
                target_xpath = './/xliff:target'
            
            units = root.findall(unit_xpath, namespaces)
            logger.info(f"找到 {len(units)} 个翻译单元")
            
            translation_units = []
            
            for unit in units:
                unit_id = unit.get('id', '')
                source_elem = unit.find(source_xpath, namespaces)
                target_elem = unit.find(target_xpath, namespaces)
                
                if source_elem is not None:
                    # 构建带占位符的文本结构
                    placeholder_info = self._build_placeholder_text(source_elem)
                    
                    if placeholder_info['text_with_placeholders']:
                        translation_units.append({
                            'id': unit_id,
                            'source': placeholder_info['text_with_placeholders'],
                            'placeholder_info': placeholder_info,
                            'target': ''.join(target_elem.itertext()).strip() if target_elem is not None else '',
                            'unit_element': unit
                        })
            
            logger.info(f"解析到 {len(translation_units)} 个翻译单元（使用占位符）")
            return translation_units
            
        except Exception as e:
            logger.error(f"解析 XLIFF 文件失败: {e}")
            return []

    def _build_placeholder_text(self, source_elem) -> Dict[str, Any]:
        """
        构建带占位符的文本结构
        
        Args:
            source_elem: source元素
            
        Returns:
            Dict: 包含占位符信息的字典
        """
        # 使用特殊占位符 ☼ (U+263C WHITE SUN WITH RAYS)
        # 这个符号很少见，不太可能出现在正常文本中
        PLACEHOLDER = "♂"
        
        text_parts = []
        text_with_placeholders = ""
        format_tags = []
        
        # 处理开头的直接文本
        if source_elem.text:
            text_parts.append(source_elem.text)
            text_with_placeholders += source_elem.text
        
        # 处理所有子元素
        for child in source_elem:
            # 处理格式标签
            if child.tag.endswith('}bpt') or child.tag.endswith('}ept') or child.tag.endswith('}ph'):
                format_tags.append({
                    'tag': child.tag,
                    'attrib': child.attrib.copy(),
                    'content': child.text or '',
                    'position': len(text_parts)  # 记录在哪个文本段之后
                })
                # 在格式标签位置插入占位符
                text_with_placeholders += PLACEHOLDER
            
            # 处理标签后的文本
            if child.tail:
                text_parts.append(child.tail)
                text_with_placeholders += child.tail
        
        return {
            'text_parts': text_parts,
            'text_with_placeholders': text_with_placeholders,
            'format_tags': format_tags,
            'placeholder': PLACEHOLDER
        }

    def _clean_placeholder_spaces(self, translated_text: str, placeholder: str) -> str:
        """
        智能清理占位符前后的空格
        
        Args:
            translated_text: 翻译后的文本
            placeholder: 占位符
            
        Returns:
            str: 清理后的文本
        """
        import re
        
        # 情况1：占位符前后都有空格，删除后面的一个空格
        # 例如：♂ XCMG ♂ Personnel -> ♂ XCMG ♂Personnel
        pattern = f"{placeholder} (.+?) {placeholder}"
        replacement = f"{placeholder}\\1 {placeholder}"
        cleaned_text = re.sub(pattern, replacement, translated_text)
        
        # 其他情况（♂前面有空格后面没有，♂后面有空格前面没有，♂前后都没有空格）
        # 都不删除空格，保持原样
        
        logger.debug(f"智能清理占位符空格: '{translated_text}' -> '{cleaned_text}'")
        return cleaned_text

    def apply_translation_with_placeholders(self, target_elem, translated_text: str, placeholder_info: Dict[str, Any]) -> bool:
        """
        将翻译文本应用到目标元素，使用占位符恢复格式标签
        
        Args:
            target_elem: 目标元素
            translated_text: 翻译后的文本
            placeholder_info: 占位符信息
            
        Returns:
            bool: 是否成功
        """
        try:
            # 清空目标元素
            target_elem.clear()
            
            placeholder = placeholder_info['placeholder']
            format_tags = placeholder_info['format_tags']
            
            # 清理占位符后面的多余空格
            cleaned_text = self._clean_placeholder_spaces(translated_text, placeholder)
            
            if not format_tags:
                # 没有格式标签，直接设置文本
                target_elem.text = cleaned_text
                return True
            
            # 按占位符分割翻译文本
            translated_parts = cleaned_text.split(placeholder)
            logger.debug(f"翻译文本按占位符分割: {len(translated_parts)} 部分")
            
            # 重建XML结构
            current_part_index = 0
            
            for i, format_tag in enumerate(format_tags):
                # 添加格式标签前的文本
                if current_part_index < len(translated_parts):
                    text_before = translated_parts[current_part_index]
                    if target_elem.text is None:
                        target_elem.text = text_before
                    else:
                        if len(target_elem) == 0:
                            target_elem.text = (target_elem.text or "") + text_before
                        else:
                            last_child = target_elem[-1]
                            last_child.tail = (last_child.tail or "") + text_before
                    current_part_index += 1
                
                # 添加格式标签
                new_tag = ET.SubElement(target_elem, format_tag['tag'], format_tag['attrib'])
                if format_tag['content']:
                    new_tag.text = format_tag['content']
            
            # 添加剩余的文本
            if current_part_index < len(translated_parts):
                remaining_text = translated_parts[current_part_index]
                if target_elem.text is None:
                    target_elem.text = remaining_text
                else:
                    if len(target_elem) == 0:
                        target_elem.text = (target_elem.text or "") + remaining_text
                    else:
                        last_child = target_elem[-1]
                        last_child.tail = (last_child.tail or "") + remaining_text
            
            logger.debug(f"占位符格式映射成功，翻译文本长度: {len(cleaned_text)}")
            return True
            
        except Exception as e:
            logger.error(f"占位符格式映射失败: {e}")
            return False
    
    def update_xliff_translations(self, xliff_file: str, translations: Dict[str, str]) -> bool:
        """
        更新 XLIFF 文件中的翻译内容
        
        Args:
            xliff_file: XLIFF 文件路径
            translations: 翻译结果字典 {unit_id: translated_text}
            
        Returns:
            bool: 是否成功
        """
        try:
            tree = ET.parse(xliff_file)
            root = tree.getroot()
            
            updated_count = 0
            
            # 检查 XLIFF 版本并定义相应的命名空间
            if 'urn:oasis:names:tc:xliff:document:2.0' in root.tag:
                logger.info("更新 XLIFF 2.0 格式")
                namespaces = {'xliff': 'urn:oasis:names:tc:xliff:document:2.0'}
                unit_xpath = './/xliff:unit'
                source_xpath = './/xliff:source'
                target_xpath = './/xliff:target'
                target_ns = f'{{{namespaces["xliff"]}}}target'
            else:
                logger.info("更新 XLIFF 1.2 格式")
                namespaces = {'xliff': 'urn:oasis:names:tc:xliff:document:1.2'}
                unit_xpath = './/xliff:trans-unit'
                source_xpath = './/xliff:source'
                target_xpath = './/xliff:target'
                target_ns = f'{{{namespaces["xliff"]}}}target'
            
            # 更新每个翻译单元
            units = root.findall(unit_xpath, namespaces)
            logger.info(f"找到 {len(units)} 个翻译单元进行更新")
            
            for unit in units:
                unit_id = unit.get('id', '')
                
                if unit_id in translations:
                    # 查找或创建目标元素
                    target_elem = unit.find(target_xpath, namespaces)
                    if target_elem is None:
                        # 如果没有目标元素，创建一个
                        source_elem = unit.find(source_xpath, namespaces)
                        if source_elem is not None:
                            target_elem = ET.SubElement(unit, target_ns)
                    
                    if target_elem is not None:
                        # 保持target元素的完整结构，只替换其中的文本内容
                        translated_text = translations[unit_id]
                        
                        # 方法：保持所有<bpt>和<ept>标签，只替换它们之间的文本
                        # 首先保存所有格式标签
                        format_tags = []
                        for child in target_elem:
                            if child.tag.endswith('}bpt') or child.tag.endswith('}ept'):
                                format_tags.append((child.tag, child.attrib.copy(), child.text))
                        
                        # 清空target元素
                        target_elem.clear()
                        
                        # 重新插入格式标签
                        for tag, attrib, text in format_tags:
                            new_tag = ET.SubElement(target_elem, tag, attrib)
                            if text:
                                new_tag.text = text
                        
                        # 将翻译文本作为纯文本插入（Okapi会在合并时处理）
                        target_elem.text = translated_text
                        
                        updated_count += 1
                        logger.debug(f"更新翻译单元 {unit_id}: {translated_text[:50]}...")
            
            # 保存更新后的文件
            tree.write(xliff_file, encoding='utf-8', xml_declaration=True)
            
            logger.info(f"更新了 {updated_count} 个翻译单元")
            return True
            
        except Exception as e:
            logger.error(f"更新 XLIFF 翻译失败: {e}")
            return False

    def update_xliff_translations_with_placeholders(self, xliff_file: str, translations: Dict[str, Dict]) -> bool:
        """
        更新 XLIFF 文件中的翻译内容，使用占位符恢复格式标签
        
        Args:
            xliff_file: XLIFF 文件路径
            translations: 翻译结果字典 {unit_id: {'text': translated_text, 'placeholder_info': placeholder_info}}
            
        Returns:
            bool: 是否成功
        """
        try:
            tree = ET.parse(xliff_file)
            root = tree.getroot()
            
            updated_count = 0
            
            # 检查 XLIFF 版本并定义相应的命名空间
            if 'urn:oasis:names:tc:xliff:document:2.0' in root.tag:
                logger.info("更新 XLIFF 2.0 格式（使用占位符）")
                namespaces = {'xliff': 'urn:oasis:names:tc:xliff:document:2.0'}
                unit_xpath = './/xliff:unit'
                source_xpath = './/xliff:source'
                target_xpath = './/xliff:target'
                target_ns = f'{{{namespaces["xliff"]}}}target'
            else:
                logger.info("更新 XLIFF 1.2 格式（使用占位符）")
                namespaces = {'xliff': 'urn:oasis:names:tc:xliff:document:1.2'}
                unit_xpath = './/xliff:trans-unit'
                source_xpath = './/xliff:source'
                target_xpath = './/xliff:target'
                target_ns = f'{{{namespaces["xliff"]}}}target'
            
            # 更新每个翻译单元
            units = root.findall(unit_xpath, namespaces)
            logger.info(f"找到 {len(units)} 个翻译单元进行更新")
            logger.info(f"翻译结果字典包含 {len(translations)} 个条目")
            
            for unit in units:
                unit_id = unit.get('id', '')
                
                if unit_id in translations:
                    # 查找或创建目标元素
                    target_elem = unit.find(target_xpath, namespaces)
                    if target_elem is None:
                        # 如果没有目标元素，创建一个
                        source_elem = unit.find(source_xpath, namespaces)
                        if source_elem is not None:
                            target_elem = ET.SubElement(unit, target_ns)
                    
                    if target_elem is not None:
                        translation_data = translations[unit_id]
                        
                        # 检查数据结构
                        if isinstance(translation_data, dict) and 'text' in translation_data:
                            # 占位符模式
                            translated_text = translation_data['text']
                            placeholder_info = translation_data['placeholder_info']
                            
                            # 使用占位符方法
                            success = self.apply_translation_with_placeholders(
                                target_elem, translated_text, placeholder_info
                            )
                        else:
                            # 简单模式
                            translated_text = translation_data
                            target_elem.text = translated_text
                            success = True
                        
                        if success:
                            updated_count += 1
                            logger.debug(f"更新翻译单元 {unit_id}: {translated_text[:50]}...")
                        else:
                            logger.warning(f"更新翻译单元 {unit_id} 失败")
            
            # 保存更新后的文件
            tree.write(xliff_file, encoding='utf-8', xml_declaration=True)
            
            logger.info(f"更新了 {updated_count} 个翻译单元（使用占位符）")
            return True
            
        except Exception as e:
            logger.error(f"更新 XLIFF 翻译失败: {e}")
            return False


class OkapiWordTranslator:
    """使用 Okapi Framework 的 Word 文档翻译器"""
    
    def __init__(self, okapi_home: str = "/opt/okapi", use_placeholders: bool = True):
        """
        初始化翻译器
        
        Args:
            okapi_home: Okapi 安装目录
            use_placeholders: 是否使用占位符模式
        """
        self.okapi_integration = DockerOkapiIntegration(okapi_home)
        self.translation_service = None  # 翻译服务将在后续设置
        self.use_placeholders = use_placeholders  # 占位符模式
        
        logger.info(f"Okapi Word 翻译器初始化完成，使用占位符: {use_placeholders}")
    
    def set_translation_service(self, translation_service):
        """设置翻译服务"""
        self.translation_service = translation_service
        logger.info("翻译服务已设置")
    
    def translate_document(self, input_file: str, output_file: str,
                          source_lang: str = "zh", target_lang: str = "en") -> bool:
        """
        翻译 Word 文档
        
        Args:
            input_file: 输入文件路径
            output_file: 输出文件路径
            source_lang: 源语言
            target_lang: 目标语言
            
        Returns:
            bool: 是否成功
        """
        if not self.translation_service:
            raise OkapiIntegrationError("翻译服务未设置")
        
        try:
            with self.okapi_integration.temp_workspace() as workspace:
                # 文件路径
                xliff_file = os.path.join(workspace, "extracted.xliff")
                translated_xliff = os.path.join(workspace, "translated.xliff")
                
                # 步骤1：Word 文档 → XLIFF
                logger.info("🔄 步骤1: 提取 Word 文档到 XLIFF...")
                success = self.okapi_integration.extract_to_xliff(
                    input_file, xliff_file, source_lang, target_lang
                )
                if not success:
                    return False
                
                # 步骤2：翻译 XLIFF 内容
                logger.info("🔄 步骤2: 翻译 XLIFF 内容...")
                if self.use_placeholders:
                    logger.info("使用占位符模式")
                    success = self._translate_xliff_content_with_placeholders(
                        xliff_file, translated_xliff, source_lang, target_lang
                    )
                else:
                    logger.info("使用简单模式")
                    success = self._translate_xliff_content(
                        xliff_file, translated_xliff, source_lang, target_lang
                    )
                if not success:
                    return False
                
                # 步骤2.5：根据目标语言调整XLIFF中的字体信息
                logger.info("🔄 步骤2.5: 根据目标语言调整XLIFF字体...")
                font_success = self._adjust_xliff_font(translated_xliff, target_lang)
                if font_success:
                    logger.info("✅ XLIFF字体调整成功")
                else:
                    logger.warning("XLIFF字体调整失败，但继续处理")
                
                # 步骤3：XLIFF → Word 文档
                logger.info("🔄 步骤3: 合并翻译后的 XLIFF 到 Word...")
                success = self.okapi_integration.merge_from_xliff(
                    input_file, translated_xliff, output_file
                )
                
                if success:
                    # 步骤4：根据目标语言调整Word文档字体
                    logger.info("🔄 步骤4: 根据目标语言调整Word文档字体...")
                    font_success = self._adjust_word_document_font(output_file, target_lang)
                    if font_success:
                        logger.info("✅ Word文档字体调整成功")
                    else:
                        logger.warning("Word文档字体调整失败，但翻译已完成")
                    
                    # 只有在字体调整完成后才返回成功
                    # 这样可以确保前端状态更新时，所有处理都已完成
                    logger.info("🎯 所有处理完成：翻译 + 字体调整")
                    return True
                else:
                    logger.error("❌ Word文档合并失败，无法进行字体调整")
                    return False
                
        except Exception as e:
            logger.error(f"翻译过程出错: {e}")
            return False
    
    def _translate_xliff_content(self, xliff_file: str, translated_xliff: str,
                                source_lang: str, target_lang: str) -> bool:
        """
        翻译 XLIFF 文件中的内容
        
        Args:
            xliff_file: 原始 XLIFF 文件
            translated_xliff: 翻译后的 XLIFF 文件
            source_lang: 源语言
            target_lang: 目标语言
            
        Returns:
            bool: 是否成功
        """
        try:
            # 解析 XLIFF 内容
            translation_units = self.okapi_integration.parse_xliff_content(xliff_file)
            
            if not translation_units:
                logger.warning("没有找到需要翻译的内容")
                # 复制原文件
                shutil.copy2(xliff_file, translated_xliff)
                return True
            
            # 批量翻译
            translations = {}
            batch_texts = []
            batch_ids = []
            code_mappings = []
            
            # 直接使用纯文本进行翻译
            for unit in translation_units:
                batch_texts.append(unit['source'])
                batch_ids.append(unit['id'])
                code_mappings.append({})
            
            logger.info(f"开始批量翻译 {len(batch_texts)} 个文本单元...")
            
            # 调用翻译服务
            translated_texts = self.translation_service.batch_translate(
                batch_texts, source_lang, target_lang
            )
            
            if len(translated_texts) != len(batch_texts):
                logger.error("翻译结果数量不匹配")
                return False
            
            # 直接使用翻译结果，不恢复格式标签（让Okapi处理格式）
            for i, unit_id in enumerate(batch_ids):
                translations[unit_id] = translated_texts[i]
            
            # 复制原文件并更新翻译
            shutil.copy2(xliff_file, translated_xliff)
            success = self.okapi_integration.update_xliff_translations(
                translated_xliff, translations
            )
            
            if success:
                logger.info(f"✅ XLIFF 内容翻译完成，共翻译 {len(translations)} 个单元")
            
            return success
            
        except Exception as e:
            logger.error(f"翻译 XLIFF 内容失败: {e}")
            return False

    def _translate_xliff_content_with_placeholders(self, xliff_file: str, translated_xliff: str,
                                                  source_lang: str, target_lang: str) -> bool:
        """
        翻译 XLIFF 文件中的内容，使用占位符保持格式标签
        
        Args:
            xliff_file: 原始 XLIFF 文件
            translated_xliff: 翻译后的 XLIFF 文件
            source_lang: 源语言
            target_lang: 目标语言
            
        Returns:
            bool: 是否成功
        """
        try:
            # 使用占位符解析方法
            translation_units = self.okapi_integration.parse_xliff_with_placeholders(xliff_file)
            
            if not translation_units:
                logger.warning("没有找到需要翻译的内容")
                shutil.copy2(xliff_file, translated_xliff)
                return True
            
            # 批量翻译
            translations = {}
            placeholder_infos = {}
            batch_texts = []
            batch_ids = []
            
            # 提取带占位符的文本进行翻译
            for unit in translation_units:
                batch_texts.append(unit['source'])
                batch_ids.append(unit['id'])
                placeholder_infos[unit['id']] = unit['placeholder_info']
            
            logger.info(f"开始批量翻译 {len(batch_texts)} 个文本单元（带占位符）...")
            
            # 调用翻译服务
            translated_texts = self.translation_service.batch_translate(
                batch_texts, source_lang, target_lang
            )
            
            if len(translated_texts) != len(batch_texts):
                logger.error("翻译结果数量不匹配")
                return False
            
            # 保存翻译结果和占位符信息
            for i, unit_id in enumerate(batch_ids):
                translations[unit_id] = {
                    'text': translated_texts[i],
                    'placeholder_info': placeholder_infos[unit_id]
                }
            
            # 复制原文件并更新翻译
            shutil.copy2(xliff_file, translated_xliff)
            success = self.okapi_integration.update_xliff_translations_with_placeholders(
                translated_xliff, translations
            )
            
            if success:
                logger.info(f"✅ XLIFF 内容翻译完成，共翻译 {len(translations)} 个单元（使用占位符）")
            
            return success
            
        except Exception as e:
            logger.error(f"翻译 XLIFF 内容失败: {e}")
            return False

    def _adjust_xliff_font(self, xliff_file: str, target_lang: str) -> bool:
        """
        在XLIFF文件中根据目标语言调整字体信息
        
        Args:
            xliff_file: XLIFF文件路径
            target_lang: 目标语言
            
        Returns:
            bool: 是否成功
        """
        try:
            # 检查目标语言是否为English
            if not target_lang or target_lang.lower() not in ['english', 'en', 'eng']:
                logger.debug(f"目标语言不是English，保持原字体: {target_lang}")
                return True
            
            logger.info(f"目标语言是English，开始调整XLIFF字体为Times New Roman")
            
            # 检查文件是否存在
            if not os.path.exists(xliff_file):
                logger.error(f"XLIFF文件不存在: {xliff_file}")
                return False
            
            # 读取XLIFF文件内容进行分析
            with open(xliff_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            logger.info(f"XLIFF文件大小: {len(content)} 字符")
            logger.info(f"XLIFF文件前1000字符: {content[:1000]}")
            
            # 查找字体相关的信息 - 支持XLIFF 1.2和2.0格式
            font_patterns = [
                # 标准字体属性
                r'font-family="([^"]*)"',
                r'font="([^"]*)"',
                r'family="([^"]*)"',
                r'typeface="([^"]*)"',
                # XLIFF 1.2特定属性
                r'ns\d+:font="([^"]*)"',
                r'ns\d+:family="([^"]*)"',
                # 通用属性查找
                r'[a-zA-Z-]*font[a-zA-Z-]*="([^"]*)"',
                r'[a-zA-Z-]*family[a-zA-Z-]*="([^"]*)"',
            ]
            
            updated_count = 0
            modified_content = content
            
            # 策略1：查找标准字体属性
            for pattern in font_patterns:
                matches = re.findall(pattern, content, re.IGNORECASE)
                for match in matches:
                    if match.lower() not in ['times new roman', 'timesnewroman', 'times', 'new roman']:
                        old_font = match
                        # 替换字体信息
                        modified_content = re.sub(
                            f'{pattern[:-1]}"{re.escape(old_font)}"',
                            f'{pattern[:-1]}"Times New Roman"',
                            modified_content,
                            flags=re.IGNORECASE
                        )
                        logger.info(f"字体属性从 '{old_font}' 更改为 'Times New Roman'")
                        updated_count += 1
            
            # 策略2：查找包含字体信息的文本内容
            font_names_in_text = [
                'arial', 'simsun', 'simhei', 'microsoft', 'calibri', 'verdana', 'tahoma',
                '宋体', '黑体', '微软雅黑', '新宋体', '仿宋', '楷体'
            ]
            
            for font_name in font_names_in_text:
                if font_name in content.lower():
                    # 替换文本中的字体名称
                    modified_content = re.sub(
                        font_name, 'Times New Roman', modified_content, flags=re.IGNORECASE
                    )
                    logger.info(f"文本内容字体从 '{font_name}' 更改为 'Times New Roman'")
                    updated_count += 1
            
            # 策略3：查找Word特定的格式标签
            word_format_patterns = [
                r'<w:rFonts w:ascii="([^"]*)"',
                r'<w:rFonts w:eastAsia="([^"]*)"',
                r'<w:rFonts w:hAnsi="([^"]*)"',
                r'<w:rFonts w:cs="([^"]*)"',
            ]
            
            for pattern in word_format_patterns:
                matches = re.findall(pattern, content, re.IGNORECASE)
                for match in matches:
                    if match.lower() not in ['times new roman', 'timesnewroman', 'times', 'new roman']:
                        old_font = match
                        # 替换Word格式标签中的字体
                        modified_content = re.sub(
                            f'{pattern[:-1]}"{re.escape(old_font)}"',
                            f'{pattern[:-1]}"Times New Roman"',
                            modified_content,
                            flags=re.IGNORECASE
                        )
                        logger.info(f"Word格式标签字体从 '{old_font}' 更改为 'Times New Roman'")
                        updated_count += 1
            
            # 策略4：查找样式定义中的字体
            style_patterns = [
                r'<w:style w:name="[^"]*"[^>]*>.*?<w:rFonts[^>]*w:ascii="([^"]*)"',
                r'<w:style w:name="[^"]*"[^>]*>.*?<w:rFonts[^>]*w:eastAsia="([^"]*)"',
                r'<w:style w:name="[^"]*"[^>]*>.*?<w:rFonts[^>]*w:hAnsi="([^"]*)"',
            ]
            
            for pattern in style_patterns:
                matches = re.findall(pattern, content, re.IGNORECASE | re.DOTALL)
                for match in matches:
                    if match.lower() not in ['times new roman', 'timesnewroman', 'times', 'new roman']:
                        old_font = match
                        # 替换样式中的字体
                        modified_content = re.sub(
                            f'{pattern[:-1]}"{re.escape(old_font)}"',
                            f'{pattern[:-1]}"Times New Roman"',
                            modified_content,
                            flags=re.IGNORECASE | re.DOTALL
                        )
                        logger.info(f"样式定义字体从 '{old_font}' 更改为 'Times New Roman'")
                        updated_count += 1
            
            # 策略5：查找并处理我们添加的字体标记
            font_marker_pattern = r'<font:Times New Roman>(.*?)</font:Times New Roman>'
            font_marker_matches = re.findall(font_marker_pattern, content)
            
            if font_marker_matches:
                logger.info(f"找到 {len(font_marker_matches)} 个字体标记")
                # 移除字体标记，但保留文本内容
                modified_content = re.sub(font_marker_pattern, r'\1', modified_content)
                updated_count += len(font_marker_matches)
                logger.info("已移除字体标记，文本内容保留")
                
                # 在XLIFF中添加字体样式信息
                # 查找所有trans-unit标签，为包含英文的target添加字体属性
                trans_unit_pattern = r'(<ns0:target[^>]*>)(.*?)(</ns0:target>)'
                
                def add_font_to_target(match):
                    target_tag = match.group(1)
                    target_content = match.group(2)
                    closing_tag = match.group(3)
                    
                    # 检查目标文本是否包含英文（可能是翻译后的文本）
                    if re.search(r'[a-zA-Z]', target_content):
                        # 在target标签中添加字体属性
                        if 'xml:lang="english"' in target_tag or 'target-language="english"' in content:
                            # 添加字体样式属性
                            font_attr = ' ns1:font="Times New Roman"'
                            if font_attr not in target_tag:
                                target_tag = target_tag.replace('>', font_attr + '>')
                                logger.info(f"为目标文本添加字体属性: Times New Roman")
                                updated_count += 1
                    
                    return target_tag + target_content + closing_tag
                
                modified_content = re.sub(trans_unit_pattern, add_font_to_target, modified_content, flags=re.DOTALL)
            
            # 如果内容有变化，保存文件
            if modified_content != content:
                with open(xliff_file, 'w', encoding='utf-8') as f:
                    f.write(modified_content)
                logger.info(f"XLIFF字体调整完成，更新了 {updated_count} 处字体信息")
                return True
            else:
                logger.info("未找到需要调整的字体信息")
                # 输出更多调试信息
                logger.info("尝试查找文件中的其他格式信息...")
                
                # 查找所有可能包含格式信息的标签
                format_tags = re.findall(r'<[^>]*>', content)
                format_info = []
                for tag in format_tags:
                    if any(keyword in tag.lower() for keyword in ['font', 'family', 'style', 'format']):
                        format_info.append(tag)
                
                if format_info:
                    logger.info(f"找到 {len(format_info)} 个可能包含格式信息的标签:")
                    for i, tag in enumerate(format_info[:10]):  # 只显示前10个
                        logger.info(f"  {i+1}: {tag}")
                else:
                    logger.info("未找到任何格式相关的标签")
                
                return True
                
        except Exception as e:
            logger.error(f"调整XLIFF字体失败: {str(e)}")
            return False

    def _adjust_word_document_font(self, docx_file: str, target_lang: str) -> bool:
        """
        根据目标语言调整Word文档字体
        
        Args:
            docx_file: Word文档路径
            target_lang: 目标语言
            
        Returns:
            bool: 是否成功
        """
        try:
            # 检查目标语言是否为English
            if not target_lang or target_lang.lower() not in ['english', 'en', 'eng']:
                logger.debug(f"目标语言不是English，保持原字体: {target_lang}")
                return True
            
            logger.info(f"目标语言是English，开始调整字体为Times New Roman")
            
            # 检查文件是否存在
            if not os.path.exists(docx_file):
                logger.error(f"Word文档不存在: {docx_file}")
                return False
            
            # 使用python-docx直接修改Word文档
            try:
                from docx import Document
                
                # 加载文档
                doc = Document(docx_file)
                
                # 遍历所有段落和表格，调整字体
                updated_count = 0
                
                # 调整段落字体
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        # 检查run是否包含英文文本
                        if run.text and re.search(r'[a-zA-Z]', run.text):
                            # 如果字体不是Times New Roman，则更改
                            if not run.font.name or run.font.name != 'Times New Roman':
                                old_font = run.font.name or '默认字体'
                                run.font.name = 'Times New Roman'
                                logger.debug(f"段落字体从 '{old_font}' 更改为 'Times New Roman': {run.text[:30]}...")
                                updated_count += 1
                
                # 调整表格字体
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    # 检查run是否包含英文文本
                                    if run.text and re.search(r'[a-zA-Z]', run.text):
                                        # 如果字体不是Times New Roman，则更改
                                        if not run.font.name or run.font.name != 'Times New Roman':
                                            old_font = run.font.name or '默认字体'
                                            run.font.name = 'Times New Roman'
                                            logger.debug(f"表格字体从 '{old_font}' 更改为 'Times New Roman': {run.text[:30]}...")
                                            updated_count += 1
                
                # 保存文档
                doc.save(docx_file)
                logger.info(f"Word文档字体调整完成，更新了 {updated_count} 个包含英文的run")
                return True
                
            except ImportError:
                logger.error("python-docx未安装，无法调整Word文档字体")
                return False
            except Exception as e:
                logger.error(f"调整Word文档字体失败: {str(e)}")
                return False
                
        except Exception as e:
            logger.error(f"调整Word文档字体失败: {str(e)}")
            return False


class OkapiPptxTranslator:
    """使用 Okapi Framework 的 PPTX 文档翻译器"""
    
    def __init__(self, okapi_home: str = "/opt/okapi", use_placeholders: bool = True):
        """
        初始化翻译器
        
        Args:
            okapi_home: Okapi 安装目录
            use_placeholders: 是否使用占位符模式
        """
        self.okapi_integration = DockerOkapiIntegration(okapi_home)
        self.translation_service = None  # 翻译服务将在后续设置
        self.use_placeholders = use_placeholders  # 占位符模式
        
        logger.info(f"Okapi PPTX 翻译器初始化完成，使用占位符: {use_placeholders}")
    
    def set_translation_service(self, translation_service):
        """设置翻译服务"""
        self.translation_service = translation_service
        logger.info("翻译服务已设置")
    
    def translate_document(self, input_file: str, output_file: str,
                          source_lang: str = "zh", target_lang: str = "en") -> bool:
        """
        翻译 PPTX 文档
        
        Args:
            input_file: 输入文件路径
            output_file: 输出文件路径
            source_lang: 源语言
            target_lang: 目标语言
            
        Returns:
            bool: 是否成功
        """
        if not self.translation_service:
            raise OkapiIntegrationError("翻译服务未设置")
        
        try:
            with self.okapi_integration.temp_workspace() as workspace:
                # 文件路径
                xliff_file = os.path.join(workspace, "extracted.xliff")
                translated_xliff = os.path.join(workspace, "translated.xliff")
                
                # 步骤1：PPTX 文档 → XLIFF
                logger.info("🔄 步骤1: 提取 PPTX 文档到 XLIFF...")
                success = self.okapi_integration.extract_to_xliff(
                    input_file, xliff_file, source_lang, target_lang
                )
                if not success:
                    return False
                
                # 步骤2：翻译 XLIFF 内容
                logger.info("🔄 步骤2: 翻译 XLIFF 内容...")
                if self.use_placeholders:
                    logger.info("使用占位符模式")
                    success = self._translate_xliff_content_with_placeholders(
                        xliff_file, translated_xliff, source_lang, target_lang
                    )
                else:
                    logger.info("使用简单模式")
                    success = self._translate_xliff_content(
                        xliff_file, translated_xliff, source_lang, target_lang
                    )
                if not success:
                    return False
                
                # 步骤3：XLIFF → PPTX 文档
                logger.info("🔄 步骤3: 合并翻译后的 XLIFF 到 PPTX...")
                success = self.okapi_integration.merge_from_xliff(
                    input_file, translated_xliff, output_file
                )
                
                if success:
                    logger.info("🎯 PPTX 翻译完成")
                    return True
                else:
                    logger.error("❌ PPTX 文档合并失败")
                    return False
                
        except Exception as e:
            logger.error(f"翻译过程出错: {e}")
            return False
    
    def _translate_xliff_content(self, xliff_file: str, translated_xliff: str,
                                source_lang: str, target_lang: str) -> bool:
        """
        翻译 XLIFF 文件中的内容
        
        Args:
            xliff_file: 原始 XLIFF 文件
            translated_xliff: 翻译后的 XLIFF 文件
            source_lang: 源语言
            target_lang: 目标语言
            
        Returns:
            bool: 是否成功
        """
        try:
            # 解析 XLIFF 内容
            translation_units = self.okapi_integration.parse_xliff_content(xliff_file)
            
            if not translation_units:
                logger.warning("没有找到需要翻译的内容")
                # 复制原文件
                shutil.copy2(xliff_file, translated_xliff)
                return True
            
            # 批量翻译
            translations = {}
            batch_texts = []
            batch_ids = []
            
            # 直接使用纯文本进行翻译
            for unit in translation_units:
                batch_texts.append(unit['source'])
                batch_ids.append(unit['id'])
            
            logger.info(f"开始批量翻译 {len(batch_texts)} 个文本单元...")
            
            # 调用翻译服务
            translated_texts = self.translation_service.batch_translate(
                batch_texts, source_lang, target_lang
            )
            
            if len(translated_texts) != len(batch_texts):
                logger.error("翻译结果数量不匹配")
                return False
            
            # 直接使用翻译结果，不恢复格式标签（让Okapi处理格式）
            for i, unit_id in enumerate(batch_ids):
                translations[unit_id] = translated_texts[i]
            
            # 复制原文件并更新翻译
            shutil.copy2(xliff_file, translated_xliff)
            success = self.okapi_integration.update_xliff_translations(
                translated_xliff, translations
            )
            
            if success:
                logger.info(f"✅ XLIFF 内容翻译完成，共翻译 {len(translations)} 个单元")
            
            return success
            
        except Exception as e:
            logger.error(f"翻译 XLIFF 内容失败: {e}")
            return False

    def _translate_xliff_content_with_placeholders(self, xliff_file: str, translated_xliff: str,
                                                  source_lang: str, target_lang: str) -> bool:
        """
        翻译 XLIFF 文件中的内容，使用占位符保持格式标签
        
        Args:
            xliff_file: 原始 XLIFF 文件
            translated_xliff: 翻译后的 XLIFF 文件
            source_lang: 源语言
            target_lang: 目标语言
            
        Returns:
            bool: 是否成功
        """
        try:
            # 使用占位符解析方法
            translation_units = self.okapi_integration.parse_xliff_with_placeholders(xliff_file)
            
            if not translation_units:
                logger.warning("没有找到需要翻译的内容")
                shutil.copy2(xliff_file, translated_xliff)
                return True
            
            # 批量翻译
            translations = {}
            placeholder_infos = {}
            batch_texts = []
            batch_ids = []
            
            # 提取带占位符的文本进行翻译
            for unit in translation_units:
                batch_texts.append(unit['source'])
                batch_ids.append(unit['id'])
                placeholder_infos[unit['id']] = unit['placeholder_info']
            
            logger.info(f"开始批量翻译 {len(batch_texts)} 个文本单元（带占位符）...")
            
            # 调用翻译服务
            translated_texts = self.translation_service.batch_translate(
                batch_texts, source_lang, target_lang
            )
            
            if len(translated_texts) != len(batch_texts):
                logger.error("翻译结果数量不匹配")
                return False
            
            # 保存翻译结果和占位符信息
            for i, unit_id in enumerate(batch_ids):
                translations[unit_id] = {
                    'text': translated_texts[i],
                    'placeholder_info': placeholder_infos[unit_id]
                }
            
            # 复制原文件并更新翻译
            shutil.copy2(xliff_file, translated_xliff)
            success = self.okapi_integration.update_xliff_translations_with_placeholders(
                translated_xliff, translations
            )
            
            if success:
                logger.info(f"✅ XLIFF 内容翻译完成，共翻译 {len(translations)} 个单元（使用占位符）")
            
            return success
            
        except Exception as e:
            logger.error(f"翻译 XLIFF 内容失败: {e}")
            return False


# 便捷函数
def create_okapi_translator(okapi_home: str = "/opt/okapi", use_placeholders: bool = True) -> OkapiWordTranslator:
    """创建 Okapi 翻译器实例"""
    return OkapiWordTranslator(okapi_home, use_placeholders)


def verify_okapi_installation(okapi_home: str = "/opt/okapi") -> bool:
    """验证 Okapi 安装是否正确"""
    try:
        # 检查目录是否存在
        if not os.path.exists(okapi_home):
            logger.error(f"❌ Okapi 目录不存在: {okapi_home}")
            return False
        
        # 尝试创建集成实例（会自动查找 JAR 文件）
        integration = DockerOkapiIntegration(okapi_home)
        logger.info("✅ Okapi 安装验证通过")
        return True
    except Exception as e:
        logger.error(f"❌ Okapi 安装验证失败: {e}")
        return False 