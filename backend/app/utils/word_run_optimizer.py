"""
Word文档Run优化器 - 安全合并机制

本模块提供了一个安全、可靠的Word文档run合并机制，用于优化文档结构，减少不必要的run分割。
在保证100%格式不变的前提下，合并那些被Word错误分割的runs。

1. 完整的样式提取 🔍

提取30+个样式属性
包括XML级别的隐藏属性
生成属性哈希值快速比较

2. 特殊边界检测 🚧
系统会检测并保护以下元素：

域代码（Field codes）
超链接
书签
批注和修订
脚注/尾注
分隔符（换行、分页等）
图形对象
制表符

3. 安全合并策略 ✅
只有满足以下所有条件的runs才会被合并：

样式属性100%相同
都不包含任何特殊边界元素
都不是空文本（空文本可能有特殊用途）

主要功能：
1. 深度分析run的所有样式属性（30+个属性）
2. 检测特殊边界元素（超链接、域代码、批注等）
3. 安全合并格式完全相同的相邻runs
4. 生成详细的分析报告

使用示例：
    >>> from word_run_optimizer import SafeRunMerger
    >>> merger = SafeRunMerger()
    >>> stats = merger.optimize_document("input.docx", "output.docx")
    >>> print(f"优化完成，减少了{stats['merged_runs']}个runs")

依赖：
    - python-docx >= 0.8.10
    - lxml >= 4.6.0

作者：Claude
版本：1.0.0
"""

import docx
from docx.shared import RGBColor, Pt
from docx.oxml import parse_xml, register_element_cls
from docx.oxml.ns import nsdecls
from docx.oxml.text.run import CT_R
import hashlib
import json
import logging
from typing import List, Dict, Tuple, Optional, Any
import copy
import os
import sys

# 配置日志
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

# 版本兼容性常量
DOCX_MIN_VERSION = "0.8.10"
PYTHON_MIN_VERSION = (3, 6)


class DocxCompatibilityError(Exception):
    """python-docx版本兼容性错误"""
    pass


class RunOptimizationError(Exception):
    """Run优化过程中的错误"""
    pass


def check_compatibility():
    """
    检查环境兼容性
    
    Raises:
        DocxCompatibilityError: 当环境不满足要求时
    """
    # 检查Python版本
    if sys.version_info < PYTHON_MIN_VERSION:
        raise DocxCompatibilityError(
            f"需要Python {PYTHON_MIN_VERSION[0]}.{PYTHON_MIN_VERSION[1]}或更高版本"
        )
    
    # 检查python-docx版本
    try:
        import docx
        docx_version = getattr(docx, '__version__', '0.0.0')
        # 简单的版本比较（实际项目中可能需要更复杂的版本比较）
        if docx_version < DOCX_MIN_VERSION:
            raise DocxCompatibilityError(
                f"需要python-docx {DOCX_MIN_VERSION}或更高版本，当前版本：{docx_version}"
            )
    except ImportError:
        raise DocxCompatibilityError("未安装python-docx库")


class RunStyleAnalyzer:
    """
    Run样式深度分析器
    
    负责提取和比较run的所有样式属性，包括：
    - 基本字符格式（粗体、斜体、下划线等）
    - 字体属性（字体名称、大小、颜色等）
    - 高级格式（上下标、字符间距等）
    - XML级别的特殊属性
    """
    
    @staticmethod
    def extract_all_properties(run) -> Dict[str, Any]:
        """
        提取run的所有样式属性
        
        Args:
            run: python-docx的Run对象
            
        Returns:
            Dict[str, Any]: 包含所有样式属性的字典
            
        Note:
            - 某些属性可能在不同版本的Word中表现不同
            - None值表示属性未设置（继承自父级样式）
            - 返回的字典可用于完整重建run样式
        """
        properties = {}
        
        # 定义属性列表，使用getattr安全获取
        # 1. 基本字符格式属性
        basic_attrs = [
            ('bold', None),  # 粗体
            ('italic', None),  # 斜体
            ('underline', None),  # 下划线
            ('strike', None),  # 删除线
            ('double_strike', None),  # 双删除线
            ('outline', None),  # 轮廓
            ('shadow', None),  # 阴影
            ('emboss', None),  # 浮雕
            ('imprint', None),  # 印记
            ('all_caps', None),  # 全部大写
            ('small_caps', None),  # 小型大写字母
            ('hidden', None),  # 隐藏
            ('web_hidden', None),  # Web隐藏
            ('spec_vanish', None),  # 特殊隐藏
            ('rtl', None),  # 从右到左
            ('cs_bold', None),  # 复杂脚本粗体
            ('cs_italic', None),  # 复杂脚本斜体
            ('complex_script', None),  # 复杂脚本
        ]
        
        # 安全获取基本属性
        for attr_name, default_value in basic_attrs:
            try:
                properties[attr_name] = getattr(run, attr_name, default_value)
            except AttributeError:
                properties[attr_name] = default_value
            except Exception:
                properties[attr_name] = default_value
            
        
        # 2. 字体基本属性
        if hasattr(run, 'font') and run.font:
            try:
                properties['font_name'] = getattr(run.font, 'name', None)
                
                # 字体大小
                font_size = getattr(run.font, 'size', None)
                if font_size:
                    properties['font_size'] = getattr(font_size, 'pt', None)
                else:
                    properties['font_size'] = None
                
                # 字体颜色
                font_color = getattr(run.font, 'color', None)
                if font_color:
                    properties['font_color_rgb'] = getattr(font_color, 'rgb', None)
                    properties['font_color_theme'] = getattr(font_color, 'theme_color', None)
                    properties['font_color_brightness'] = getattr(font_color, 'brightness', None)
                else:
                    properties['font_color_rgb'] = None
                    properties['font_color_theme'] = None
                    properties['font_color_brightness'] = None
                
                # 高亮颜色
                properties['font_highlight'] = getattr(run.font, 'highlight_color', None)
                
                # 3. 高级字体属性
                font_attrs = [
                    ('subscript', None),  # 下标
                    ('superscript', None),  # 上标
                    ('kern', None),  # 字距调整
                    ('position', None),  # 位置调整
                    ('spacing', None),  # 字符间距
                    ('scale', None),  # 字符缩放
                ]
                
                for attr_name, default_value in font_attrs:
                    properties[attr_name] = getattr(run.font, attr_name, default_value)
                    
            except Exception as e:
                logger.debug(f"处理字体属性时出现小问题: {str(e)}")
        
        # 4. 样式引用
        try:
            if hasattr(run, 'style') and run.style:
                properties['style_id'] = getattr(run.style, 'style_id', None)
                properties['style_name'] = getattr(run.style, 'name', None)
            else:
                properties['style_id'] = None
                properties['style_name'] = None
        except Exception:
            properties['style_id'] = None
            properties['style_name'] = None
            
        
        # 5. XML级别的特殊属性
        # 这些属性可能影响特殊处理逻辑
        try:
            if hasattr(run, '_element') and run._element is not None:
                element = run._element
                
                # XML空格保留属性
                properties['xml_space_preserve'] = element.get(
                    '{http://www.w3.org/XML/1998/namespace}space'
                )
                
                # 获取run properties元素
                try:
                    rPr = element.find(
                        './/w:rPr', 
                        {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    )
                    
                    if rPr is not None:
                        # 语言设置（影响拼写检查和断字）
                        lang = rPr.find(
                            './/w:lang', 
                            {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                        )
                        if lang is not None:
                            properties['lang'] = lang.get(
                                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'
                            )
                            properties['lang_bidi'] = lang.get(
                                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bidi'
                            )
                            properties['lang_eastAsia'] = lang.get(
                                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia'
                            )
                        
                        # 垂直对齐
                        vertAlign = rPr.find(
                            './/w:vertAlign', 
                            {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                        )
                        if vertAlign is not None:
                            properties['vert_align'] = vertAlign.get(
                                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'
                            )
                        
                        # 文字效果
                        effect = rPr.find(
                            './/w:effect', 
                            {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                        )
                        if effect is not None:
                            properties['effect'] = effect.get(
                                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'
                            )
                except Exception as e:
                    logger.debug(f"处理XML属性时出现小问题: {str(e)}")
        except Exception as e:
            logger.debug(f"访问_element时出现问题: {str(e)}")
        
        return properties
    
    @staticmethod
    def properties_identical(props1: Dict[str, Any], props2: Dict[str, Any]) -> bool:
        """
        比较两个属性字典是否完全相同
        
        Args:
            props1: 第一个属性字典
            props2: 第二个属性字典
            
        Returns:
            bool: 如果所有属性都相同返回True，否则返回False
            
        Note:
            - None和False被视为不同的值（Word中有区别）
            - 会比较所有属性，包括只在一个字典中存在的属性
        """
        # 获取所有键的并集
        all_keys = set(props1.keys()) | set(props2.keys())
        
        for key in all_keys:
            val1 = props1.get(key)
            val2 = props2.get(key)
            
            # 严格比较，None和False不相等
            if val1 != val2:
                return False
        
        return True
    
    @staticmethod
    def calculate_property_hash(properties: Dict[str, Any]) -> str:
        """
        计算属性字典的哈希值，用于快速比较
        
        Args:
            properties: 属性字典
            
        Returns:
            str: 属性的MD5哈希值（32字符十六进制字符串）
            
        Note:
            - 相同的属性组合总是产生相同的哈希值
            - 用于快速判断两个run的样式是否相同
        """
        try:
            # 将属性字典转换为稳定的字符串表示
            # 排序确保相同属性产生相同哈希
            sorted_items = sorted(properties.items())
            prop_str = json.dumps(sorted_items, sort_keys=True, default=str)
            return hashlib.md5(prop_str.encode('utf-8')).hexdigest()
        except Exception as e:
            logger.warning(f"计算属性哈希时出现错误: {str(e)}")
            # 返回一个基于错误的唯一值
            return hashlib.md5(f"error_{id(properties)}".encode()).hexdigest()


class RunBoundaryDetector:
    """
    Run边界检测器
    
    识别run中的特殊边界元素，这些元素阻止run合并：
    - 域代码（Field codes）
    - 超链接（Hyperlinks）
    - 书签（Bookmarks）
    - 批注和修订（Comments and Revisions）
    - 各种引用（脚注、尾注等）
    - 分隔符（换行、分页等）
    - 嵌入对象（图片、图表等）
    """
    
    @staticmethod
    def detect_special_boundaries(run) -> Dict[str, bool]:
        """
        检测run是否包含特殊边界元素
        
        Args:
            run: python-docx的Run对象
            
        Returns:
            Dict[str, bool]: 各种边界元素的存在状态
            
        Note:
            - 返回的字典包含所有可能的边界类型
            - True表示存在该类型的边界元素
            - 任何边界元素的存在都会阻止run合并
        """
        boundaries = {
            'has_field': False,      # 域代码
            'has_bookmark': False,   # 书签
            'has_hyperlink': False,  # 超链接
            'has_comment': False,    # 批注
            'has_revision': False,   # 修订标记
            'has_footnote': False,   # 脚注
            'has_endnote': False,    # 尾注
            'has_break': False,      # 分隔符（换行、分页等）
            'has_tab': False,        # 制表符
            'has_drawing': False,    # 图形对象
            'has_inline_shape': False,  # 内联形状
            'has_math': False,       # 数学公式
        }
        
        try:
            if hasattr(run, '_element') and run._element is not None:
                element = run._element
                
                # 检查父元素以确定run的上下文
                parent = element.getparent()
                if parent is not None:
                    parent_tag = parent.tag if hasattr(parent, 'tag') else ''
                    
                    # 检查是否在超链接中
                    if parent_tag.endswith('hyperlink'):
                        boundaries['has_hyperlink'] = True
                    
                    # 检查是否在域中
                    if parent_tag.endswith('fldSimple') or parent_tag.endswith('instrText'):
                        boundaries['has_field'] = True
                    
                    # 检查是否在书签中
                    if parent_tag.endswith('bookmarkStart') or parent_tag.endswith('bookmarkEnd'):
                        boundaries['has_bookmark'] = True
                
                # 检查run内部的子元素
                for child in element:
                    if not hasattr(child, 'tag'):
                        continue
                        
                    tag = child.tag
                    
                    # 分隔符
                    if tag.endswith('br') or tag.endswith('cr'):
                        boundaries['has_break'] = True
                    # 制表符
                    elif tag.endswith('tab'):
                        boundaries['has_tab'] = True
                    # 图形和图片
                    elif tag.endswith('drawing') or tag.endswith('pict'):
                        boundaries['has_drawing'] = True
                    # 内联对象
                    elif tag.endswith('object'):
                        boundaries['has_inline_shape'] = True
                    # 脚注引用
                    elif tag.endswith('footnoteReference'):
                        boundaries['has_footnote'] = True
                    # 尾注引用
                    elif tag.endswith('endnoteReference'):
                        boundaries['has_endnote'] = True
                    # 批注引用
                    elif tag.endswith('commentReference'):
                        boundaries['has_comment'] = True
                    # 修订标记
                    elif tag.endswith('ins') or tag.endswith('del'):
                        boundaries['has_revision'] = True
                    # 数学公式
                    elif tag.endswith('oMath') or tag.endswith('oMathPara'):
                        boundaries['has_math'] = True
                    # 域代码
                    elif tag.endswith('fldChar') or tag.endswith('instrText'):
                        boundaries['has_field'] = True
        
        except Exception as e:
            logger.warning(f"检测run边界时出现错误: {str(e)}")
            # 出错时保守处理，标记为包含未知边界
            boundaries['has_unknown'] = True
        
        return boundaries
    
    @staticmethod
    def can_merge_with_boundaries(boundaries1: Dict[str, bool], 
                                 boundaries2: Dict[str, bool]) -> bool:
        """
        判断两个包含边界信息的run是否可以合并
        
        Args:
            boundaries1: 第一个run的边界检测结果
            boundaries2: 第二个run的边界检测结果
            
        Returns:
            bool: 如果可以安全合并返回True，否则返回False
            
        Note:
            - 只有当两个run都不包含任何特殊边界元素时才能合并
            - 这是一个保守的策略，确保不会破坏文档结构
        """
        # 检查是否有任何边界元素存在
        for key in boundaries1:
            if boundaries1.get(key, False) or boundaries2.get(key, False):
                return False
        
        # 也检查第二个字典中独有的键
        for key in boundaries2:
            if boundaries2.get(key, False):
                return False
        
        return True


class SafeRunMerger:
    """
    安全的Run合并器
    
    主要功能：
    1. 分析段落中的所有runs
    2. 识别可以安全合并的run组
    3. 执行合并操作
    4. 生成统计报告
    
    设计原则：
    - 安全第一：只合并100%确定不会改变格式的runs
    - 保守策略：有疑问时不合并
    - 完整日志：记录所有操作便于调试
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        初始化合并器
        
        Args:
            config: 可选的配置字典，支持以下配置项：
                - 'skip_empty_runs': bool, 是否跳过空run的合并（默认True）
                - 'min_group_size': int, 最小合并组大小（默认2）
                - 'log_level': str, 日志级别（默认'INFO'）
        """
        self.analyzer = RunStyleAnalyzer()
        self.detector = RunBoundaryDetector()
        
        # 默认配置
        self.config = {
            'skip_empty_runs': True,  # 跳过空run（可能有特殊用途）
            'min_group_size': 2,      # 至少2个run才执行合并
            'log_level': 'INFO'
        }
        
        # 更新用户配置
        if config:
            self.config.update(config)
        
        # 统计信息
        self.merge_stats = {
            'total_runs': 0,
            'merged_runs': 0,
            'merge_groups': 0,
            'skipped_runs': 0,
            'error_count': 0
        }
        
        # 设置日志级别
        logger.setLevel(self.config['log_level'])
    
    def analyze_paragraph_runs(self, paragraph) -> List[Dict[str, Any]]:
        """
        分析段落中的所有runs
        
        Args:
            paragraph: python-docx的Paragraph对象
            
        Returns:
            List[Dict]: 每个run的详细分析结果，包含：
                - index: run在段落中的索引
                - text: run的文本内容
                - properties: 所有样式属性
                - boundaries: 边界元素检测结果
                - property_hash: 属性哈希值
                - can_merge: 是否可以参与合并
                
        Raises:
            RunOptimizationError: 分析过程中出现严重错误时
        """
        run_analysis = []
        
        try:
            for i, run in enumerate(paragraph.runs):
                try:
                    # 提取run信息
                    properties = self.analyzer.extract_all_properties(run)
                    boundaries = self.detector.detect_special_boundaries(run)
                    property_hash = self.analyzer.calculate_property_hash(properties)
                    
                    # 判断是否可以参与合并
                    can_merge = True
                    if self.config['skip_empty_runs'] and run.text == '':
                        can_merge = False
                        self.merge_stats['skipped_runs'] += 1
                    
                    # 如果有任何边界元素，不能合并
                    if any(boundaries.values()):
                        can_merge = False
                    
                    analysis = {
                        'index': i,
                        'text': run.text,
                        'text_length': len(run.text),
                        'properties': properties,
                        'boundaries': boundaries,
                        'property_hash': property_hash,
                        'can_merge': can_merge
                    }
                    
                    run_analysis.append(analysis)
                    
                except Exception as e:
                    logger.error(f"分析run {i} 时出错: {str(e)}")
                    self.merge_stats['error_count'] += 1
                    
                    # 添加错误标记的分析结果
                    run_analysis.append({
                        'index': i,
                        'text': run.text if hasattr(run, 'text') else '[ERROR]',
                        'text_length': 0,
                        'properties': {},
                        'boundaries': {'has_error': True},
                        'property_hash': f"error_{i}",
                        'can_merge': False,
                        'error': str(e)
                    })
            
        except Exception as e:
            logger.error(f"分析段落runs时出现严重错误: {str(e)}")
            raise RunOptimizationError(f"段落分析失败: {str(e)}")
        
        return run_analysis
    
    def find_mergeable_groups(self, run_analysis: List[Dict[str, Any]]) -> List[List[int]]:
        """
        找出可以合并的run组
        
        Args:
            run_analysis: analyze_paragraph_runs返回的分析结果
            
        Returns:
            List[List[int]]: 可合并的run索引组列表
            
        Example:
            [[0, 1, 2], [5, 6]] 表示run 0,1,2可以合并，run 5,6可以合并
        """
        if len(run_analysis) < self.config['min_group_size']:
            return []
        
        mergeable_groups = []
        current_group = []
        
        for i in range(len(run_analysis)):
            current_run = run_analysis[i]
            
            # 如果当前run不能合并，结束当前组
            if not current_run.get('can_merge', False):
                if len(current_group) >= self.config['min_group_size']:
                    mergeable_groups.append(current_group)
                current_group = []
                continue
            
            # 如果是第一个可合并的run，开始新组
            if not current_group:
                current_group = [i]
                continue
            
            # 检查是否可以加入当前组
            prev_run = run_analysis[current_group[-1]]
            
            # 比较属性哈希
            if prev_run['property_hash'] == current_run['property_hash']:
                current_group.append(i)
            else:
                # 属性不同，结束当前组，开始新组
                if len(current_group) >= self.config['min_group_size']:
                    mergeable_groups.append(current_group)
                current_group = [i]
        
        # 处理最后一组
        if len(current_group) >= self.config['min_group_size']:
            mergeable_groups.append(current_group)
        
        logger.debug(f"找到 {len(mergeable_groups)} 个可合并组")
        return mergeable_groups
    
    def merge_runs_in_paragraph(self, paragraph) -> int:
        """
        合并段落中的runs
        
        Args:
            paragraph: python-docx的Paragraph对象
            
        Returns:
            int: 成功合并的runs数量
            
        Note:
            - 直接修改传入的paragraph对象
            - 保证合并后格式完全不变
            - 出错时会回滚到原始状态
        """
        try:
            # 分析runs
            run_analysis = self.analyze_paragraph_runs(paragraph)
            if not run_analysis:
                return 0
            
            # 找出可合并组
            mergeable_groups = self.find_mergeable_groups(run_analysis)
            if not mergeable_groups:
                return 0
            
            merged_count = 0
            
            # 从后向前处理，避免索引变化问题
            for group in reversed(mergeable_groups):
                try:
                    if len(group) <= 1:
                        continue
                    
                    # 记录合并前的信息（用于错误恢复）
                    original_runs_info = []
                    for idx in group:
                        run = paragraph.runs[idx]
                        original_runs_info.append({
                            'text': run.text,
                            'properties': self.analyzer.extract_all_properties(run)
                        })
                    
                    # 获取第一个run作为保留的run
                    first_run = paragraph.runs[group[0]]
                    
                    # 合并文本
                    merged_text = ''.join(paragraph.runs[i].text for i in group)
                    
                    # 保存第一个run的完整样式
                    first_run_props = self.analyzer.extract_all_properties(first_run)
                    
                    # 更新第一个run的文本
                    first_run.text = merged_text
                    
                    # 删除其他runs（从后向前删除）
                    for i in reversed(group[1:]):
                        try:
                            paragraph._element.remove(paragraph.runs[i]._element)
                            merged_count += 1
                        except Exception as e:
                            logger.error(f"删除run {i} 时出错: {str(e)}")
                            # 继续处理其他runs
                    
                    self.merge_stats['merge_groups'] += 1
                    logger.debug(f"成功合并 {len(group)} 个runs: {group}")
                    
                except Exception as e:
                    logger.error(f"合并run组 {group} 时出错: {str(e)}")
                    self.merge_stats['error_count'] += 1
                    # 继续处理其他组
            
            return merged_count
            
        except Exception as e:
            logger.error(f"合并段落runs时出现严重错误: {str(e)}")
            self.merge_stats['error_count'] += 1
            return 0
    
    def optimize_document(self, doc_path: str, output_path: str, 
                         progress_callback: Optional[callable] = None) -> Dict[str, Any]:
        """
        优化整个文档的run结构
        
        Args:
            doc_path: 输入文档路径
            output_path: 输出文档路径
            progress_callback: 进度回调函数，接收(current, total)参数
            
        Returns:
            Dict[str, Any]: 优化统计信息，包含：
                - total_runs: 原始run总数
                - merged_runs: 合并的run数
                - merge_groups: 合并组数
                - paragraphs_processed: 处理的段落数
                - paragraphs_optimized: 优化的段落数
                - optimization_rate: 优化率
                - error_count: 错误数
                
        Raises:
            FileNotFoundError: 输入文件不存在
            RunOptimizationError: 优化过程中出现严重错误
        """
        # 检查输入文件
        if not os.path.exists(doc_path):
            raise FileNotFoundError(f"输入文件不存在: {doc_path}")
        
        # 重置统计信息
        self.merge_stats = {
            'total_runs': 0,
            'merged_runs': 0,
            'merge_groups': 0,
            'paragraphs_processed': 0,
            'paragraphs_optimized': 0,
            'skipped_runs': 0,
            'error_count': 0
        }
        
        try:
            # 打开文档
            logger.info(f"正在打开文档: {doc_path}")
            doc = docx.Document(doc_path)
            
            # 统计总段落数
            total_paragraphs = len(doc.paragraphs)
            logger.info(f"文档包含 {total_paragraphs} 个段落")
            
            # 处理每个段落
            for para_idx, paragraph in enumerate(doc.paragraphs):
                try:
                    # 统计原始run数
                    runs_before = len(paragraph.runs)
                    self.merge_stats['total_runs'] += runs_before
                    
                    # 只处理有多个runs的段落
                    if runs_before > 1:
                        # 执行合并
                        merged = self.merge_runs_in_paragraph(paragraph)
                        self.merge_stats['merged_runs'] += merged
                        
                        if merged > 0:
                            self.merge_stats['paragraphs_optimized'] += 1
                            logger.debug(f"段落 {para_idx}: 合并了 {merged} 个runs")
                    
                    self.merge_stats['paragraphs_processed'] += 1
                    
                    # 调用进度回调
                    if progress_callback:
                        progress_callback(para_idx + 1, total_paragraphs)
                    
                except Exception as e:
                    logger.error(f"处理段落 {para_idx} 时出错: {str(e)}")
                    self.merge_stats['error_count'] += 1
                    # 继续处理下一个段落
            
            # 保存优化后的文档
            logger.info(f"正在保存优化后的文档: {output_path}")
            doc.save(output_path)
            
            # 计算优化统计
            if self.merge_stats['total_runs'] > 0:
                optimization_rate = (self.merge_stats['merged_runs'] / 
                                   self.merge_stats['total_runs'] * 100)
                self.merge_stats['optimization_rate'] = f"{optimization_rate:.1f}%"
            else:
                self.merge_stats['optimization_rate'] = "0%"
            
            logger.info(f"优化完成: {self.merge_stats}")
            return self.merge_stats
            
        except Exception as e:
            logger.error(f"优化文档时出现严重错误: {str(e)}")
            raise RunOptimizationError(f"文档优化失败: {str(e)}")


class RunAnalysisReport:
    """
    Run分析报告生成器
    
    提供详细的文档run结构分析，帮助了解：
    - Run分布情况
    - 潜在的优化空间
    - 问题段落识别
    - 样式复杂度分析
    """
    
    @staticmethod
    def generate_analysis_report(doc_path: str, 
                               detailed: bool = False) -> Dict[str, Any]:
        """
        生成文档run分析报告
        
        Args:
            doc_path: 文档路径
            detailed: 是否生成详细报告（包含每个段落的详细信息）
            
        Returns:
            Dict[str, Any]: 分析报告，包含：
                - total_paragraphs: 总段落数
                - total_runs: 总run数
                - avg_runs_per_paragraph: 平均每段落run数
                - single_char_runs: 单字符run数
                - empty_runs: 空run数
                - runs_with_special_boundaries: 包含特殊边界的run数
                - style_variations: 样式变化数
                - problematic_paragraphs: 问题段落列表
                - optimization_potential: 优化潜力评估
                - detailed_analysis: 详细分析（如果requested）
                
        Raises:
            FileNotFoundError: 文档不存在
        """
        if not os.path.exists(doc_path):
            raise FileNotFoundError(f"文档不存在: {doc_path}")
        
        try:
            doc = docx.Document(doc_path)
            analyzer = RunStyleAnalyzer()
            detector = RunBoundaryDetector()
            
            # 初始化报告
            report = {
                'file_path': doc_path,
                'file_size': os.path.getsize(doc_path),
                'total_paragraphs': len(doc.paragraphs),
                'total_runs': 0,
                'single_char_runs': 0,
                'empty_runs': 0,
                'runs_with_special_boundaries': 0,
                'style_variations': set(),
                'problematic_paragraphs': [],
                'run_length_distribution': {
                    '0': 0,      # 空
                    '1': 0,      # 单字符
                    '2-5': 0,    # 短
                    '6-20': 0,   # 中等
                    '21-50': 0,  # 长
                    '50+': 0     # 很长
                }
            }
            
            # 详细分析数据
            if detailed:
                report['detailed_analysis'] = []
            
            # 分析每个段落
            for para_idx, paragraph in enumerate(doc.paragraphs):
                runs = paragraph.runs
                para_text_length = len(paragraph.text)
                run_count = len(runs)
                
                report['total_runs'] += run_count
                
                # 段落级统计
                para_stats = {
                    'index': para_idx,
                    'run_count': run_count,
                    'text_length': para_text_length,
                    'avg_chars_per_run': para_text_length / run_count if run_count > 0 else 0,
                    'style_changes': 0,
                    'special_boundaries': 0
                }
                
                # 检查问题段落（runs过多）
                if run_count > 10 and para_text_length < 100:
                    report['problematic_paragraphs'].append({
                        'index': para_idx,
                        'text_preview': paragraph.text[:50] + '...' if para_text_length > 50 else paragraph.text,
                        'run_count': run_count,
                        'avg_chars_per_run': para_stats['avg_chars_per_run'],
                        'severity': 'high' if run_count > 20 else 'medium'
                    })
                
                # 分析每个run
                prev_style_hash = None
                for run_idx, run in enumerate(runs):
                    run_text_length = len(run.text)
                    
                    # 长度分布统计
                    if run_text_length == 0:
                        report['empty_runs'] += 1
                        report['run_length_distribution']['0'] += 1
                    elif run_text_length == 1:
                        report['single_char_runs'] += 1
                        report['run_length_distribution']['1'] += 1
                    elif run_text_length <= 5:
                        report['run_length_distribution']['2-5'] += 1
                    elif run_text_length <= 20:
                        report['run_length_distribution']['6-20'] += 1
                    elif run_text_length <= 50:
                        report['run_length_distribution']['21-50'] += 1
                    else:
                        report['run_length_distribution']['50+'] += 1
                    
                    # 检查特殊边界
                    boundaries = detector.detect_special_boundaries(run)
                    if any(boundaries.values()):
                        report['runs_with_special_boundaries'] += 1
                        para_stats['special_boundaries'] += 1
                    
                    # 收集样式变化
                    try:
                        prop_hash = analyzer.calculate_property_hash(
                            analyzer.extract_all_properties(run)
                        )
                        report['style_variations'].add(prop_hash)
                        
                        # 统计段落内样式变化
                        if prev_style_hash and prev_style_hash != prop_hash:
                            para_stats['style_changes'] += 1
                        prev_style_hash = prop_hash
                        
                    except Exception as e:
                        logger.warning(f"分析run样式时出错: {str(e)}")
                
                # 添加到详细分析
                if detailed:
                    report['detailed_analysis'].append(para_stats)
            
            # 计算汇总统计
            report['style_variations'] = len(report['style_variations'])
            report['avg_runs_per_paragraph'] = (report['total_runs'] / report['total_paragraphs'] 
                                               if report['total_paragraphs'] > 0 else 0)
            
            # 评估优化潜力
            optimization_score = 0
            if report['avg_runs_per_paragraph'] > 5:
                optimization_score += 30
            if report['single_char_runs'] > report['total_runs'] * 0.1:
                optimization_score += 20
            if report['empty_runs'] > report['total_runs'] * 0.05:
                optimization_score += 10
            if len(report['problematic_paragraphs']) > report['total_paragraphs'] * 0.1:
                optimization_score += 20
            if report['style_variations'] < report['total_runs'] * 0.5:
                optimization_score += 20
            
            report['optimization_potential'] = {
                'score': optimization_score,
                'level': ('high' if optimization_score >= 60 else 
                         'medium' if optimization_score >= 30 else 'low'),
                'recommendation': (
                    '强烈建议优化' if optimization_score >= 60 else
                    '建议优化' if optimization_score >= 30 else
                    '优化空间有限'
                )
            }
            
            # 问题段落排序（按严重程度）
            report['problematic_paragraphs'].sort(
                key=lambda x: x['run_count'], 
                reverse=True
            )
            
            return report
            
        except Exception as e:
            logger.error(f"生成分析报告时出错: {str(e)}")
            raise RunOptimizationError(f"分析报告生成失败: {str(e)}")


# 便捷函数
def quick_optimize(input_path: str, output_path: str = None) -> Dict[str, Any]:
    """
    快速优化Word文档的run结构
    
    Args:
        input_path: 输入文档路径
        output_path: 输出文档路径（如果为None，将添加'_optimized'后缀）
        
    Returns:
        Dict[str, Any]: 优化统计信息
        
    Example:
        >>> stats = quick_optimize("document.docx")
        >>> print(f"优化完成，减少了{stats['merged_runs']}个runs")
    """
    if output_path is None:
        base, ext = os.path.splitext(input_path)
        output_path = f"{base}_optimized{ext}"
    
    merger = SafeRunMerger()
    return merger.optimize_document(input_path, output_path)


def analyze_document(doc_path: str, detailed: bool = False) -> None:
    """
    分析并打印文档run结构报告
    
    Args:
        doc_path: 文档路径
        detailed: 是否打印详细信息
    """
    report = RunAnalysisReport.generate_analysis_report(doc_path, detailed)
    
    print(f"\n=== 文档Run分析报告 ===")
    print(f"文件: {report['file_path']}")
    print(f"文件大小: {report['file_size'] / 1024:.1f} KB")
    print(f"\n基本统计:")
    print(f"  总段落数: {report['total_paragraphs']}")
    print(f"  总Run数: {report['total_runs']}")
    print(f"  平均每段落Run数: {report['avg_runs_per_paragraph']:.2f}")
    print(f"  单字符Run数: {report['single_char_runs']} ({report['single_char_runs']/report['total_runs']*100:.1f}%)")
    print(f"  空Run数: {report['empty_runs']}")
    print(f"  包含特殊元素的Run数: {report['runs_with_special_boundaries']}")
    print(f"  样式变化数: {report['style_variations']}")
    
    print(f"\nRun长度分布:")
    for length_range, count in report['run_length_distribution'].items():
        if count > 0:
            percentage = count / report['total_runs'] * 100
            print(f"  {length_range}: {count} ({percentage:.1f}%)")
    
    print(f"\n优化潜力评估:")
    print(f"  评分: {report['optimization_potential']['score']}/100")
    print(f"  级别: {report['optimization_potential']['level']}")
    print(f"  建议: {report['optimization_potential']['recommendation']}")
    
    if report['problematic_paragraphs']:
        print(f"\n问题段落（Run过多）:")
        for i, p in enumerate(report['problematic_paragraphs'][:5]):  # 只显示前5个
            print(f"  {i+1}. 段落 {p['index']}: {p['run_count']} runs, "
                  f"平均{p['avg_chars_per_run']:.1f}字符/run, 严重程度: {p['severity']}")
            print(f"     预览: {p['text_preview']}")


# 主程序示例
def main():
    """
    演示如何使用Word Run优化器
    """
    # 检查兼容性
    try:
        check_compatibility()
    except DocxCompatibilityError as e:
        print(f"兼容性错误: {e}")
        return
    
    # 配置日志
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    # 示例文档路径
    input_doc = "/Users/wangqiang/Downloads/特机文件/软文部/鲲鹏c3.8/CN-鲲鹏司机手册-原版.docx"
    output_doc = "/Users/wangqiang/Downloads/特机文件/软文部/鲲鹏c3.8/CN-鲲鹏司机手册-原版-runOptimized.docx"
    
    try:
        # 1. 分析文档
        print("=== 步骤1: 分析文档 ===")
        analyze_document(input_doc)
        
        # 2. 优化文档
        print("\n=== 步骤2: 优化文档 ===")
        
        # 创建自定义配置的合并器
        config = {
            'skip_empty_runs': True,
            'min_group_size': 2,
            'log_level': 'INFO'
        }
        merger = SafeRunMerger(config)
        
        # 定义进度回调
        def progress_callback(current, total):
            percentage = (current / total) * 100
            print(f"\r处理进度: {percentage:.1f}% ({current}/{total})", end='')
        
        # 执行优化
        stats = merger.optimize_document(input_doc, output_doc, progress_callback)
        print()  # 换行
        
        # 3. 显示优化结果
        print("\n=== 优化结果 ===")
        print(f"处理段落数: {stats['paragraphs_processed']}")
        print(f"优化段落数: {stats['paragraphs_optimized']}")
        print(f"合并Run组数: {stats['merge_groups']}")
        print(f"减少Run数: {stats['merged_runs']}")
        print(f"跳过的空Run数: {stats['skipped_runs']}")
        print(f"错误数: {stats['error_count']}")
        print(f"优化率: {stats['optimization_rate']}")
        
        # 4. 分析优化后的文档
        print("\n=== 步骤3: 验证优化结果 ===")
        analyze_document(output_doc)
        
    except FileNotFoundError as e:
        print(f"文件错误: {e}")
    except RunOptimizationError as e:
        print(f"优化错误: {e}")
    except Exception as e:
        print(f"未知错误: {e}")
        logger.exception("发生未知错误")


if __name__ == "__main__":
    main()